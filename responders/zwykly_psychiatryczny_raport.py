"""
responders/zwykly_psychiatryczny_raport.py
Moduł obsługujący CAŁY pipeline raportu psychiatrycznego.
"""

import os
import io
import re
import json
import base64
import random
import logging
import requests
import concurrent.futures
from datetime import datetime, timedelta
from flask import current_app

from core.ai_client import call_deepseek, MODEL_TYLER
from core.config import HF_API_URL, HF_STEPS, HF_GUIDANCE, HF_TIMEOUT, MAX_DLUGOSC_EMAIL
from core.logging_reporter import get_logger
from core.hf_token_manager import get_active_tokens, mark_dead

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROMPTS_DIR = os.path.join(BASE_DIR, "prompts")
RAPORT_JSON = os.path.join(PROMPTS_DIR, "zwykly_raport.json")
SUBSTITUTE_IMAGE_PATH = os.path.join(BASE_DIR, "images", "zastepczy.jpg")

# ─────────────────────────────────────────────────────────────────────────────
# PARSOWANIE JSON — pancerne funkcje
# ─────────────────────────────────────────────────────────────────────────────


def _strip_json_markdown(raw: str) -> str:
    """Wyciąga JSON z surowego tekstu (usuwa markdown, przecinki, białe znaki)."""
    if not raw:
        return ""
    raw = raw.strip()
    # Szukamy wszystkiego między pierwszą a ostatnią klamrą/nawiasem
    match = re.search(r"(\{.*\}|\[.*\])", raw, re.DOTALL)
    if match:
        return match.group(1)
    # Jeśli brak klamer, czyścimy z markdownu
    clean = raw.strip().lstrip("`, \n\t,")
    if clean.lower().startswith("json"):
        clean = clean[4:].strip()
    return clean


def _fix_unicode_escapes(raw: str) -> str:
    """Naprawia podwójnie escape'owane sekwencje Unicode (\\\\uXXXX → \\uXXXX).
    NIE używa decode('unicode_escape') — niszczyłoby polskie znaki UTF-8."""
    return re.sub(r"\\\\u([0-9a-fA-F]{4})", r"\\u\1", raw)


def _extract_best_json(raw: str) -> tuple:
    """Wyciąga pierwszy kompletny obiekt JSON z tekstu.

    Używa json.JSONDecoder.raw_decode() który parsuje pierwszy obiekt JSON
    i zwraca (obj, end_position). W przeciwieństwie do json.loads() NIE rzuca
    "Extra data" — po prostu zwraca pierwszy obiekt, a resztę ignoruje.

    Jeśli raw_decode zawiedzie (uszkodzony JSON), szuka największego
    poprawnego fragmentu JSON w tekście.
    """
    decoder = json.JSONDecoder()
    best_obj = None
    best_text = None
    best_len = 0

    # Próba 1: raw_decode od początku — parsuje pierwszy obiekt JSON
    # nawet jeśli po nim są dodatkowe dane ("Extra data")
    try:
        obj, end = decoder.raw_decode(raw)
        return obj, raw[:end]
    except json.JSONDecodeError:
        pass  # Uszkodzony JSON — przejdź do szukania fragmentu

    # Próba 2: szukaj największego obiektu JSON w tekście
    for match in re.finditer(r"[\[{]", raw):
        start = match.start()
        try:
            obj, end = decoder.raw_decode(raw[start:])
            if end > best_len:
                best_len = end
                best_obj = obj
                best_text = raw[start : start + end]
        except json.JSONDecodeError:
            continue
    if isinstance(best_obj, list) and best_obj:
        current_app.logger.warning(
            "[psych-raport] _extract_best_json: root list zamiast obiektu — używam pierwszego elementu"
        )
        best_obj = best_obj[0]
    return best_obj, best_text


def _repair_truncated_json(raw: str) -> str:
    """Naprawia ucięty lub zniszczony JSON."""
    raw = raw.strip()
    if raw.startswith(","):
        raw = "{" + raw.lstrip(",").strip()
    if not raw.endswith("}") and not raw.endswith("]"):
        # Spróbuj dodać brakujące nawiasy
        if "{" in raw:
            raw += "}"
        elif "[" in raw:
            raw += "]"
    return raw


def _strip_trailing_text(raw: str) -> str:
    """Usuwa WSZYSTKO po ostatniej klamrze/nawiasie (JSON + komentarz tekstowy)."""
    # Szukamy ostatniej } lub ]
    last_brace = max(
        raw.rfind("}"),
        raw.rfind("]"),
    )
    if last_brace > 0:
        return raw[: last_brace + 1]
    return raw


def _normalize_json_text(raw: str) -> str:
    """Normalizuje raw response przed parsowaniem."""
    raw = raw.replace("\r\n", "\n")
    # Najpierw usuń tekst POZA JSON ("} KOMENTARZ DLA CIEBIE")
    raw = _strip_trailing_text(raw)
    raw = re.sub(r"//[^\n]*", "", raw)  # Usuń komentarze //
    raw = re.sub(r",\s*([}\]])", r"\1", raw)  # Usuń przecinki przed } i ]
    raw = _fix_unicode_escapes(raw)
    return raw.strip()


# ─────────────────────────────────────────────────────────────────────────────
# PARSOWANIE BEZPIECZNE
# ─────────────────────────────────────────────────────────────────────────────


def _parse_json_safe(raw: str, section: str) -> dict | list | None:
    """Parsuje JSON z fallbackami na każdy poziom."""
    if not raw or len(raw.strip()) < 2:
        return None

    clean = _normalize_json_text(_strip_json_markdown(raw))
    if not clean:
        return None

    # Próba 1: bezpośrednie parsowanie
    try:
        result = json.loads(clean)
        current_app.logger.info("[psych-raport] JSON OK sekcja=%s", section)
        return result
    except json.JSONDecodeError as e:
        current_app.logger.warning(
            "[psych-raport] JSON błąd sekcja=%s: %s — próba ekstrakcji", section, e
        )

    # Próba 2: ekstrakcja największego JSON fragmentu
    extracted, extracted_text = _extract_best_json(clean)
    if extracted is not None:
        if isinstance(extracted, list) and extracted:
            # Przeanalizuj listę — może zawierać dicty
            dict_items = [item for item in extracted if isinstance(item, dict)]
            if dict_items:
                current_app.logger.warning(
                    "[psych-raport] JSON ekstrakcja sekcja=%s zwrócono listę (%d dictów) — scalanie",
                    section,
                    len(dict_items),
                )
                # Scalaj wszystkie dicty w jedną
                merged = {}
                for item in dict_items:
                    merged.update(item)
                if merged:
                    extracted = merged
            else:
                current_app.logger.warning(
                    "[psych-raport] JSON ekstrakcja sekcja=%s zwrócono listę —biorę pierwszy element",
                    section,
                )
                extracted = extracted[0]
        current_app.logger.warning(
            "[psych-raport] JSON ekstrakcja sekcja=%s OK (%d znaków)",
            section,
            len(extracted_text or ""),
        )
        return extracted

    # Próba 3: naprawa i retry
    repaired = _repair_truncated_json(clean)
    try:
        result = json.loads(repaired)
        current_app.logger.warning(
            "[psych-raport] JSON naprawiony sekcja=%s (ucięty output)", section
        )
        return result
    except Exception:
        pass

    # Próba 4: ast.literal_eval — radzi sobie z niektórymi wariantami złego JSON
    try:
        import ast

        result = ast.literal_eval(clean)
        if isinstance(result, (dict, list)):
            current_app.logger.warning(
                "[psych-raport] JSON naprawiony sekcja=%s (ast.literal_eval)", section
            )
            return result
    except Exception:
        pass

    # Fallback: jeśli to czyste słowa (bez JSON), zapakuj w raw text wrapper
    if len(clean) > 10 and "{" not in clean and "[" not in clean:
        current_app.logger.warning(
            "[psych-raport] Sekcja %s: fallback tekstowy (brak JSON)", section
        )
        return {"__raw_text__": clean}

    current_app.logger.error(
        "[psych-raport] JSON nienaprawialny sekcja=%s (raw_len=%d)", section, len(raw)
    )
    return {"__error__": "parse_failed", "raw": raw[:500]}


# ─────────────────────────────────────────────────────────────────────────────
# SEKCJE WYNIKÓW I WALIDACJA
# ─────────────────────────────────────────────────────────────────────────────


def validate_section(data, required_keys):
    """Waliduje czy wymagane pola są wypełnione (nie są puste, None, __BRAK__)."""
    if not isinstance(data, dict):
        return False
    for key in required_keys:
        value = data.get(key)
        if not value or str(value).strip() in ("", "__BRAK__"):
            return False
    return True


def count_empty_fields(data, optional_keys):
    """Liczy ile pól ze sturmientu optional_keys jest puste lub __BRAK__."""
    if not isinstance(data, dict):
        return len(optional_keys)
    empty_count = 0
    for key in optional_keys:
        value = data.get(key)
        if not value or str(value).strip() in ("", "__BRAK__"):
            empty_count += 1
    return empty_count


def _section_result(data, status: str) -> dict:
    return {"data": data, "status": status}


def _unwrap_section(section):
    if isinstance(section, dict) and "data" in section and "status" in section:
        return section["data"]
    return section


def _is_wrapped_section(section):
    return isinstance(section, dict) and (
        "__error__" in section or "__raw_text__" in section
    )


def _section_status(section):
    if isinstance(section, dict) and "status" in section:
        return section["status"]
    return "unknown"


def _wrap_section(raw, section_name: str) -> dict:
    if isinstance(raw, dict) and "data" in raw and "status" in raw:
        return raw
    if raw is None:
        current_app.logger.warning(
            "[psych-raport] sekcja %s: brak wyniku (None)", section_name
        )
        return {"data": {"__error__": "empty_result", "raw": "None"}, "status": "error"}
    if raw == {}:
        current_app.logger.warning(
            "[psych-raport] sekcja %s: pusty dict — traktuję jako błąd", section_name
        )
        return {"data": {"__error__": "empty_result", "raw": "{}"}, "status": "error"}
    if isinstance(raw, dict) and raw.get("__error__"):
        return {"data": raw, "status": "error"}
    if isinstance(raw, dict) and raw.get("__raw_text__"):
        return {"data": raw, "status": "fallback"}
    return {"data": raw, "status": "ok"}


# ─────────────────────────────────────────────────────────────────────────────
# RETRY Z WYŻSZYM MAX_TOKENS
# ─────────────────────────────────────────────────────────────────────────────


def _call_with_retry(system, user, max_tokens=1000):
    """Wywołuje DeepSeek z retry gdy odpowiedź jest podejrzanie krótka."""
    res = call_deepseek(system, user, MODEL_TYLER, max_tokens=max_tokens)
    # Jeśli odpowiedź < 5 znaków lub to same nawiasy — retry ze zwiększonym limitem
    if not res or len(res.strip()) <= 5 or res.strip() in ("{", "}", "[", "]"):
        current_app.logger.warning(
            "[psych-raport] Odpowiedź ucięta/pusta, retry max_tokens=%d", max_tokens * 2
        )
        res = call_deepseek(system, user, MODEL_TYLER, max_tokens=max_tokens * 2)
    return res


# ─────────────────────────────────────────────────────────────────────────────
# WYMUSZANIE JSON STARTU
# ─────────────────────────────────────────────────────────────────────────────

_JSON_FORCE_SUFFIX = "\n\nOdpowiedź TYLKO w formacie JSON. Pierwszym znakiem MUSI być { lub [. Ostatnim znakiem MUSI być } lub ]. Zakaz tekstu poza nawiasami."
_JSON_FORCE_SYSTEM = (
    "KRYTYCZNE: Twoja CAŁKOWITA odpowiedź to WYŁĄCZNIE czysty JSON. "
    "Pierwszym znakiem odpowiedzi MUSI być { lub [. Ostatnim znakiem MUSI być } lub ]. "
    "Absolutny zakaz pisania czegokolwiek przed { — żadnej prozy. "
    "Absolutny zakaz pisania czegokolwiek po } — żadnych notek. "
)


def _u(user_prompt: str) -> str:
    """Wymusza start odpowiedzi od '{' w user promptcie."""
    return user_prompt + _JSON_FORCE_SUFFIX


def _s(system_prompt: str) -> str:
    """Dodaje wymóg startu od '{' do system promptu."""
    if not system_prompt:
        return _JSON_FORCE_SYSTEM
    return system_prompt + "\n" + _JSON_FORCE_SYSTEM


# ─────────────────────────────────────────────────────────────────────────────
# ŁADOWANIE KONFIGURACJI
# ─────────────────────────────────────────────────────────────────────────────


def _load_cfg() -> dict:
    """Wczytuje konfigurację raportu z JSON."""
    try:
        with open(RAPORT_JSON, encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd ładowania cfg: %s", e)
        return {}


def _load_substitute_image() -> dict | None:
    """Ładuje obrazek zastępczy ze ścieżki."""
    # Używamy logging.getLogger zamiast current_app.logger,
    # bo ta funkcja może być wywołana z wątku (ThreadPoolExecutor)
    # gdzie nie ma kontekstu aplikacji Flask.
    log = logging.getLogger(__name__)
    full_path = SUBSTITUTE_IMAGE_PATH
    if not os.path.exists(full_path):
        log.error("[psych-raport] Brak zastepczy.jpg: %s", full_path)
        return None
    try:
        with open(full_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return {
            "base64": b64,
            "content_type": "image/jpeg",
            "filename": "zastepczy.jpg",
        }
    except Exception as e:
        log.error("[psych-raport] Błąd ładowania zastepczy.jpg: %s", e)
        return None


# ─────────────────────────────────────────────────────────────────────────────
# SEKCJE
# ─────────────────────────────────────────────────────────────────────────────


def _sekcja_pacjent(cfg: dict, body: str, sender_name: str) -> dict:
    """Sekcja 1 — dane pacjenta + powód przyjęcia + cytaty.
    Wywołuje deepseek_1_pacjent (pełny schemat) jako główne źródło.
    Jeśli powod_przyjecia lub cytaty_z_przyjecia wychodzą puste,
    robi OSOBNE wywołania deepseek_1b i deepseek_1c jako awaryjne uzupełnienie.
    """
    try:
        # ── Główne wywołanie: deepseek_1_pacjent ─────────────────────────────
        pacjent_cfg = cfg.get("deepseek_1_pacjent", {})
        if not pacjent_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_1_pacjent"
            )
            return _section_result(
                {"__error__": "missing_configuration", "raw": "no deepseek_1_pacjent"},
                "error",
            )

        system = pacjent_cfg.get("system", "")
        schema = pacjent_cfg.get("schema", {})
        instrukcje = pacjent_cfg.get("instrukcje", "")

        user = (
            f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
            f"SENDER_NAME: {sender_name or 'pacjent'}\n\n"
            f"INSTRUKCJE:\n{instrukcje}\n\n"
            f"SCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
        )

        raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja pacjent: brak odpowiedzi AI"
            )
            return _section_result(
                {"__error__": "no_ai_response", "raw": ""},
                "error",
            )

        # Retry gdy wynik zbyt krótki
        MIN_DANE_PACJENTA_LEN = 100
        if len(raw.strip()) < MIN_DANE_PACJENTA_LEN:
            current_app.logger.warning(
                "[psych-raport] dane_pacjenta zbyt krótkie (%d znaków) — retry z max_tokens=5000",
                len(raw.strip()),
            )
            raw2 = call_deepseek(_s(system), _u(user), MODEL_TYLER, max_tokens=5000)
            if raw2 and len(raw2.strip()) > len(raw.strip()):
                raw = raw2

        result = _parse_json_safe(raw, "dane_pacjenta")
        if isinstance(result, dict) and result.get("__error__"):
            current_app.logger.warning(
                "[psych-raport] Sekcja pacjent: JSON nienaprawialny"
            )
            return _section_result(result, "error")

        if isinstance(result, dict) and result.get("__raw_text__"):
            current_app.logger.warning(
                "[psych-raport] Sekcja pacjent: otrzymano surowy tekst zamiast JSON"
            )
            return _section_result(result, "fallback")

        # Walidacja: sprawdź czy mamy wszystkie kluczowe pola (nie tylko powod_przyjecia)
        KLUCZOWE_POLA_PACJENTA = [
            "imie_nazwisko",
            "wiek",
            "adres",
            "zawod",
            "stan_cywilny",
            "powod_przyjecia",
        ]
        empty_count = (
            count_empty_fields(result, KLUCZOWE_POLA_PACJENTA)
            if isinstance(result, dict)
            else len(KLUCZOWE_POLA_PACJENTA)
        )

        if isinstance(result, dict) and empty_count >= 3:  # Zbyt dużo pól puste
            current_app.logger.warning(
                "[psych-raport] Sekcja pacjent: kryterium minimum nie spełnione (%d/%d puste) — retry",
                empty_count,
                len(KLUCZOWE_POLA_PACJENTA),
            )
            # Retry z poprawionym promptem — wymuszajacy pełne dane
            retry_system = (
                system
                + "\n\nKRYTYCZNE: imie_nazwisko, wiek, adres, zawod, stan_cywilny MUSZĄ być wypełnione (NIGDY __BRAK__). "
                + "Jeśli brak danych z emaila — WYMYŚL absurdalnie nawiązując do treści."
            )
            raw_retry = call_deepseek(
                _s(retry_system), _u(user), MODEL_TYLER, max_tokens=5000
            )
            if raw_retry:
                result_retry = _parse_json_safe(raw_retry, "dane_pacjenta")
                if isinstance(result_retry, dict) and not result_retry.get("__error__"):
                    empty_count_retry = count_empty_fields(
                        result_retry, KLUCZOWE_POLA_PACJENTA
                    )
                    if empty_count_retry < empty_count:
                        current_app.logger.info(
                            "[psych-raport] Sekcja pacjent: dane uzupełnione po retry (%d → %d puste)",
                            empty_count,
                            empty_count_retry,
                        )
                        result = result_retry
                    else:
                        current_app.logger.warning(
                            "[psych-raport] Sekcja pacjent: retry nie poprawiło sytuacji"
                        )

        # Obsługa listy zamiast dict — bierz pierwszy element jeśli to dict
        if isinstance(result, list):
            current_app.logger.warning(
                "[psych-raport] dane_pacjenta zwrócono jako list (len=%d) — próba wyciągnięcia dictów",
                len(result),
            )
            dict_items = [item for item in result if isinstance(item, dict)]
            if dict_items:
                merged = {}
                for item in dict_items:
                    merged.update(item)
                result = merged
                current_app.logger.warning(
                    "[psych-raport] dane_pacjenta: list→dict OK (scalono %d dictów)",
                    len(dict_items),
                )
            else:
                current_app.logger.warning(
                    "[psych-raport] dane_pacjenta: list bez dict — ERROR"
                )
                return _section_result(
                    {"__error__": "no_dict_in_list", "raw": str(result)[:500]},
                    "error",
                )

        # Normalizacja kluczy
        if isinstance(result, dict):
            KEY_MAP = {
                "imie": "imie_nazwisko",
                "name": "imie_nazwisko",
                "nazwisko": "imie_nazwisko",
                "wiek": "wiek",
                "age": "wiek",
                "adres": "adres",
                "address": "adres",
                "zawod": "zawod",
                "job": "zawod",
                "stan_cywilny": "stan_cywilny",
                "marital_status": "stan_cywilny",
                "numer_ubezpieczenia": "numer_ubezpieczenia",
                "insurance": "numer_ubezpieczenia",
                "data_przyjecia": "data_przyjecia",
                "admission_date": "data_przyjecia",
                # Obsługa literówki w starym prompcie ('powod_prijecia' bez 'ę')
                "powod_prijecia": "powod_przyjecia",
            }
            for wrong, right in KEY_MAP.items():
                if wrong in result and right not in result:
                    result[right] = result.pop(wrong)
                    current_app.logger.info(
                        "[psych-raport] dane_pacjenta: znormalizowano '%s' → '%s'",
                        wrong,
                        right,
                    )

        # ── Awaryjne wywołanie 1b: powód przyjęcia ────────────────────────────
        # Detekcja: jeśli imie_nazwisko zawiera email (@) — to błąd
        if isinstance(result, dict):
            imie_val = result.get("imie_nazwisko", "")
            if imie_val and "@" in str(imie_val):
                current_app.logger.warning(
                    "[psych-raport] dane_pacjenta: imie_nazwisko to EMAIL (%s) — to jest błąd parsowania!",
                    imie_val,
                )
                # AI wziął sender_name literal bez transformacji — retry z wymaganą transformacją
                retry_system = (
                    system
                    + "\n\nKRYTYCZNE: imie_nazwisko NIE MOŻE być emailem! "
                    + "Jeśli otrzymujesz email jak sender_name, TRANSFORMUJ go w wymyślone imię i nazwisko. "
                    + f"np. zamiast '{sender_name}' napisz coś jak 'Mikolaj Xyz' lub absurdalne imię na bazie emaila."
                )
                raw_retry = call_deepseek(
                    _s(retry_system), _u(user), MODEL_TYLER, max_tokens=5000
                )
                if raw_retry:
                    result_retry = _parse_json_safe(raw_retry, "dane_pacjenta")
                    if isinstance(result_retry, dict) and not result_retry.get(
                        "__error__"
                    ):
                        imie_retry = result_retry.get("imie_nazwisko", "")
                        if imie_retry and "@" not in str(imie_retry):
                            current_app.logger.info(
                                "[psych-raport] dane_pacjenta: imie_nazwisko poprawione po retry"
                            )
                            result = result_retry

        powod_empty = not result.get("powod_przyjecia") or str(
            result.get("powod_przyjecia", "")
        ).strip() in ("", "__BRAK__")
        if powod_empty:
            cfg_1b = cfg.get("deepseek_1b_powod_przyjecia", {})
            if cfg_1b:
                current_app.logger.info(
                    "[psych-raport] powod_przyjecia pusty — wywołuję deepseek_1b"
                )
                try:
                    sys_1b = cfg_1b.get("system", "")
                    ins_1b = cfg_1b.get("instrukcje", "")
                    usr_1b = (
                        f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                        f"SENDER_NAME: {sender_name or 'pacjent'}\n\n"
                        f"INSTRUKCJE:\n{ins_1b}\n\n"
                        f"SCHEMAT JSON:\n{json.dumps(cfg_1b.get('schema', {}), ensure_ascii=False, indent=2)}"
                    )
                    raw_1b = _call_with_retry(_s(sys_1b), _u(usr_1b), max_tokens=3000)
                    if raw_1b:
                        parsed_1b = _parse_json_safe(raw_1b, "powod_przyjecia")
                        if isinstance(parsed_1b, dict):
                            # Obsługa literówki 'powod_prijecia'
                            if (
                                "powod_prijecia" in parsed_1b
                                and "powod_przyjecia" not in parsed_1b
                            ):
                                parsed_1b["powod_przyjecia"] = parsed_1b.pop(
                                    "powod_prijecia"
                                )
                            powod_val = parsed_1b.get(
                                "powod_przyjecia"
                            ) or parsed_1b.get("powod_prijecia")
                            if powod_val:
                                result["powod_przyjecia"] = powod_val
                                current_app.logger.info(
                                    "[psych-raport] deepseek_1b powod_przyjecia OK"
                                )
                except Exception as e1b:
                    current_app.logger.warning(
                        "[psych-raport] deepseek_1b błąd: %s", e1b
                    )

        # ── Awaryjne wywołanie 1c: cytaty z izby przyjęć ─────────────────────
        cytaty_empty = (
            not result.get("cytaty_z_przyjecia")
            or (
                isinstance(result.get("cytaty_z_przyjecia"), list)
                and len(result["cytaty_z_przyjecia"]) == 0
            )
            or str(result.get("cytaty_z_przyjecia", "")).strip() in ("", "__BRAK__")
        )
        if cytaty_empty:
            cfg_1c = cfg.get("deepseek_1c_cytaty", {})
            if cfg_1c:
                current_app.logger.info(
                    "[psych-raport] cytaty_z_przyjecia puste — wywołuję deepseek_1c"
                )
                try:
                    sys_1c = cfg_1c.get("system", "")
                    ins_1c = cfg_1c.get("instrukcje", "")
                    usr_1c = (
                        f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                        f"SENDER_NAME: {sender_name or 'pacjent'}\n\n"
                        f"INSTRUKCJE:\n{ins_1c}\n\n"
                        f"SCHEMAT JSON:\n{json.dumps(cfg_1c.get('schema', {}), ensure_ascii=False, indent=2)}"
                    )
                    raw_1c = _call_with_retry(_s(sys_1c), _u(usr_1c), max_tokens=4000)
                    if raw_1c:
                        parsed_1c = _parse_json_safe(raw_1c, "cytaty_z_przyjecia")
                        if isinstance(parsed_1c, dict):
                            cytaty_val = parsed_1c.get("cytaty_z_przyjecia")
                            if cytaty_val:
                                result["cytaty_z_przyjecia"] = cytaty_val
                                current_app.logger.info(
                                    "[psych-raport] deepseek_1c cytaty OK"
                                )
                except Exception as e1c:
                    current_app.logger.warning(
                        "[psych-raport] deepseek_1c błąd: %s", e1c
                    )

        status = "ok"
        if isinstance(result, dict) and result.get("__raw_text__"):
            status = "fallback"
        current_app.logger.info("[psych-raport] Sekcja pacjent %s", status)
        if isinstance(result, dict):
            return _section_result(result, status)
        return _section_result({"dane_pacjenta": result}, status)

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji pacjent: %s", e)
        return _section_result(
            {"__error__": "exception", "raw": str(e)[:500]},
            "error",
        )


def _sekcja_depozyt_leki(cfg: dict, body: str, nouns_dict: dict) -> dict:
    """Sekcja 2 — depozyt + leki.
    Próbuje najpierw osobnych wywołań deepseek_2a (depozyt) i deepseek_2b (farmakologia).
    Jeśli nie istnieją w konfiguracji, spada na deepseek_2_depozyt_leki (stary klucz).
    """
    nouns_str = (
        ", ".join(list(nouns_dict.values())[:10])
        if nouns_dict
        else "przedmioty codzienne"
    )

    def _normalize_dep(result: dict) -> dict:
        """Normalizuje klucze słownika depozytu/farmakologii."""
        KEY_MAP = {
            "deposit": "depozyt",
            "przedmioty": "lista_przedmiotow",
            "items": "lista_przedmiotow",
            "protokol": "protokol_depozytu",
            "protocol": "protokol_depozytu",
            "leki": "leki",
            "drugs": "leki",
            "medications": "leki",
            "nota_farmaceutyczna": "nota_farmaceutyczna",
            "pharmacy_note": "nota_farmaceutyczna",
        }
        for wrong, right in KEY_MAP.items():
            if wrong in result and right not in result:
                result[right] = result.pop(wrong)
        return result

    merged = {}

    # ── Próba 1: osobne wywołania 2a + 2b ────────────────────────────────────
    cfg_2a = cfg.get("deepseek_2a_depozyt", {})
    cfg_2b = cfg.get("deepseek_2b_farmakologia", {})

    if cfg_2a:
        try:
            ins_2a = cfg_2a.get("instrukcje", "")
            usr_2a = (
                f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                f"RZECZOWNIKI Z EMAILA:\n{nouns_str}\n\n"
                f"INSTRUKCJE:\n{ins_2a}"
            )
            raw_2a = _call_with_retry(
                _s(cfg_2a.get("system", "")), _u(usr_2a), max_tokens=4000
            )
            if raw_2a:
                parsed_2a = _parse_json_safe(raw_2a, "depozyt_2a")
                if isinstance(parsed_2a, dict):
                    parsed_2a = _normalize_dep(parsed_2a)
                    # Wyciągnij podklucz depozyt jeśli zagnieżdżony
                    dep_val = parsed_2a.get("depozyt", parsed_2a)
                    if isinstance(dep_val, dict):
                        merged["depozyt"] = dep_val
                    current_app.logger.info("[psych-raport] deepseek_2a depozyt OK")
        except Exception as e2a:
            current_app.logger.warning("[psych-raport] deepseek_2a błąd: %s", e2a)

    if cfg_2b:
        try:
            ins_2b = cfg_2b.get("instrukcje", "")
            usr_2b = (
                f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                f"RZECZOWNIKI Z EMAILA:\n{nouns_str}\n\n"
                f"INSTRUKCJE:\n{ins_2b}"
            )
            raw_2b = _call_with_retry(
                _s(cfg_2b.get("system", "")), _u(usr_2b), max_tokens=4000
            )
            if raw_2b:
                parsed_2b = _parse_json_safe(raw_2b, "farmakologia_2b")
                if isinstance(parsed_2b, dict):
                    parsed_2b = _normalize_dep(parsed_2b)
                    farm_val = parsed_2b.get("farmakologia", parsed_2b)
                    if isinstance(farm_val, dict):
                        merged["farmakologia"] = farm_val
                    current_app.logger.info(
                        "[psych-raport] deepseek_2b farmakologia OK"
                    )
        except Exception as e2b:
            current_app.logger.warning("[psych-raport] deepseek_2b błąd: %s", e2b)

    # Jeśli udało się zebrać obie sekcje — zwróć
    if merged.get("depozyt") and merged.get("farmakologia"):
        current_app.logger.info("[psych-raport] Sekcja depozyt+leki OK (2a+2b)")
        return merged

    # ── Fallback: stary klucz deepseek_2_depozyt_leki ────────────────────────
    try:
        dep_cfg = cfg.get("deepseek_2_depozyt_leki", {})
        if not dep_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_2_depozyt_leki (fallback)"
            )
            return merged  # zwróć to co mamy z 2a/2b

        system = dep_cfg.get("system", "")
        schema = dep_cfg.get("schema", {})
        instrukcje = dep_cfg.get("instrukcje", "")

        user = (
            f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
            f"RZECZOWNIKI Z EMAILA:\n{nouns_str}\n\n"
            f"INSTRUKCJE:\n{instrukcje}\n\n"
            f"SCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
        )

        raw = _call_with_retry(_s(system), _u(user), max_tokens=6000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja depozyt fallback: brak AI"
            )
            return merged

        result = _parse_json_safe(raw, "depozyt_fallback")
        if result is None:
            return merged

        if isinstance(result, dict):
            result = _normalize_dep(result)
            # Uzupełnij tylko brakujące
            if not merged.get("depozyt") and result.get("depozyt"):
                merged["depozyt"] = result["depozyt"]
            if not merged.get("farmakologia") and result.get("farmakologia"):
                merged["farmakologia"] = result["farmakologia"]
            # Jeśli wynik jest płaski (bez podkluczy depozyt/farmakologia) — przekaż całość
            if not merged.get("depozyt") and not merged.get("farmakologia"):
                merged = result

        current_app.logger.info("[psych-raport] Sekcja depozyt+leki OK (fallback)")
        return merged if isinstance(merged, dict) else {"depozyt": merged}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji depozyt: %s", e)
        return merged


def _sekcja_tydzien(
    cfg: dict, body: str, leki: list, tydzien: int, data_przyjecia: str
) -> list:
    """Sekcja 3/4 — dni hospitalizacji."""
    try:
        tydzien_key = "deepseek_3_tydzien1" if tydzien == 1 else "deepseek_4_tydzien2"
        tydzien_cfg = cfg.get(tydzien_key, {})
        if not tydzien_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji %s", tydzien_key
            )
            return []

        system = tydzien_cfg.get("system", "")
        schema = tydzien_cfg.get("schema", {})
        instrukcje = tydzien_cfg.get("instrukcje", "")

        leki_str = ", ".join(
            [l.get("nazwa", "") for l in (leki or []) if isinstance(l, dict)]
        )[:200]
        user = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\nTYDZIEN: {tydzien}\nLEKI: {leki_str or 'brak danych'}\nDATA PRZYJECIA: {data_przyjecia}\n\nINSTRUKCJE:\n{instrukcje}\n\nSCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"

        raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] %s: brak odpowiedzi AI", tydzien_key
            )
            return []

        result = _parse_json_safe(raw, tydzien_key)
        if result is None:
            current_app.logger.warning(
                "[psych-raport] %s: JSON nienaprawialny", tydzien_key
            )
            return []

        # Normalizacja: oczekujemy listy dni
        if isinstance(result, dict):
            # Może być zapakowane w klucz "dni" lub "hospitalizacja"
            for key in (
                "dni",
                "hospitalizacja",
                "days",
                "records",
                "hospitalizacja_tydzien_1",
                "hospitalizacja_tydzien_2",
            ):
                if key in result and isinstance(result[key], list):
                    result = result[key]
                    break
            else:
                # Jeśli dict ale nie ma listy, spróbuj wyciągnąć wartości
                for v in result.values():
                    if isinstance(v, list):
                        result = v
                        break

        if not isinstance(result, list):
            current_app.logger.warning(
                "[psych-raport] %s: oczekiwano listy, dostałem %s — wartość: %.100s",
                tydzien_key,
                type(result).__name__,
                result,
            )
            return []

        # Normalizacja kluczy w każdym dniu
        DAY_KEY_MAP = {
            "day": "dzien",
            "date": "data",
            "event": "zdarzenie",
            "drug": "lek",
            "medication": "lek",
            "condition": "stan_pacjenta",
            "state": "stan_pacjenta",
            "doctor_note": "nota_lekarska",
            "note": "nota_lekarska",
        }
        for d in result:
            if isinstance(d, dict):
                for wrong, right in DAY_KEY_MAP.items():
                    if wrong in d and right not in d:
                        d[right] = d.pop(wrong)

        current_app.logger.info(
            "[psych-raport] %s OK (%d dni)", tydzien_key, len(result)
        )
        return result

    except Exception as e:
        current_app.logger.error(
            "[psych-raport] Błąd %s: %s",
            tydzien_key if "tydzien_key" in dir() else "tydzien",
            e,
        )
        return []


def _sekcja_wypis(cfg: dict, body: str, data_przyjecia: str) -> dict:
    """Sekcja 5 — wypis."""
    try:
        wypis_cfg = cfg.get("deepseek_5_wypis", {})
        if not wypis_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_5_wypis"
            )
            return _section_result(
                {"__error__": "missing_configuration", "raw": "no deepseek_5_wypis"},
                "error",
            )

        system = wypis_cfg.get("system", "")
        schema = wypis_cfg.get("schema", {})
        instrukcje = wypis_cfg.get("instrukcje", "")

        user = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\nDATA PRZYJECIA: {data_przyjecia}\n\nINSTRUKCJE:\n{instrukcje}\n\nSCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"

        raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja wypis: brak odpowiedzi AI"
            )
            return _section_result({"__error__": "no_ai_response", "raw": ""}, "error")

        result = _parse_json_safe(raw, "wypis")
        if result is None:
            current_app.logger.warning(
                "[psych-raport] Sekcja wypis: JSON nienaprawialny"
            )
            return _section_result(
                {"__error__": "parse_failed", "raw": raw[:500]}, "error"
            )

        # Normalizacja: oczekujemy zagnieżdżenia w kluczu "wypis"
        if isinstance(result, dict) and "wypis" in result:
            result = result["wypis"]

        # Normalizacja kluczy
        if isinstance(result, dict):
            KEY_MAP = {
                "discharge_day": "dzien_wypisu",
                "dzien": "dzien_wypisu",
                "discharge_reason": "powod_wypisu",
                "powod": "powod_wypisu",
                "reason": "powod_wypisu",
                "discharge_condition": "stan_przy_wypisie",
                "stan": "stan_przy_wypisie",
                "condition": "stan_przy_wypisie",
                "post_discharge_recommendations": "zalecenia_po_wypisie",
                "zalecenia": "zalecenia_po_wypisie",
                "recommendations": "zalecenia_po_wypisie",
                "farewell": "opis_pozegnania",
                "pozegnanie": "opis_pozegnania",
            }
            for wrong, right in KEY_MAP.items():
                if wrong in result and right not in result:
                    result[right] = result.pop(wrong)
                    current_app.logger.info(
                        "[psych-raport] wypis: znormalizowano '%s' → '%s'", wrong, right
                    )

        current_app.logger.info("[psych-raport] Sekcja wypis OK")
        return result if isinstance(result, dict) else {"wypis": result}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji wypis: %s", e)
        return _section_result({"__error__": "exception", "raw": str(e)[:500]}, "error")


def _sekcja_diagnozy(cfg: dict, body: str, previous_body: str) -> dict:
    """Sekcja 6 — diagnozy łacińskie + objawy.
    Próbuje najpierw deepseek_6a (diagnozy główne) i deepseek_6b (objawy).
    Jeśli nie istnieją, spada na deepseek_6_diagnozy_lacina (stary klucz).
    """
    historia = (previous_body or "brak historii choroby")[:MAX_DLUGOSC_EMAIL]
    merged = {}

    # ── Próba 1: osobne wywołania 6a + 6b ────────────────────────────────────
    cfg_6a = cfg.get("deepseek_6a_diagnozy_glowne", {})
    cfg_6b = cfg.get("deepseek_6b_objawy", {})

    if cfg_6a:
        try:
            ins_6a = cfg_6a.get("instrukcje", "")
            usr_6a = (
                f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                f"HISTORIA CHOROBY:\n{historia}\n\n"
                f"INSTRUKCJE:\n{ins_6a}"
            )
            raw_6a = _call_with_retry(
                _s(cfg_6a.get("system", "")), _u(usr_6a), max_tokens=4000
            )
            if raw_6a:
                parsed_6a = _parse_json_safe(raw_6a, "diagnozy_6a")
                if isinstance(parsed_6a, dict) and not parsed_6a.get("__error__"):
                    merged.update(parsed_6a)
                    current_app.logger.info("[psych-raport] deepseek_6a diagnozy OK")
        except Exception as e6a:
            current_app.logger.warning("[psych-raport] deepseek_6a błąd: %s", e6a)

    if cfg_6b:
        try:
            ins_6b = cfg_6b.get("instrukcje", "")
            usr_6b = (
                f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                f"INSTRUKCJE:\n{ins_6b}"
            )
            raw_6b = _call_with_retry(
                _s(cfg_6b.get("system", "")), _u(usr_6b), max_tokens=3000
            )
            if raw_6b:
                parsed_6b = _parse_json_safe(raw_6b, "objawy_6b")
                if isinstance(parsed_6b, dict) and not parsed_6b.get("__error__"):
                    # Uzupełnij merged tylko kluczami których jeszcze nie ma
                    for k, v in parsed_6b.items():
                        if k not in merged:
                            merged[k] = v
                    current_app.logger.info("[psych-raport] deepseek_6b objawy OK")
        except Exception as e6b:
            current_app.logger.warning("[psych-raport] deepseek_6b błąd: %s", e6b)

    # Jeśli mamy cokolwiek — sprawdź czy wystarczające
    if merged.get("diagnoza_wstepna"):
        # Normalizacja kluczy
        KEY_MAP = {
            "primary_diagnosis": "diagnoza_wstepna",
            "diagnoza": "diagnoza_wstepna",
            "diagnosis": "diagnoza_wstepna",
            "rozpoznanie": "diagnoza_wstepna",
            "additional_diagnosis": "diagnoza_dodatkowa",
            "diagnoza_dod": "diagnoza_dodatkowa",
            "comorbidity": "choroba_wspolistniejaca",
            "choroba_wspol": "choroba_wspolistniejaca",
            "symptoms": "objawy",
            "symptomy": "objawy",
            "objaw": "objawy",
        }
        for wrong, right in KEY_MAP.items():
            if wrong in merged and right not in merged:
                merged[right] = merged.pop(wrong)
        current_app.logger.info("[psych-raport] Sekcja diagnozy OK (6a+6b)")
        return merged

    # ── Fallback: stary klucz deepseek_6_diagnozy_lacina ────────────────────────
    if not merged.get("diagnoza_wstepna"):
        try:
            diagnozy_cfg = cfg.get("deepseek_6_diagnozy_lacina", {})
            if not diagnozy_cfg:
                current_app.logger.warning(
                    "[psych-raport] Brak konfiguracji deepseek_6_diagnozy_lacina (fallback)"
                )
                return merged

            system = diagnozy_cfg.get("system", "")
            schema = diagnozy_cfg.get("schema", {})
            instrukcje = diagnozy_cfg.get("instrukcje", "")

            user = (
                f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
                f"HISTORIA CHOROBY:\n{historia}\n\n"
                f"INSTRUKCJE:\n{instrukcje}\n\n"
                f"SCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
            )

            raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
            if not raw:
                current_app.logger.warning(
                    "[psych-raport] Sekcja diagnozy fallback: brak AI"
                )
                return merged

            result = _parse_json_safe(raw, "diagnozy_fallback")
            if result is None:
                return merged

            if isinstance(result, dict):
                KEY_MAP = {
                    "primary_diagnosis": "diagnoza_wstepna",
                    "diagnoza": "diagnoza_wstepna",
                    "diagnosis": "diagnoza_wstepna",
                    "rozpoznanie": "diagnoza_wstepna",
                    "additional_diagnosis": "diagnoza_dodatkowa",
                    "diagnoza_dod": "diagnoza_dodatkowa",
                    "comorbidity": "choroba_wspolistniejaca",
                    "choroba_wspol": "choroba_wspolistniejaca",
                    "symptoms": "objawy",
                    "symptomy": "objawy",
                    "objaw": "objawy",
                }
                for wrong, right in KEY_MAP.items():
                    if wrong in result and right not in result:
                        result[right] = result.pop(wrong)
                # Uzupełnij merged brakującymi polami
                for k, v in result.items():
                    if k not in merged:
                        merged[k] = v

            current_app.logger.info("[psych-raport] Sekcja diagnozy OK (fallback)")
            return merged if isinstance(merged, dict) else {"diagnoza_wstepna": merged}

        except Exception as e:
            current_app.logger.error("[psych-raport] Błąd sekcji diagnozy: %s", e)
            return merged
    else:
        current_app.logger.info("[psych-raport] Sekcja diagnozy OK (6a+6b)")
        return merged


def _sekcja_zalecenia(cfg: dict, body: str, dni_1_7: list, dni_8_14: list) -> dict:
    """Sekcja 7 — zalecenia + notatki (5 osobnych DeepSeek)."""
    try:
        zalecenia_cfg = cfg.get("deepseek_7_zalecenia_notatki", {})
        if not zalecenia_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_7_zalecenia_notatki"
            )
            return _section_result(
                {
                    "__error__": "missing_configuration",
                    "raw": "no deepseek_7_zalecenia_notatki",
                },
                "error",
            )

        system = zalecenia_cfg.get("system", "")
        schema = zalecenia_cfg.get("schema", {})
        instrukcje = zalecenia_cfg.get("instrukcje", "")

        # Podsumowanie hospitalizacji dla kontekstu
        dni_str = ""
        for d in (dni_1_7 or [])[:3] + (dni_8_14 or [])[:3]:
            if isinstance(d, dict):
                dni_str += (
                    f"Dzień {d.get('dzien', '?')}: {d.get('zdarzenie', '')[:100]}\n"
                )

        user = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\nPRZEBIEG HOSPITALIZACJI:\n{dni_str[:500] or 'brak danych'}\n\nINSTRUKCJE:\n{instrukcje}\n\nSCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"

        raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja zalecenia: brak odpowiedzi AI"
            )
            return _section_result({"__error__": "no_ai_response", "raw": ""}, "error")

        result = _parse_json_safe(raw, "zalecenia")
        if result is None:
            current_app.logger.warning(
                "[psych-raport] Sekcja zalecenia: JSON nienaprawialny"
            )
            return _section_result(
                {"__error__": "parse_failed", "raw": raw[:500]}, "error"
            )

        # Normalizacja kluczy
        if isinstance(result, dict):
            KEY_MAP = {
                "tyler_recommendations": "zalecenia_tylera",
                "recommendations": "zalecenia_tylera",
                "tyler_zadania": "zalecenia_tylera",
                "naglowek": "naglowek",
                "header": "naglowek",
                "zadanie": "zadanie_1",
                "task_1": "zadanie_1",
                "zadanie_1": "zadanie_1",
                "podpis": "podpis",
                "signature": "podpis",
                "rokowanie": "rokowanie",
                "prognosis": "rokowanie",
                "incydenty": "incydenty_specjalne",
                "incidents": "incydenty_specjalne",
                "special_incidents": "incydenty_specjalne",
                "nurse_notes": "notatki_pielegniarek",
                "notatki_pielegniarek": "notatki_pielegniarek",
                "cleaner_notes": "notatki_sprzataczki",
                "notatki_sprzataczki": "notatki_sprzataczki",
            }
            for wrong, right in KEY_MAP.items():
                if wrong in result and right not in result:
                    result[right] = result.pop(wrong)
                    current_app.logger.info(
                        "[psych-raport] zalecenia: znormalizowano '%s' → '%s'",
                        wrong,
                        right,
                    )

            # Naprawa: AI czasem zwraca zalecenia_tylera jako listę notatek pielęgniarek
            # zamiast dict {naglowek, zadanie_1, zadanie_2, zadanie_3, podpis}
            zt = result.get("zalecenia_tylera")
            if isinstance(zt, list):
                current_app.logger.warning(
                    "[psych-raport] zalecenia_tylera jest listą (%d el.) — próba rekonstrukcji dict",
                    len(zt),
                )
                # Sprawdź czy to lista notatek pielęgniarek zamiast zadań
                if zt and isinstance(zt[0], dict) and "imie_pielegniarki" in zt[0]:
                    # AI pomyliło klucze — przenieś listę do notatek_pielegniarek
                    if not result.get("notatki_pielegniarek"):
                        result["notatki_pielegniarek"] = zt
                    result["zalecenia_tylera"] = {}
                    current_app.logger.warning(
                        "[psych-raport] zalecenia_tylera → przeniesiono do notatki_pielegniarek"
                    )
                else:
                    # Lista zadań — przebuduj jako dict
                    rebuilt = {}
                    for i, item in enumerate(zt[:3], 1):
                        if isinstance(item, dict):
                            task = (
                                item.get("zadanie")
                                or item.get("tresc")
                                or item.get(f"zadanie_{i}", "")
                            )
                        else:
                            task = str(item)
                        rebuilt[f"zadanie_{i}"] = task
                    result["zalecenia_tylera"] = rebuilt

        current_app.logger.info("[psych-raport] Sekcja zalecenia OK")
        return result if isinstance(result, dict) else {"zalecenia_tylera": result}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji zalecenia: %s", e)
        return {}


def _sekcja_flux_prompty(
    cfg: dict,
    body: str,
    nouns_dict: dict,
    sender_name: str,
    gender: str,
    test_mode: bool = False,
) -> dict:
    """Sekcja 8 — prompty FLUX."""
    try:
        flux_cfg = cfg.get("deepseek_8_flux_prompty", {})
        if not flux_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_8_flux_prompty"
            )
            return {}

        system = flux_cfg.get("system", "")
        schema = flux_cfg.get("schema", {})
        instrukcje = flux_cfg.get("instrukcje", "")

        nouns_str = (
            ", ".join(list(nouns_dict.values())[:8])
            if nouns_dict
            else "przedmioty codzienne"
        )
        user = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\nRZECZOWNIKI Z EMAILA:\n{nouns_str}\n\nSENDER_NAME: {sender_name or 'pacjent'}\nGENDER: {gender or 'patient'}\n\nINSTRUKCJE:\n{instrukcje}\n\nSCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"

        raw = _call_with_retry(_s(system), _u(user), max_tokens=1500)
        if not raw:
            current_app.logger.warning("[psych-raport] Sekcja flux: brak odpowiedzi AI")
            return {}

        result = _parse_json_safe(raw, "flux_prompty")
        if result is None:
            current_app.logger.warning(
                "[psych-raport] Sekcja flux: JSON nienaprawialny"
            )
            return {}

        # Normalizacja kluczy
        if isinstance(result, dict):
            KEY_MAP = {
                "pacjent_prompt": "prompt_pacjent",
                "patient_prompt": "prompt_pacjent",
                "prompt_pacjenta": "prompt_pacjent",
                "przedmioty_prompt": "prompt_przedmioty",
                "objects_prompt": "prompt_przedmioty",
                "prompt_przedmiotow": "prompt_przedmioty",
            }
            for wrong, right in KEY_MAP.items():
                if wrong in result and right not in result:
                    result[right] = result.pop(wrong)
                    current_app.logger.info(
                        "[psych-raport] flux: znormalizowano '%s' → '%s'", wrong, right
                    )

        current_app.logger.info("[psych-raport] Sekcja flux OK")
        return result if isinstance(result, dict) else {"prompt_pacjent": str(result)}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji flux: %s", e)
        return {}


def _sekcja_leczenie_specjalne(
    cfg: dict, body: str, dni_1_7: list, dni_8_14: list
) -> dict:
    """Sekcja 9 — leczenie specjalne (deepseek_9_leczenie_specjalne).
    Zwraca dict z kluczem 'leczenie_specjalne' zawierającym listę metod leczenia.
    """
    try:
        leczenie_cfg = cfg.get("deepseek_9_leczenie_specjalne", {})
        if not leczenie_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_9_leczenie_specjalne"
            )
            return {}

        system = leczenie_cfg.get("system", "")
        schema = leczenie_cfg.get("schema", {})
        instrukcje = leczenie_cfg.get("instrukcje", "")

        # Podsumowanie przebiegu leczenia dla kontekstu
        dni_str = ""
        for d in (dni_1_7 or [])[:4] + (dni_8_14 or [])[:4]:
            if isinstance(d, dict):
                dni_str += (
                    f"Dzień {d.get('dzien', '?')}: {d.get('zdarzenie', '')[:100]}\n"
                )

        user = (
            f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
            f"PRZEBIEG HOSPITALIZACJI:\n{dni_str[:600] or 'brak danych'}\n\n"
            f"INSTRUKCJE:\n{instrukcje}\n\n"
            f"SCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
        )

        raw = _call_with_retry(_s(system), _u(user), max_tokens=3000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja leczenie_specjalne: brak odpowiedzi AI"
            )
            return {}

        result = _parse_json_safe(raw, "leczenie_specjalne")
        if result is None:
            current_app.logger.warning(
                "[psych-raport] Sekcja leczenie_specjalne: JSON nienaprawialny"
            )
            return {}

        # Normalizacja: oczekujemy listy metod lub dict z kluczem leczenie_specjalne
        if isinstance(result, dict):
            for key in (
                "leczenie_specjalne",
                "metody_leczenia",
                "treatments",
                "methods",
            ):
                if key in result:
                    val = result[key]
                    if isinstance(val, list):
                        current_app.logger.info(
                            "[psych-raport] leczenie_specjalne: wyciągnięto z klucza '%s' (%d el.)",
                            key,
                            len(val),
                        )
                        return {"leczenie_specjalne": val}
            # Jeśli dict ale bez listy — zwróć jako słownik metod
            current_app.logger.info(
                "[psych-raport] leczenie_specjalne: dict → zachowano jako dict"
            )
            return {"leczenie_specjalne": result}

        if isinstance(result, list):
            current_app.logger.info(
                "[psych-raport] leczenie_specjalne OK (%d metod)", len(result)
            )
            return {"leczenie_specjalne": result}

        current_app.logger.warning(
            "[psych-raport] leczenie_specjalne: nieoczekiwany typ %s",
            type(result).__name__,
        )
        return {}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji leczenie_specjalne: %s", e)
        return {}


def _sekcja_relacje_swiadkow(cfg: dict, body: str, raport: dict) -> dict:
    """Sekcja relacji świadków (DeepSeek)."""
    try:
        swiadkowie_cfg = cfg.get("deepseek_3_relacje_swiadkow", {})
        if not swiadkowie_cfg:
            current_app.logger.warning(
                "[psych-raport] Brak konfiguracji deepseek_3_relacje_swiadkow"
            )
            return {"relacje_swiadkow": []}

        system = swiadkowie_cfg.get("system", "")
        schema = swiadkowie_cfg.get("schema", {})
        instrukcje = swiadkowie_cfg.get("instrukcje", "")

        # Kontekst z raportu — zabezpieczenie na wypadek gdyby raport nie był dict
        if not isinstance(raport, dict):
            current_app.logger.warning(
                "[psych-raport] raport nie jest dict (type=%s) — używam pustego",
                type(raport).__name__,
            )
            raport = {}
        dane_pacjenta = raport.get("dane_pacjenta", {})
        if not isinstance(dane_pacjenta, dict):
            current_app.logger.warning(
                "[psych-raport] relacje_swiadkow: dane_pacjenta zły typ: %s — wartość: %.100s",
                type(dane_pacjenta).__name__,
                dane_pacjenta,
            )
            dane_pacjenta = {}
        pacjent = dane_pacjenta.get("imie_nazwisko", "pacjent")
        diagnoza = raport.get("diagnoza_wstepna", {})
        if isinstance(diagnoza, dict):
            diagnoza_str = diagnoza.get("nazwa_lacinska", "") or diagnoza.get(
                "nazwa_polska", ""
            )
        else:
            diagnoza_str = str(diagnoza)

        user = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}\n\nPACJENT: {pacjent}\nDIAGNOZA: {diagnoza_str or 'nieznana'}\n\nINSTRUKCJE:\n{instrukcje}\n\nSCHEMAT JSON:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"

        raw = _call_with_retry(_s(system), _u(user), max_tokens=4000)
        if not raw:
            current_app.logger.warning(
                "[psych-raport] Sekcja świadkowie: brak odpowiedzi AI"
            )
            return {"relacje_swiadkow": []}

        result = _parse_json_safe(raw, "relacje_swiadkow")
        if result is None:
            current_app.logger.warning(
                "[psych-raport] Sekcja świadkowie: JSON nienaprawialny"
            )
            return {"relacje_swiadkow": []}

        # Normalizacja: oczekujemy listy świadków
        if isinstance(result, dict):
            for key in (
                "relacje_swiadkow",
                "swiadkowie",
                "witnesses",
                "witness_statements",
            ):
                if key in result and isinstance(result[key], list):
                    result = result[key]
                    break
            else:
                # Jeśli dict ale nie ma listy, spróbuj wyciągnąć wartości
                for v in result.values():
                    if isinstance(v, list):
                        result = v
                        break

        if not isinstance(result, list):
            current_app.logger.warning(
                "[psych-raport] świadkowie: oczekiwano listy, dostałem %s — wartość: %.100s",
                type(result).__name__,
                result,
            )
            return {"relacje_swiadkow": []}

        # Normalizacja kluczy w każdym świadku
        WITNESS_KEY_MAP = {
            "name": "imie_swiadka",
            "imie": "imie_swiadka",
            "witness_name": "imie_swiadka",
            "occupation": "zawod",
            "job": "zawod",
            "date": "data",
            "statement": "tresc",
            "testimony": "tresc",
            "zeznanie": "tresc",
            "content": "tresc",
        }
        for w in result:
            if isinstance(w, dict):
                for wrong, right in WITNESS_KEY_MAP.items():
                    if wrong in w and right not in w:
                        w[right] = w.pop(wrong)

        current_app.logger.info(
            "[psych-raport] Sekcja świadkowie OK (%d relacji)", len(result)
        )
        return {"relacje_swiadkow": result}

    except Exception as e:
        current_app.logger.error("[psych-raport] Błąd sekcji świadkowie: %s", e)
        return {"relacje_swiadkow": []}


def _hf_credit_exhausted(resp) -> bool:
    """Sprawdza czy odpowiedź 402 wskazuje na globalne wyczerpanie kredytów."""
    try:
        data = resp.json()
        return data.get("error", "").find("exhausted") != -1
    except Exception:
        return False


def _substitute_or_none(label: str) -> str | None:
    """Zwraca obrazek zastępczy lub None."""
    substitute = _load_substitute_image()
    if substitute:
        return substitute.get("base64")
    return None


def _generate_flux(
    prompt: str,
    label: str,
    steps: int = 28,
    guidance: float = 7.0,
    width: int = 1024,
    height: int = 1024,
    test_mode: bool = False,
) -> str | None:
    """Generuje obrazek FLUX — zwraca base64 JPG lub None."""
    # Używamy logging.getLogger zamiast current_app.logger,
    # bo ta funkcja może być wywołana z wątku (ThreadPoolExecutor)
    # gdzie nie ma kontekstu aplikacji Flask.
    log = logging.getLogger(__name__)

    if test_mode:
        substitute = _load_substitute_image()
        if substitute:
            return substitute.get("base64")
        return None

    tokens = get_active_tokens()
    if not tokens:
        log.error("[psych-flux] Brak tokenów HF dla %s", label)
        substitute = _load_substitute_image()
        return substitute.get("base64") if substitute else None

    log.info("[psych-flux] %s — prompt %.120s...", label, prompt)

    payload = {
        "inputs": prompt,
        "parameters": {
            "num_inference_steps": steps,
            "guidance_scale": guidance,
            "width": width,
            "height": height,
            "seed": random.randint(0, 2**32 - 1),
        },
    }

    for name, token in tokens:
        headers = {"Authorization": f"Bearer {token}", "Accept": "image/png"}
        try:
            resp = requests.post(
                HF_API_URL, headers=headers, json=payload, timeout=HF_TIMEOUT
            )
            if resp.status_code == 200:
                log.info(
                    "[psych-flux] %s OK token=%s (%dB)", label, name, len(resp.content)
                )
                try:
                    from PIL import Image as PILImage

                    pil = PILImage.open(io.BytesIO(resp.content)).convert("RGB")
                    buf = io.BytesIO()
                    pil.save(buf, format="JPEG", quality=92, optimize=True)
                    return base64.b64encode(buf.getvalue()).decode("ascii")
                except Exception as e:
                    log.warning("[psych-flux] PNG→JPG błąd: %s", e)
                    return base64.b64encode(resp.content).decode("ascii")
            elif resp.status_code == 402:
                mark_dead(name)
                log.warning(
                    "[psych-flux] %s 402 token=%s — wyczerpane kredyty, dodano do czarnej listy",
                    label,
                    name,
                )
                if _hf_credit_exhausted(resp):
                    log.warning(
                        "[psych-flux] %s 402 wskazuje na globalne wyczerpanie kredytów — kończę próby",
                        label,
                    )
                    break
            elif resp.status_code in (401, 403):
                mark_dead(name)
                log.warning(
                    "[psych-flux] %s HTTP %d token=%s — nieważny, dodano do czarnej listy",
                    label,
                    resp.status_code,
                    name,
                )
            elif resp.status_code == 429:
                log.warning("[psych-flux] %s 429 token=%s → następny", label, name)
            else:
                log.warning(
                    "[psych-flux] %s HTTP %d token=%s", label, resp.status_code, name
                )
        except Exception as e:
            log.warning("[psych-flux] %s wyjątek token=%s: %s", label, name, e)

    log.error(
        "[psych-flux] %s — wszystkie tokeny zawiodły — używam zastepczy.jpg", label
    )
    return _substitute_or_none(label)


def _generate_photos_parallel(
    prompt_pacjent: str, prompt_przedmioty: str, test_mode: bool = False
) -> tuple:
    """Generuje oba zdjęcia równolegle — zwraca (photo_1_dict, photo_2_dict)."""
    # Używamy logging.getLogger zamiast current_app.logger,
    # bo ta funkcja może być wywołana z wątku (ThreadPoolExecutor)
    # gdzie nie ma kontekstu aplikacji Flask.
    log = logging.getLogger(__name__)

    def gen_pacjent():
        b64 = _generate_flux(prompt_pacjent, "photo_pacjent", test_mode=test_mode)
        if b64:
            return {
                "base64": b64,
                "content_type": "image/jpeg",
                "filename": f"psych_pacjent_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg",
            }
        return None

    def gen_przedmioty():
        b64 = _generate_flux(prompt_przedmioty, "photo_przedmioty", test_mode=test_mode)
        if b64:
            return {
                "base64": b64,
                "content_type": "image/jpeg",
                "filename": f"psych_przedmioty_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg",
            }
        return None

    # Szybkie wyjście gdy brak aktywnych tokenów — nie czekaj 5 minut
    if not get_active_tokens():
        log.warning(
            "[psych-flux] Brak aktywnych tokenów — pomijam FLUX, używam zastepczy.jpg"
        )
        sub = _load_substitute_image()
        sub_dict = None
        if sub:
            sub_dict = {
                "base64": sub.get("base64"),
                "content_type": "image/jpeg",
                "filename": f"zastepczy_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jpg",
            }
        return sub_dict, sub_dict

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            fut1 = executor.submit(gen_pacjent)
            fut2 = executor.submit(gen_przedmioty)
            p1 = fut1.result(timeout=120)
            p2 = fut2.result(timeout=120)
            return p1, p2
    except Exception as e:
        log.error("[psych-flux] Błąd równoległy: %s", e)
        return None, None


def _build_docx(
    raport: dict,
    photo_pacjent_b64: str | None,
    photo_przedmioty_b64: str | None,
    cfg: dict,
) -> str | None:
    """Buduje DOCX z całym raportem — zwraca base64 lub None."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        current_app.logger.error("[psych-docx] Brak python-docx")
        return None

    try:
        return _build_docx_inner(raport, photo_pacjent_b64, photo_przedmioty_b64, cfg)
    except Exception as e:
        current_app.logger.error(
            "[psych-docx] Nieoczekiwany błąd budowania DOCX: %s", e, exc_info=True
        )
        return None


def _build_docx_inner(
    raport: dict,
    photo_pacjent_b64: str | None,
    photo_przedmioty_b64: str | None,
    cfg: dict,
) -> str | None:
    """Wewnętrzna implementacja — wywoływana przez _build_docx z ochronnym try/except."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    szpital = cfg.get("szpital", {})
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Dodaj stopkę z numerami stron (XML — python-docx nie ma add_field na Run)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    footer = section.footer
    footer_paragraph = (
        footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    )
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.clear()

    def _add_page_field(paragraph, field_name):
        """Wstawia pole Word (PAGE lub NUMPAGES) do akapitu przez XML."""
        run = paragraph.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " " + field_name + " "
        run._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "separate")
        run._r.append(fld2)
        fld3 = OxmlElement("w:fldChar")
        fld3.set(qn("w:fldCharType"), "end")
        run._r.append(fld3)

    _add_page_field(footer_paragraph, "PAGE")
    sep_run = footer_paragraph.add_run(" / ")
    sep_run.font.size = Pt(9)
    sep_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _add_page_field(footer_paragraph, "NUMPAGES")

    RED = RGBColor(0x99, 0x1A, 0x1A)
    DARK = RGBColor(0x0D, 0x0D, 0x0D)
    GREY = RGBColor(0x66, 0x66, 0x66)
    LGREY = RGBColor(0x99, 0x99, 0x99)

    def heading(text, level=1, color=DARK, size=14):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.size = Pt(size)
            run.font.color.rgb = color
        return h

    def para(text, bold=False, italic=False, color=DARK, size=10, align=None):
        p = doc.add_paragraph()
        if align:
            p.alignment = align
        r = p.add_run(str(text))
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color
        return p

    def _docx_unwrap(value):
        if isinstance(value, dict):
            if "__error__" in value:
                return "[BŁĄD GENEROWANIA]"
            if "__raw_text__" in value:
                return value["__raw_text__"]
        return value

    def _docx_render_error_or_raw(value):
        if isinstance(value, dict):
            if "__error__" in value:
                para(
                    "[BŁĄD GENEROWANIA SEKCJI]",
                    italic=True,
                    color=LGREY,
                    size=9,
                )
                return True
            if "__raw_text__" in value:
                para(value["__raw_text__"], italic=True, color=GREY, size=9)
                return True
        return False

    def field(label, value, label_color=DARK, val_color=DARK, size=10):
        value = _docx_unwrap(value)
        if value in (None, "", [], {}, "__BRAK__"):
            value = "[brak danych]"
            val_color = LGREY
        p = doc.add_paragraph()
        rl = p.add_run(f"{label}: ")
        rl.bold = True
        rl.font.size = Pt(size)
        rl.font.color.rgb = label_color
        rv = p.add_run(str(value))
        rv.font.size = Pt(size)
        rv.font.color.rgb = val_color

    def separator():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run("─" * 72)
        r.font.size = Pt(7)
        r.font.color.rgb = LGREY

    def insert_photo(b64: str, caption: str, width_cm: float = 14.0):
        if not b64:
            return
        try:
            img_bytes = base64.b64decode(b64)
            stream = io.BytesIO(img_bytes)
            doc.add_picture(stream, width=Cm(width_cm))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.font.size = Pt(8)
                r.font.italic = True
                r.font.color.rgb = GREY
        except Exception as e:
            current_app.logger.warning("[psych-docx] Błąd wstawiania zdjęcia: %s", e)

    # ══════════════════════════════════════════════════════════════════════════
    # NAGŁÓWEK SZPITALA
    # ══════════════════════════════════════════════════════════════════════════
    h1 = doc.add_heading(
        szpital.get("nazwa", "Szpital Psychiatryczny im. Tylera Durdena"), 1
    )
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h1.runs:
        r.font.size = Pt(14)
        r.font.color.rgb = DARK

    p_adr = doc.add_paragraph(szpital.get("adres", ""))
    p_adr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_adr.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    p_odd = doc.add_paragraph(szpital.get("oddzial", ""))
    p_odd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_odd.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    doc.add_paragraph()

    tyt = doc.add_heading("HISTORIA CHOROBY — KARTA PRZYJĘCIA I HOSPITALIZACJI", 2)
    tyt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in tyt.runs:
        r.font.size = Pt(12)

    nr = raport.get("numer_historii_choroby", "NY-2026-00000")
    data_przyj = raport.get("data_przyjecia", datetime.now().strftime("%d.%m.%Y"))
    nr_p = doc.add_paragraph(
        f"Nr: {nr}  |  Data przyjęcia: {data_przyj}  |  "
        f"Lekarz prowadzący: {szpital.get('lekarz', 'Dr. T. Durden')}"
    )
    nr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in nr_p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 1 — DANE PACJENTA
    # ══════════════════════════════════════════════════════════════════════════
    heading("I. DANE PACJENTA", 2, RED, 11)
    dp = raport.get("dane_pacjenta", {})
    if isinstance(dp, dict) and "__error__" in dp:
        para(
            "[BŁĄD GENEROWANIA DANYCH PACJENTA]",
            italic=True,
            color=LGREY,
            size=9,
        )
    elif isinstance(dp, dict) and "__raw_text__" in dp:
        para(dp["__raw_text__"], italic=True, color=GREY, size=9)
    else:
        if not isinstance(dp, dict):
            current_app.logger.warning(
                "[psych-docx] dane_pacjenta zły typ: %s — wartość: %.100s",
                type(dp).__name__,
                dp,
            )
            dp = {}
        field("Imię i nazwisko", dp.get("imie_nazwisko", ""))
        field("Wiek", dp.get("wiek", ""))
        field("Adres", dp.get("adres", ""))
        field("Zawód", dp.get("zawod", ""))
        field("Stan cywilny", dp.get("stan_cywilny", ""))
        field("Nr ubezpieczenia", dp.get("numer_ubezpieczenia", ""))
    doc.add_paragraph()

    if photo_pacjent_b64:
        heading("DOKUMENTACJA FOTOGRAFICZNA — PRZYJĘCIE", 3, GREY, 9)
        insert_photo(
            photo_pacjent_b64,
            "Fot. 1 — Pacjent w kaftanie bezpieczeństwa. Oddział B. Materiał dowodowy.",
        )
        doc.add_paragraph()

    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 2 — POWÓD PRZYJĘCIA
    # ══════════════════════════════════════════════════════════════════════════
    heading("II. POWÓD PRZYJĘCIA", 2, RED, 11)
    powod = raport.get("powod_przyjecia")
    powod = _docx_unwrap(powod)
    if not powod or powod == "__BRAK__":
        para(
            "[brak danych — sekcja nie została wygenerowana]",
            italic=True,
            color=LGREY,
            size=9,
        )
    else:
        para(powod, size=10)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 3 — CYTATY Z PRZYJĘCIA
    # ══════════════════════════════════════════════════════════════════════════
    heading("III. CYTATY Z IZBY PRZYJĘĆ", 2, RED, 11)
    cytaty = raport.get("cytaty_z_przyjecia")
    if isinstance(cytaty, dict):
        if "__error__" in cytaty:
            para(
                "[BŁĄD GENEROWANIA]",
                italic=True,
                color=LGREY,
                size=9,
            )
        elif "__raw_text__" in cytaty:
            para(cytaty["__raw_text__"], italic=True, color=GREY, size=9)
        else:
            para(
                "[brak danych — cytaty nie zostały wygenerowane]",
                italic=True,
                color=LGREY,
                size=9,
            )
    elif not cytaty or cytaty == "__BRAK__":
        para(
            "[brak danych — cytaty nie zostały wygenerowane]",
            italic=True,
            color=LGREY,
            size=9,
        )
    elif isinstance(cytaty, list):
        for c in cytaty:
            if c and str(c).strip() not in ("", "__BRAK__"):
                para(str(c), italic=True, color=GREY, size=9)
    else:
        para(str(cytaty), italic=True, color=GREY, size=9)
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 4 — DEPOZYT
    # ══════════════════════════════════════════════════════════════════════════
    heading("IV. PROTOKÓŁ DEPOZYTU — PRZEDMIOTY SKONFISKOWANE", 2, RED, 11)
    dep = raport.get("depozyt", {})
    if _docx_render_error_or_raw(dep):
        doc.add_paragraph()
    else:
        if not isinstance(dep, dict):
            current_app.logger.warning(
                "[psych-docx] depozyt zły typ: %s — wartość: %.100s",
                type(dep).__name__,
                dep,
            )
            dep = {}
        if isinstance(dep, dict):
            lista = dep.get("lista_przedmiotow", [])
            proto = dep.get("protokol_depozytu", "")
        if isinstance(lista, str):
            lista = [x.strip() for x in lista.split(",") if x.strip()]
        if lista and lista != ["__BRAK__"]:
            for item in lista:
                if str(item).strip() in ("", "__BRAK__"):
                    continue
                p_item = doc.add_paragraph(style="List Bullet")
                r_item = p_item.add_run(str(item))
                r_item.font.size = Pt(10)
                r_item.font.color.rgb = DARK
        else:
            para(
                "[brak danych — lista przedmiotów nie została wygenerowana]",
                italic=True,
                color=LGREY,
                size=9,
            )
        if proto and proto != "__BRAK__":
            doc.add_paragraph()
            para(proto, italic=True, color=GREY, size=9)
    doc.add_paragraph()

    if photo_przedmioty_b64:
        heading("DOKUMENTACJA FOTOGRAFICZNA — DOWODY RZECZOWE", 3, GREY, 9)
        insert_photo(
            photo_przedmioty_b64,
            "Fot. 2 — Przedmioty skonfiskowane przy przyjęciu. "
            "Protokół dowodów rzeczowych, Oddział B.",
        )
        doc.add_paragraph()

    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 5 — FARMAKOLOGIA
    # ══════════════════════════════════════════════════════════════════════════
    heading("V. FARMAKOLOGIA — PEŁNA LISTA LEKÓW ZASTOSOWANYCH", 2, RED, 11)
    farm = raport.get("farmakologia", {})
    if _docx_render_error_or_raw(farm):
        doc.add_paragraph()
        leki_lista = []
    else:
        # BUGFIX: farmakologia może być stringiem
        if not isinstance(farm, dict):
            current_app.logger.warning(
                "[psych-docx] farmakologia zły typ: %s — wartość: %.100s",
                type(farm).__name__,
                farm,
            )
            farm = {}
        leki_lista = farm.get("leki", [])

    # Filtruj leki z wartością __BRAK__
    leki_lista = [
        l
        for l in leki_lista
        if isinstance(l, dict) and l.get("nazwa", "") != "__BRAK__"
    ]

    if leki_lista:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, label in enumerate(
            ["Nazwa leku", "Przedmioty odebrane", "Wskazanie", "Dawkowanie"]
        ):
            hdr[i].text = ""
            r = hdr[i].paragraphs[0].add_run(label)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RED
        for lek in leki_lista:
            if not isinstance(lek, dict):
                current_app.logger.warning(
                    "[psych-docx] leki_lista: element zły typ: %s — wartość: %.80s",
                    type(lek).__name__,
                    lek,
                )
                continue
            row = t.add_row().cells
            row[0].text = str(lek.get("nazwa", ""))
            row[1].text = str(lek.get("rzeczownik_zrodlowy", ""))
            row[2].text = str(lek.get("wskazanie", ""))
            row[3].text = str(lek.get("dawkowanie", ""))
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        doc.add_paragraph()
    else:
        para(
            "[brak danych — lista leków nie została wygenerowana]",
            italic=True,
            color=LGREY,
            size=9,
        )
        doc.add_paragraph()

    nota_farm = farm.get("nota_farmaceutyczna", "") if isinstance(farm, dict) else ""
    if nota_farm:
        para(nota_farm, italic=True, color=GREY, size=9)

    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 6 — HOSPITALIZACJA (14 dni)
    # ══════════════════════════════════════════════════════════════════════════
    heading("VI. PRZEBIEG HOSPITALIZACJI — DNI 1, 3, 14, 30", 2, RED, 11)

    dni_all = (raport.get("hospitalizacja_tydzien_1", []) or []) + (
        raport.get("hospitalizacja_tydzien_2", []) or []
    )

    for d in dni_all:
        if not isinstance(d, dict):
            current_app.logger.warning(
                "[psych-docx] hospitalizacja: dzień zły typ: %s — wartość: %.80s",
                type(d).__name__,
                d,
            )
            continue
        dzien = d.get("dzien", "")
        data = d.get("data", "")
        zdarz = d.get("zdarzenie", "")
        lek = d.get("lek", "")
        stan = d.get("stan_pacjenta", "")
        nota = d.get("nota_lekarska", "")

        p_day = doc.add_paragraph(style="List Bullet")
        r_day = p_day.add_run(f"Dzień {dzien}  [{data}]")
        r_day.bold = True
        r_day.font.size = Pt(10)
        r_day.font.color.rgb = RED

        lines = []
        if zdarz and zdarz != "__BRAK__":
            lines.append(f"Zdarzenie: {zdarz}")
        if lek and lek != "__BRAK__":
            lines.append(f"Lek: {lek}")
        if stan and stan != "__BRAK__":
            lines.append(f"Stan: {stan}")
        for line in lines:
            p_line = doc.add_paragraph()
            p_line.paragraph_format.left_indent = Pt(24)
            r_line = p_line.add_run(line)
            r_line.font.size = Pt(9)
            r_line.font.color.rgb = DARK

        if nota and nota != "__BRAK__":
            p_nota = doc.add_paragraph()
            p_nota.paragraph_format.left_indent = Pt(24)
            r_nota = p_nota.add_run(f"↳ {nota}")
            r_nota.italic = True
            r_nota.font.size = Pt(8)
            r_nota.font.color.rgb = GREY

    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 7 — WYPIS
    # ══════════════════════════════════════════════════════════════════════════
    heading("VII. KARTA WYPISU", 2, RED, 11)
    wypis = raport.get("wypis", {})
    if _docx_render_error_or_raw(wypis):
        doc.add_paragraph()
    else:
        if not isinstance(wypis, dict):
            current_app.logger.warning(
                "[psych-docx] wypis zły typ: %s — wartość: %.100s",
                type(wypis).__name__,
                wypis,
            )
            wypis = {}
        if isinstance(wypis, dict):
            field("Dzień wypisu", wypis.get("dzien_wypisu", ""))
            field("Powód wypisu", wypis.get("powod_wypisu", ""))
            doc.add_paragraph()
            para(wypis.get("stan_przy_wypisie", ""), size=10)
            doc.add_paragraph()
            heading("Zalecenia po wypisie:", 3, DARK, 10)
            zal = wypis.get("zalecenia_po_wypisie", [])
        if isinstance(zal, list):
            for z in zal:
                p_z = doc.add_paragraph(style="List Bullet")
                p_z.add_run(str(z)).font.size = Pt(10)
        elif zal:
            para(str(zal), size=10)
        doc.add_paragraph()
        if wypis.get("opis_pozegnania"):
            para(wypis["opis_pozegnania"], italic=True, color=GREY, size=9)
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 8 — DIAGNOZY
    # ══════════════════════════════════════════════════════════════════════════
    heading("VIII. DIAGNOZA PSYCHIATRYCZNA", 2, RED, 11)

    dw = raport.get("diagnoza_wstepna", {})
    if isinstance(dw, dict):
        heading("Diagnoza Wstępna:", 3, DARK, 10)
        p_diag = doc.add_paragraph()
        r1 = p_diag.add_run(dw.get("nazwa_lacinska", ""))
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RED
        if dw.get("nazwa_polska"):
            r2 = p_diag.add_run(f" (pol. {dw['nazwa_polska']})")
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK
        if dw.get("kod_dsm"):
            field("Kod DSM", dw["kod_dsm"], size=9)
        if dw.get("opis_kliniczny"):
            para(dw["opis_kliniczny"], size=10)
    doc.add_paragraph()

    dd = raport.get("diagnoza_dodatkowa", {})
    if isinstance(dd, dict):
        heading("Diagnoza Dodatkowa:", 3, DARK, 10)
        p_dd = doc.add_paragraph()
        r1 = p_dd.add_run(dd.get("nazwa_lacinska", ""))
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RED
        if dd.get("nazwa_polska"):
            r2 = p_dd.add_run(f" (pol. {dd['nazwa_polska']})")
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK
        if dd.get("opis_kliniczny"):
            para(dd["opis_kliniczny"], size=10)
    doc.add_paragraph()

    cw = raport.get("choroba_wspolistniejaca", {})
    if isinstance(cw, dict):
        heading("Choroba Współistniejąca:", 3, DARK, 10)
        p_cw = doc.add_paragraph()
        r1 = p_cw.add_run(cw.get("nazwa_lacinska", ""))
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RED
        if cw.get("nazwa_polska"):
            r2 = p_cw.add_run(f" (pol. {cw['nazwa_polska']})")
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK
        if cw.get("opis_kliniczny"):
            para(cw["opis_kliniczny"], size=10)
    doc.add_paragraph()

    objawy = raport.get("objawy", [])
    if objawy:
        heading("Objawy kliniczne:", 3, DARK, 10)
        for obj in objawy:
            p_obj = doc.add_paragraph(style="List Bullet")
            p_obj.add_run(str(obj)).font.size = Pt(10)
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 9 — ZALECENIA TYLERA
    # ══════════════════════════════════════════════════════════════════════════
    heading("IX. ZALECENIA TERAPEUTYCZNE", 2, RED, 11)
    zt = raport.get("zalecenia_tylera", {})
    if _docx_render_error_or_raw(zt):
        doc.add_paragraph()
    elif isinstance(zt, list):
        # AI zwróciło listę zamiast dict — renderuj każdy element
        for item in zt:
            if isinstance(item, dict):
                tresc = (
                    item.get("tresc")
                    or item.get("zadanie")
                    or item.get("zadanie_1")
                    or ""
                )
                if tresc and tresc != "__BRAK__":
                    p_z = doc.add_paragraph(style="List Number")
                    p_z.add_run(str(tresc)).font.size = Pt(10)
            elif item and str(item) != "__BRAK__":
                p_z = doc.add_paragraph(style="List Number")
                p_z.add_run(str(item)).font.size = Pt(10)
    elif isinstance(zt, dict):
        if zt.get("naglowek"):
            para(zt["naglowek"], bold=True, color=RED, size=10)
        for zadanie_key in ("zadanie_1", "zadanie_2", "zadanie_3"):
            zadanie_val = zt.get(zadanie_key, "")
            if zadanie_val and zadanie_val != "__BRAK__":
                p_z = doc.add_paragraph(style="List Number")
                p_z.add_run(str(zadanie_val)).font.size = Pt(10)
        if zt.get("podpis"):
            doc.add_paragraph()
            para(zt["podpis"], italic=True, color=GREY, size=9)
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 10 — ROKOWANIE
    # ══════════════════════════════════════════════════════════════════════════
    heading("X. ROKOWANIE", 2, RED, 11)
    rokowanie = (
        raport.get("rokowanie", "").strip()
        if isinstance(raport.get("rokowanie"), str)
        else ""
    )
    if not rokowanie or rokowanie == "__BRAK__":
        rokowanie = "---brak---"
    p_rok = doc.add_paragraph()
    r_rok = p_rok.add_run(rokowanie)
    r_rok.font.size = Pt(10)
    r_rok.font.color.rgb = RED
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 11 — INCYDENTY SPECJALNE
    # ══════════════════════════════════════════════════════════════════════════
    incydenty = raport.get("incydenty_specjalne", [])
    if incydenty:
        heading("XI. INCYDENTY SPECJALNE (protokoły wewnętrzne)", 2, RED, 11)
        for inc in incydenty:
            if inc == "__BRAK__":
                continue
            p_inc = doc.add_paragraph(style="List Bullet")
            p_inc.add_run(str(inc)).font.size = Pt(10)
        doc.add_paragraph()
        separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 12 — RELACJE ŚWIADKÓW
    # ══════════════════════════════════════════════════════════════════════════
    heading("XII. RELACJE ŚWIADKÓW", 2, RED, 11)
    relacje_swiadkow = raport.get("relacje_swiadkow", [])
    if relacje_swiadkow:
        heading("Świadkowie (osoby trzecie) — zeznania:", 3, DARK, 9)
        if isinstance(relacje_swiadkow, list):
            for sw in relacje_swiadkow:
                if isinstance(sw, dict):
                    imie = sw.get("imie_swiadka", "")
                    zawod = sw.get("zawod", "")
                    data = sw.get("data", "")
                    tresc = sw.get("tresc", "")
                    if tresc == "__BRAK__":
                        continue

                    # Nagłówek świadka
                    header_parts = []
                    if imie and imie != "__BRAK__":
                        header_parts.append(imie)
                    if zawod and zawod != "__BRAK__":
                        header_parts.append(f"({zawod})")
                    if data and data != "__BRAK__":
                        header_parts.append(data)

                    if header_parts:
                        p_hdr = doc.add_paragraph()
                        p_hdr.paragraph_format.left_indent = Pt(12)
                        r_hdr = p_hdr.add_run("  ".join(header_parts))
                        r_hdr.bold = True
                        r_hdr.font.size = Pt(8)
                        r_hdr.font.color.rgb = DARK

                    if tresc:
                        para(tresc, italic=True, color=GREY, size=8)
                else:
                    if str(sw) != "__BRAK__":
                        para(str(sw), italic=True, color=GREY, size=9)
        else:
            para(str(relacje_swiadkow), italic=True, color=GREY, size=9)
        doc.add_paragraph()
    else:
        para("Brak relacji świadków w dokumentacji.", italic=True, color=GREY, size=9)
    doc.add_paragraph()
    separator()

    # ══════════════════════════════════════════════════════════════════════════
    # SEKCJA 13 — PODPIS + NOTATKI PERSONELU
    # ══════════════════════════════════════════════════════════════════════════
    heading("XIII. PODPIS I NOTATKI PERSONELU", 2, RED, 11)
    field("Lekarz prowadzący", szpital.get("lekarz", "Dr. T. Durden, MD, PhD, FIGHT"))
    doc.add_paragraph()

    notatki_p = raport.get("notatki_pielegniarek") or raport.get("notatka_pielegniarki")
    if notatki_p:
        heading("Notatki pielęgniarek:", 3, DARK, 9)
        if isinstance(notatki_p, list):
            for n in notatki_p:
                if isinstance(n, dict):
                    imie = n.get("imie_pielegniarki", "")
                    data = n.get("data", "")
                    tresc = n.get("tresc", "")
                    if tresc == "__BRAK__":
                        continue
                    header_parts = [x for x in [imie, data] if x and x != "__BRAK__"]
                    if header_parts:
                        p_hdr = doc.add_paragraph()
                        p_hdr.paragraph_format.left_indent = Pt(12)
                        r_hdr = p_hdr.add_run("  ".join(header_parts))
                        r_hdr.bold = True
                        r_hdr.font.size = Pt(8)
                        r_hdr.font.color.rgb = DARK
                    if tresc:
                        para(tresc, italic=True, color=GREY, size=8)
                else:
                    if str(n) != "__BRAK__":
                        para(str(n), italic=True, color=GREY, size=9)
        else:
            para(str(notatki_p), italic=True, color=GREY, size=9)
        doc.add_paragraph()

    notatki_s = raport.get("notatki_sprzataczki") or raport.get("notatka_sprzataczki")
    if notatki_s:
        heading("Notatki sprzątaczki:", 3, DARK, 9)
        if isinstance(notatki_s, list):
            for n in notatki_s:
                if isinstance(n, dict):
                    data = n.get("data", "")
                    tresc = n.get("tresc", "")
                    if tresc == "__BRAK__":
                        continue
                    if data and data != "__BRAK__":
                        p_hdr = doc.add_paragraph()
                        p_hdr.paragraph_format.left_indent = Pt(12)
                        r_hdr = p_hdr.add_run(data)
                        r_hdr.bold = True
                        r_hdr.font.size = Pt(8)
                        r_hdr.font.color.rgb = DARK
                    if tresc:
                        para(tresc, italic=True, color=GREY, size=8)
                else:
                    if str(n) != "__BRAK__":
                        para(str(n), italic=True, color=GREY, size=9)
        else:
            para(str(notatki_s), italic=True, color=GREY, size=9)

    # ══════════════════════════════════════════════════════════════════════════
    # ZAPIS
    # ══════════════════════════════════════════════════════════════════════════
    buf = io.BytesIO()
    doc.save(buf)
    b64 = base64.b64encode(buf.getvalue()).decode("ascii")
    current_app.logger.info("[psych-docx] DOCX OK (%dKB)", len(buf.getvalue()) // 1024)
    return b64


# ─────────────────────────────────────────────────────────────────────────────
# GŁÓWNA FUNKCJA PUBLICZNA
# ─────────────────────────────────────────────────────────────────────────────


def build_raport(
    body: str,
    previous_body: str | None,
    res_text: str,
    nouns_dict: dict,
    sender_name: str = "",
    gender: str = "patient",
    test_mode: bool = False,
) -> dict:
    """Buduje kompletny raport psychiatryczny — fallbacki na każdym poziomie."""
    current_app.logger.info("[psych-raport] START build_raport")
    cfg = _load_cfg()
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    # Runda 1 — sekcje niezależne (równolegle)
    sekcja_pacjent = {"data": {}, "status": "error"}
    sekcja_dep_leki = {"data": {}, "status": "error"}
    sekcja_diagnozy = {"data": {}, "status": "error"}
    sekcja_flux = {"data": {}, "status": "error"}

    try:
        sekcja_pacjent = _wrap_section(
            _sekcja_pacjent(cfg, body, sender_name),
            "pacjent",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda1 pacjent błąd: %s", e)

    try:
        sekcja_dep_leki = _wrap_section(
            _sekcja_depozyt_leki(cfg, body, nouns_dict),
            "depozyt",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda1 depozyt błąd: %s", e)

    try:
        sekcja_diagnozy = _wrap_section(
            _sekcja_diagnozy(cfg, body, previous_body),
            "diagnozy",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda1 diagnozy błąd: %s", e)

    try:
        sekcja_flux = _wrap_section(
            _sekcja_flux_prompty(
                cfg, body, nouns_dict, sender_name, gender, test_mode=test_mode
            ),
            "flux",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda1 flux błąd: %s", e)

    sekcja_pacjent_data = _unwrap_section(sekcja_pacjent)
    sekcja_dep_leki_data = _unwrap_section(sekcja_dep_leki)
    sekcja_diagnozy_data = _unwrap_section(sekcja_diagnozy)
    sekcja_flux_data = _unwrap_section(sekcja_flux)

    # Wyznaczenie daty przyjęcia
    data_przyjecia = sekcja_pacjent_data.get(
        "data_przyjecia", datetime.now().strftime("%d.%m.%Y")
    )
    # BUGFIX: farmakologia może być stringiem gdy AI zwróci błędny typ — guard przed .get()
    _farm_tmp = sekcja_dep_leki_data.get("farmakologia", {})
    if not isinstance(_farm_tmp, dict):
        current_app.logger.warning(
            "[psych-raport] sekcja_dep_leki.farmakologia zły typ: %s — wartość: %.100s",
            type(_farm_tmp).__name__,
            _farm_tmp,
        )
        leki_lista = []
    else:
        leki_lista = _farm_tmp.get("leki", [])

    # Runda 2 — dni hospitalizacji
    dni_1_7 = []
    dni_8_14 = []
    sekcja_wypis = {"data": {}, "status": "error"}

    try:
        dni_1_7 = (
            _unwrap_section(
                _wrap_section(
                    _sekcja_tydzien(cfg, body, leki_lista, 1, data_przyjecia),
                    "tydzien1",
                )
            )
            or []
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda2 tydzien1 błąd: %s", e)

    try:
        dni_8_14 = (
            _unwrap_section(
                _wrap_section(
                    _sekcja_tydzien(cfg, body, leki_lista, 2, data_przyjecia),
                    "tydzien2",
                )
            )
            or []
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda2 tydzien2 błąd: %s", e)

    try:
        sekcja_wypis = _wrap_section(
            _sekcja_wypis(cfg, body, data_przyjecia),
            "wypis",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda2 wypis błąd: %s", e)

    # Runda 3 — zalecenia + leczenie specjalne
    sekcja_zalecenia = {"data": {}, "status": "error"}
    sekcja_leczenie_specjalne = {"data": {}, "status": "error"}
    try:
        sekcja_zalecenia = _wrap_section(
            _sekcja_zalecenia(cfg, body, dni_1_7, dni_8_14),
            "zalecenia",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda3 zalecenia błąd: %s", e)

    try:
        sekcja_leczenie_specjalne = _wrap_section(
            _sekcja_leczenie_specjalne(cfg, body, dni_1_7, dni_8_14),
            "leczenie_specjalne",
        )
    except Exception as e:
        current_app.logger.error("[psych-raport] Runda3 leczenie_specjalne błąd: %s", e)

    sekcja_wypis_data = _unwrap_section(sekcja_wypis)
    sekcja_zalecenia_data = _unwrap_section(sekcja_zalecenia)
    sekcja_leczenie_specjalne_data = _unwrap_section(sekcja_leczenie_specjalne)

    # Scalenie całości
    raport = {}

    # ── dane_pacjenta: _sekcja_pacjent zwraca PŁASKI dict — opakowujemy
    PACJENT_FIELDS = {
        "imie_nazwisko",
        "wiek",
        "adres",
        "zawod",
        "stan_cywilny",
        "numer_ubezpieczenia",
        "rodowod",
        "rodowód",
    }
    # Klucze które mogą być na płaskim poziomie obok dane_pacjenta
    RAPORT_TOP_KEYS = {
        "numer_historii_choroby",
        "data_przyjecia",
        "powod_przyjecia",
        "cytaty_z_przyjecia",
    }
    if isinstance(sekcja_pacjent_data, dict):
        if _is_wrapped_section(sekcja_pacjent_data):
            raport["dane_pacjenta"] = sekcja_pacjent_data
            raport["powod_przyjecia"] = sekcja_pacjent_data
            raport["cytaty_z_przyjecia"] = sekcja_pacjent_data
        elif "dane_pacjenta" in sekcja_pacjent_data and isinstance(
            sekcja_pacjent_data["dane_pacjenta"], dict
        ):
            raport["dane_pacjenta"] = sekcja_pacjent_data["dane_pacjenta"]
            for k, v in sekcja_pacjent_data.items():
                if k != "dane_pacjenta":
                    raport[k] = v
        else:
            # AI zwróciło płaski dict — zbieramy pola pacjenta
            dane_pac = {
                k: v for k, v in sekcja_pacjent_data.items() if k in PACJENT_FIELDS
            }
            # Zbieramy też top-level klucze raportu (numer, data, powód, cytaty)
            for k, v in sekcja_pacjent_data.items():
                if k in RAPORT_TOP_KEYS:
                    raport[k] = v
            # Zabezpieczenie: dane_pacjenta MUSI być dict, nigdy stringiem
            raport["dane_pacjenta"] = (
                dane_pac
                if dane_pac
                else {
                    "imie_nazwisko": sender_name or "pacjent",
                    "wiek": "__BRAK__",
                    "adres": "__BRAK__",
                    "zawod": "__BRAK__",
                    "stan_cywilny": "__BRAK__",
                    "numer_ubezpieczenia": "__BRAK__",
                }
            )
            for k, v in sekcja_pacjent_data.items():
                if k not in PACJENT_FIELDS and k not in RAPORT_TOP_KEYS:
                    raport[k] = v

    # ── depozyt i farmakologia
    if _is_wrapped_section(sekcja_dep_leki_data):
        raport["depozyt"] = sekcja_dep_leki_data
        raport["farmakologia"] = sekcja_dep_leki_data
    else:
        dep_raw = (
            sekcja_dep_leki_data.get("depozyt", {}) if isinstance(sekcja_dep_leki_data, dict) else {}
        )
        if not isinstance(dep_raw, dict):
            current_app.logger.warning(
                "[psych-raport] depozyt zły typ po scaleniu: %s — wartość: %.100s",
                type(dep_raw).__name__,
                dep_raw,
            )
        raport["depozyt"] = dep_raw if isinstance(dep_raw, dict) else {}

        farm_raw = (
            sekcja_dep_leki_data.get("farmakologia", {})
            if isinstance(sekcja_dep_leki_data, dict)
            else {}
        )
        if not isinstance(farm_raw, dict):
            current_app.logger.warning(
                "[psych-raport] farmakologia zły typ po scaleniu: %s — wartość: %.100s",
                type(farm_raw).__name__,
                farm_raw,
            )
        raport["farmakologia"] = farm_raw if isinstance(farm_raw, dict) else {}

    raport["hospitalizacja_tydzien_1"] = dni_1_7 if isinstance(dni_1_7, list) else []
    raport["hospitalizacja_tydzien_2"] = dni_8_14 if isinstance(dni_8_14, list) else []

    # ── wypis: _sekcja_wypis zwraca PŁASKI dict — opakowujemy
    WYPIS_FIELDS = {
        "dzien_wypisu",
        "powod_wypisu",
        "zalecenia_po_wypisie",
        "stan_przy_wypisie",
        "opis_pozegnania",
    }
    if _is_wrapped_section(sekcja_wypis_data):
        raport["wypis"] = sekcja_wypis_data
    elif (
        isinstance(sekcja_wypis_data, dict)
        and "__error__" not in sekcja_wypis_data
        and "__raw_text__" not in sekcja_wypis_data
    ):
        if "wypis" in sekcja_wypis_data and isinstance(
            sekcja_wypis_data["wypis"], dict
        ):
            raport["wypis"] = sekcja_wypis_data["wypis"]
            for k, v in sekcja_wypis_data.items():
                if k != "wypis":
                    raport[k] = v
        else:
            raport["wypis"] = {
                k: v for k, v in sekcja_wypis_data.items() if k in WYPIS_FIELDS
            }
            for k, v in sekcja_wypis_data.items():
                if k not in WYPIS_FIELDS:
                    raport[k] = v

    if _is_wrapped_section(sekcja_zalecenia_data):
        raport["zalecenia_tylera"] = sekcja_zalecenia_data

    # ── diagnozy i zalecenia: ochrona przed stringami zamiast dict
    DICT_KEYS = {
        "diagnoza_wstepna",
        "diagnoza_dodatkowa",
        "choroba_wspolistniejaca",
        "zalecenia_tylera",
        "depozyt",
        "farmakologia",
        "wypis",
        "dane_pacjenta",
    }
    # Klucze które NIE powinny nadpisywać wartości już ustawionych przez wcześniejsze sekcje
    PROTECTED_KEYS = {"dane_pacjenta", "depozyt", "farmakologia", "wypis"}
    for sekcja in (sekcja_diagnozy_data, sekcja_zalecenia_data):
        if (
            not isinstance(sekcja, dict)
            or "__error__" in sekcja
            or "__raw_text__" in sekcja
        ):
            current_app.logger.warning(
                "[psych-raport] sekcja zły typ przy scalaniu: %s",
                type(sekcja).__name__,
            )
            continue
        for k, v in sekcja.items():
            if k in DICT_KEYS and not isinstance(v, dict):
                current_app.logger.warning(
                    "[psych-raport] klucz '%s' zły typ: %s — wartość: %.100s",
                    k,
                    type(v).__name__,
                    v,
                )
                # ── BUG FIX 1-3: diagnoza_wstepna / diagnoza_dodatkowa / choroba_wspolistniejaca
                # AI zwraca string zamiast dict — zawiń tekst w opis_kliniczny żeby nie tracić treści
                if (
                    k
                    in (
                        "diagnoza_wstepna",
                        "diagnoza_dodatkowa",
                        "choroba_wspolistniejaca",
                    )
                    and isinstance(v, str)
                    and v.strip()
                ):
                    # Wyciągnij nazwę po "Diagnoza: " jeśli istnieje, resztę do opis_kliniczny
                    tekst = v.strip()
                    # Usuń prefiks "Diagnoza: " jeśli AI go dodało
                    if tekst.lower().startswith("diagnoza:"):
                        tekst = tekst[len("Diagnoza:") :].strip()
                    raport[k] = {
                        "nazwa_lacinska": "",
                        "nazwa_polska": "",
                        "kod_dsm": "",
                        "opis_kliniczny": tekst,
                    }
                    current_app.logger.warning(
                        "[psych-raport] klucz '%s' naprawiony: str→dict z opis_kliniczny",
                        k,
                    )
                # ── BUG FIX 4: zalecenia_tylera jako lista notatek pielęgniarek
                # AI myli klucze — lista z imie_pielegniarki trafia do zalecenia_tylera zamiast do notatki
                elif (
                    k == "zalecenia_tylera"
                    and isinstance(v, list)
                    and v
                    and isinstance(v[0], dict)
                    and "imie_pielegniarki" in v[0]
                ):
                    # Przenieś do notatki_pielegniarek tylko gdy tam jeszcze nic nie ma
                    if not raport.get("notatki_pielegniarek"):
                        raport["notatki_pielegniarek"] = v
                        current_app.logger.warning(
                            "[psych-raport] zalecenia_tylera (lista pielęgniarek) → przeniesiono do notatki_pielegniarek"
                        )
                    # zalecenia_tylera zostają jako pusty dict — sekcja IX nie będzie pusta z błędną treścią
                    raport[k] = {}
                # Nie nadpisuj poprawnego dict'a z wcześniejszej sekcji pustym {}
                elif isinstance(v, list) and v:
                    raport[k] = v
                elif k not in raport or not isinstance(raport.get(k), dict):
                    raport[k] = {}
            elif k in PROTECTED_KEYS and isinstance(raport.get(k), dict) and raport[k]:
                # Nie nadpisuj niepustego dict'a z wcześniejszej sekcji danymi z diagnozy/zalecen
                pass
            else:
                raport[k] = v

    # Relacje świadków (po scaleniu raportu, bo potrzebuje kontekstu)
    relacje_result = {}
    try:
        relacje_result = _sekcja_relacje_swiadkow(cfg, body, raport) or {}
    except Exception as e:
        current_app.logger.error("[psych-raport] Relacje świadków błąd: %s", e)
    raport["relacje_swiadkow"] = relacje_result.get("relacje_swiadkow", [])

    # Leczenie specjalne (deepseek_9)
    if _is_wrapped_section(sekcja_leczenie_specjalne_data):
        raport["leczenie_specjalne"] = sekcja_leczenie_specjalne_data
    else:
        raport["leczenie_specjalne"] = sekcja_leczenie_specjalne_data.get(
            "leczenie_specjalne", []
        )

    # Twarda walidacja typu dla dane_pacjenta
    dane_pacjenta = raport.get("dane_pacjenta", {})
    if not isinstance(dane_pacjenta, dict):
        if isinstance(dane_pacjenta, list) and dane_pacjenta:
            first = dane_pacjenta[0]
            if isinstance(first, dict):
                raport["dane_pacjenta"] = first
            else:
                current_app.logger.warning(
                    "[psych-raport] dane_pacjenta: root list bez dict — fallback do domyślnych",
                )
                raport["dane_pacjenta"] = {
                    "imie_nazwisko": sender_name or "pacjent",
                    "wiek": "__BRAK__",
                    "adres": "__BRAK__",
                    "zawod": "__BRAK__",
                    "stan_cywilny": "__BRAK__",
                    "numer_ubezpieczenia": "__BRAK__",
                }
        elif isinstance(dane_pacjenta, str) and dane_pacjenta.strip():
            raport["dane_pacjenta"] = {
                "imie_nazwisko": dane_pacjenta.strip(),
                "wiek": "__BRAK__",
                "adres": "__BRAK__",
                "zawod": "__BRAK__",
                "stan_cywilny": "__BRAK__",
                "numer_ubezpieczenia": "__BRAK__",
            }
        else:
            current_app.logger.warning(
                "[psych-raport] dane_pacjenta: zły typ %s — używam domyślnych",
                type(dane_pacjenta).__name__,
            )
            raport["dane_pacjenta"] = {
                "imie_nazwisko": sender_name or "pacjent",
                "wiek": "__BRAK__",
                "adres": "__BRAK__",
                "zawod": "__BRAK__",
                "stan_cywilny": "__BRAK__",
                "numer_ubezpieczenia": "__BRAK__",
            }

    # FLUX — zdjęcia równolegle
    prompt_pacjent = sekcja_flux.get("prompt_pacjent", "")
    prompt_przedmioty = sekcja_flux.get("prompt_przedmioty", "")
    photo_1, photo_2 = _generate_photos_parallel(
        prompt_pacjent, prompt_przedmioty, test_mode=test_mode
    )

    # Budowanie DOCX
    photo_1_b64 = photo_1["base64"] if photo_1 else None
    photo_2_b64 = photo_2["base64"] if photo_2 else None
    docx_b64 = _build_docx(raport, photo_1_b64, photo_2_b64, cfg)

    if not docx_b64:
        current_app.logger.error("[psych-raport] DOCX nie wygenerowany")
        return {"raport_pdf": None, "psych_photo_1": photo_1, "psych_photo_2": photo_2}

    _dp_tmp = raport.get("dane_pacjenta", {})
    imie = (
        _dp_tmp.get("imie_nazwisko", "pacjent")
        if isinstance(_dp_tmp, dict)
        else (sender_name or "pacjent")
    )
    safe = re.sub(r"[^a-zA-Z0-9_-]", "_", imie)[:30]

    raport_pdf_dict = {
        "base64": docx_b64,
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "filename": f"raport_psychiatryczny_{safe}_{ts}.docx",
    }

    current_app.logger.info(
        "[psych-raport] DONE raport=%s photo1=%s photo2=%s",
        raport_pdf_dict["filename"],
        photo_1["filename"] if photo_1 else "brak",
        photo_2["filename"] if photo_2 else "brak",
    )

    return {
        "raport_pdf": raport_pdf_dict,
        "psych_photo_1": photo_1,
        "psych_photo_2": photo_2,
    }
