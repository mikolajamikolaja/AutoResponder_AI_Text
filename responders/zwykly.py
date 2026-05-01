"""
responders/zwykly.py
Responder emocjonalny — Tyler Durden + Sokrates.

ZMIANY W TEJ WERSJI:
  1. prompt.txt → prompt.json  (czysta struktura, render programowy)
  2. DeepSeek do generowania tekstu
  3. Brak ograniczeń długości tekstu
  4. Generowanie tryptyku FLUX (3 panele Fight Club)
     - styl z zwykly_obrazek_tyler.js
     - DeepSeek generuje prompty dla każdego panelu
     - rotacja tokenów HF (HF_TOKEN, HF_TOKEN1...HF_TOKEN20)
     - jeśli tokeny wyczerpane → wysyłamy tyle ile wygenerowano
  5. Każdy panel PNG jest od razu konwertowany do JPG 95% (Pillow)
     - PNG FLUX ~2MB → JPG 95% ~300-500KB
     - nazwa: tyler_YYYYMMDD_HHMMSS_panel{N}.jpg
     - zwracany content_type: image/jpeg
  6. Nadawca dostaje: reply_html + emotka PNG + PDF emocji + tryptyk JPG
     (inline w mailu + załącznik JPG)
  7. Pole triptych_for_drive zawiera listę JPG do zapisu na Google Drive
     przez GAS (_saveTylerJpgsToDrive) — ta sama logika co smierc.py
"""

import os
import re
import io
import zipfile
import json
import html as html_module
import base64
import random
import logging
import requests
from datetime import datetime

# Bezpieczny logger modułu — działa w wątkach bez kontekstu Flask
logger = logging.getLogger(__name__)

from core.logging_reporter import get_logger

execution_logger = get_logger()

from core.ai_client import (
    call_deepseek,
    extract_clean_text,
    sanitize_model_output,
    MODEL_TYLER,
)
from core.files import read_file_base64
from core.html_builder import build_html_reply

# reportlab — budowanie PDF CV
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from responders.zwykly_psychiatryczny_raport import build_raport

# ─────────────────────────────────────────────────────────────────────────────
# ŚCIEŻKI
# ─────────────────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
EMOTKI_DIR = os.path.join(BASE_DIR, "emotki")
PDF_DIR = os.path.join(BASE_DIR, "pdf")
PROMPTS_DIR = os.path.join(BASE_DIR, "prompts")

PROMPT_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_prompt.json")
CV_CONTENT_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_cv_content.json")
CV_PHOTO_FLUX_PATH = os.path.join(PROMPTS_DIR, "zwykly_cv_photo_flux.json")
ICON_FLUX_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_icon_flux.json")
SUBSTITUTE_IMAGE_PATH = os.path.join(BASE_DIR, "images", "zastepczy.jpg")
STYLE_JS_PATH = os.path.join(PROMPTS_DIR, "zwykly_panel_wytyczne.json")
ANKIETA_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_ankieta.json")
HOROSKOP_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_horoskop.json")
KARTA_RPG_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_karta_rpg.json")
RAPORT_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_raport.json")
PLAKAT_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_plakat.json")
GRA_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_gra.json")
PSYCHIATRYCZNY_OBRAZEK_JSON_PATH = os.path.join(
    PROMPTS_DIR, "zwykly_psychiatryczny_obrazek.json"
)
NOUNS_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_znajdz_rzeczowniki.json")


# ─────────────────────────────────────────────────────────────────────────────
# POMOCNIK: rejestracja czcionek z polskimi znakami
# ─────────────────────────────────────────────────────────────────────────────


def _register_fonts() -> tuple:
    """
    Rejestruje czcionki DejaVuSans (obsługują polskie znaki) w reportlab.
    Szuka najpierw w katalogu fonts/ projektu, potem w ścieżkach systemowych.
    Zwraca (FN, FB) — nazwy czcionek normalnej i pogrubionej.
    Bezpieczne do wielokrotnego wywołania (reportlab ignoruje duplikaty).
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    FONT_DIR = os.path.join(BASE_DIR, "fonts")

    # Kolejność szukania: projekt → system Ubuntu/Debian → system ogólny
    NORMAL_PATHS = [
        os.path.join(FONT_DIR, "DejaVuSans.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
    ]
    BOLD_PATHS = [
        os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
    ]

    FN, FB = "Helvetica", "Helvetica-Bold"

    for path in NORMAL_PATHS:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("DejaVuSans", path))
                FN = "DejaVuSans"
                break
            except Exception:
                continue

    for path in BOLD_PATHS:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", path))
                FB = "DejaVuSans-Bold"
                break
            except Exception:
                continue

    return FN, FB


# ─────────────────────────────────────────────────────────────────────────────
# STAŁE — przeniesione do core/config.py
# ─────────────────────────────────────────────────────────────────────────────
from core.config import (
    MAX_DLUGOSC_EMAIL,
    HF_API_URL,
    HF_STEPS,
    HF_GUIDANCE,
    HF_TIMEOUT,
    TYLER_JPG_QUALITY,
    EMOCJA_MAP,
    FALLBACK_EMOT,
)

from core.hf_token_manager import get_active_tokens, mark_dead, hf_tokens

# ─────────────────────────────────────────────────────────────────────────────
# HELPER: pakowanie pliku do ZIP (Gmail blokuje .html, .htm, .svg)
# ─────────────────────────────────────────────────────────────────────────────


def _to_zip(content: bytes, inner_filename: str, zip_filename: str) -> dict:
    """Pakuje bytes do ZIP i zwraca dict {base64, content_type, filename}."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(inner_filename, content)
    return {
        "base64": base64.b64encode(buf.getvalue()).decode("ascii"),
        "content_type": "application/zip",
        "filename": zip_filename,
    }


# ─────────────────────────────────────────────────────────────────────────────
# WYMUSZANIE STARTU JSON od '{'
# ─────────────────────────────────────────────────────────────────────────────
_JSON_FORCE_SUFFIX = "\n\nOdpowiedź (zacznij od {):"
_JSON_FORCE_SYSTEM = (
    "ZAWSZE zacznij odpowiedź od znaku {. Zakaz jakiegokolwiek tekstu przed {."
)


def _ju(user_prompt: str) -> str:
    """Wymusza start odpowiedzi od '{' w user promptcie."""
    return user_prompt + _JSON_FORCE_SUFFIX


def _js(system_prompt: str) -> str:
    """Dodaje wymóg startu od '{' do system promptu."""
    if not system_prompt:
        return _JSON_FORCE_SYSTEM
    return system_prompt + "\n" + _JSON_FORCE_SYSTEM


# ═══════════════════════════════════════════════════════════════════════════════
# ŁADOWANIE prompt.json
# ═══════════════════════════════════════════════════════════════════════════════


def _extract_first_json_object(text: str) -> str:
    """
    Wyciąga PIERWSZY kompletny obiekt JSON ({ ... }) lub tablicę ([ ... ]) z tekstu.
    Liczy nawiasy — bezpieczniejsze niż zachłanny regex (naprawia 'Extra data').
    Obsługuje zarówno dict jak i list na najwyższym poziomie.
    """
    if not text:
        return ""
    start = None
    stack = []
    in_string = False
    escape = False
    quote_char = None

    for idx, ch in enumerate(text):
        if escape:
            escape = False
            continue
        if ch == "\\" and in_string:
            escape = True
            continue
        if in_string:
            if ch == quote_char:
                in_string = False
            continue
        if ch in ('"', "'"):
            in_string = True
            quote_char = ch
            continue
        if ch in ("{", "["):
            start = idx
            stack.append(ch)
            break

    if start is None:
        return ""

    for idx in range(start, len(text)):
        ch = text[idx]
        if escape:
            escape = False
            continue
        if ch == "\\" and in_string:
            escape = True
            continue
        if in_string:
            if ch == quote_char:
                in_string = False
            continue
        if ch in ('"', "'"):
            in_string = True
            quote_char = ch
            continue
        if ch in ("{", "["):
            stack.append(ch)
        elif ch in ("}", "]") and stack:
            top = stack[-1]
            if (top == "{" and ch == "}") or (top == "[" and ch == "]"):
                stack.pop()
                if not stack:
                    return text[start : idx + 1]

    return text[start:]


def _extract_body_html(html_text: str) -> str:
    if not html_text:
        return ""
    match = re.search(
        r"<body[^>]*>(.*?)</body>", html_text, flags=re.DOTALL | re.IGNORECASE
    )
    if match:
        return match.group(1).strip()
    return html_text.strip()


def _format_plain_text_as_html(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = html_module.escape(normalized)
    normalized = normalized.replace("\n\n", "</p><p>")
    normalized = normalized.replace("\n", "<br>")
    return f"<div class='section'><p>{normalized}</p></div>"


def _build_combined_reply_html(sections: list[str]) -> str:
    content = "\n<hr class='section-separator'>\n".join(s for s in sections if s)
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <style>
        body {{ margin: 0; padding: 20px; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; background: #ffffff; min-height: 100vh; }}
        .container {{ max-width: 600px; margin: 0 auto; background: rgba(255, 255, 255, 0.95); border-radius: 12px; padding: 30px; box-shadow: 0 8px 32px rgba(0, 0, 0, 0.1); border: 2px solid rgba(200, 220, 255, 0.3); }}
        .header {{ background: linear-gradient(135deg, #B3E5FC 0%, #C8E6C9 100%); padding: 20px; border-radius: 8px; margin-bottom: 20px; border-left: 5px solid #81C784; }}
        .content {{ font-size: 15px; line-height: 1.8; color: #333; }}
        .content p {{ margin: 15px 0; color: #000000; }}
        .section {{ margin-bottom: 20px; }}
        .section-separator {{ border: none; border-top: 1px solid #DDD; margin: 30px 0; }}
        .footer {{ margin-top: 30px; padding: 20px 15px 15px 15px; border-top: 2px solid #FFE0B2; font-size: 12px; color: #0a8a0a; text-align: center; background: linear-gradient(to bottom, transparent, rgba(255, 224, 178, 0.2)); border-radius: 6px; }}
        .footer a {{ color: #0a8a0a; text-decoration: none; border-bottom: 1px dotted #0a8a0a; }}
        .footer a:hover {{ border-bottom: 1px solid #0a8a0a; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <p style="margin: 0; color: #333; font-weight: 500;">✉️ Odpowiedź automatyczna</p>
        </div>
        <div class="content">{content}</div>
        <div class="footer">
            <p style="margin: 0 0 10px 0;">Odpowiedź wygenerowana automatycznie przez system Script + Render.<br><span style="font-size: 11px; color: #088a08;">Projekt dostępny na GitHub:<br><a href="https://github.com/legionowopawel/AutoResponder_AI_Text" style="color: #088a08; text-decoration: none;">AutoResponder_AI_Text</a></span></p>
            <p style="margin: 0; font-size: 11px; color: #088a08;">Portfolio:<br><a href="https://legionowopawel.github.io/AutoResponder_AI_Text/" style="color: #088a08; text-decoration: none;">legionowopawel.github.io/AutoResponder_AI_Text</a></p>
        </div>
    </div>
</body>
</html>"""


def _wrap_section_html(content: str, title: str | None = None) -> str:
    if not content:
        return ""
    inner = _extract_body_html(content)
    title_html = f"<h2>{html_module.escape(title)}</h2>" if title else ""
    return f"<div class='section'>{title_html}{inner}</div>"


def _wrap_plain_text_section(text: str, title: str | None = None) -> str:
    if not text:
        return ""
    section = _format_plain_text_as_html(text)
    if title:
        title_html = f"<h2>{html_module.escape(title)}</h2>"
        return f"<div class='section'>{title_html}{section}</div>"
    return section


def _collect_section_attachments(
    section_output: dict, docs: list, docx_list: list, images: list
) -> None:
    if not isinstance(section_output, dict):
        return

    def _is_html_file(item: dict) -> bool:
        """Zwraca True jeśli plik jest .htm/.html — Gmail blokuje te rozszerzenia."""
        if not isinstance(item, dict):
            return False
        fname = item.get("filename", "").lower()
        ct = item.get("content_type", "").lower()
        return (
            fname.endswith(".htm")
            or fname.endswith(".html")
            or ct in ("text/html", "text/htm")
        )

    if section_output.get("docs"):
        docs.extend(
            [
                item
                for item in section_output.get("docs", [])
                if isinstance(item, dict) and not _is_html_file(item)
            ]
        )
    if section_output.get("docx_list"):
        docx_list.extend(
            [
                item
                for item in section_output.get("docx_list", [])
                if isinstance(item, dict)
            ]
        )
    if section_output.get("images"):
        images.extend(
            [
                item
                for item in section_output.get("images", [])
                if isinstance(item, dict)
            ]
        )
    if section_output.get("image") and isinstance(section_output.get("image"), dict):
        images.append(section_output.get("image"))
    if section_output.get("htm_for_drive") and isinstance(
        section_output.get("htm_for_drive"), dict
    ):
        # htm_for_drive idzie TYLKO na Drive, nie do załączników maila
        # (Gmail blokuje .htm/.html — zamiast tego używamy .zip w _to_zip)
        pass


def _normalize_section_html_text(html_text: str) -> str:
    return _extract_body_html(html_text)


def _extract_section_html(raw_html: str) -> str:
    return _extract_body_html(raw_html)


def _sanitize_reply_html(html_text: str) -> str:
    return html_text or ""


def _render_body_sections(
    main_html: str,
    emocje_html: str,
    dociekliwy_html: str,
    scrabble_html: str = "",
) -> str:
    sections = [main_html, emocje_html, dociekliwy_html, scrabble_html]
    return _build_combined_reply_html(sections)


def _build_email_body_text(text: str) -> str:
    return _format_plain_text_as_html(text)


def _strip_leading_markdown(raw: str) -> str:
    if not raw:
        return ""
    match = re.search(r"[\{\[]", raw)
    if match:
        return raw[match.start() :]
    return raw


def _strip_json_markdown(raw: str) -> str:
    """
    Czyści markdown wokół JSON-a.
    WAŻNE: NIE usuwa prowadzącego przecinka — zachowujemy go żeby
    _parse_json_safe wiedział, że to fragment tablicy/obiektu.
    """
    if not raw:
        return ""
    raw = raw.strip()
    # Usuń code fences (```json ... ```)
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    raw = raw.strip()

    # Jeśli zaczyna się od '{' lub '[' — wyciągnij kompletny obiekt
    if raw and raw[0] in ("{", "["):
        fragment = _extract_first_json_object(raw)
        if fragment:
            return fragment
        return raw

    # Jeśli zaczyna się od ',' — fragment tablicy/obiektu; zostawiamy bez zmian
    if raw and raw[0] == ",":
        return raw

    # Szukaj pierwszego '{' lub '[' (po słowie "json", backtickach itp.)
    match = re.search(r"[\{\[]", raw)
    if match:
        candidate = raw[match.start() :]
        fragment = _extract_first_json_object(candidate)
        if fragment:
            return fragment
        return candidate

    # Ostateczność — minimalne czyszczenie (bez usuwania przecinka)
    clean = raw.lstrip("`, \n\t")
    if clean.lower().startswith("json"):
        clean = clean[4:].strip()
    return clean


def _parse_json_safe(raw: str, label: str = "json") -> dict | list | None:
    """Parsuje JSON z wielostopniowym fallbackiem — naprawia ucięte i niekompletne JSONy."""
    if not raw or len(raw.strip()) < 2:
        return None

    clean = _strip_json_markdown(raw)
    if not clean:
        return None

    # Próba 1: bezpośrednie parsowanie
    try:
        return json.loads(clean)
    except json.JSONDecodeError:
        pass

    # Próba 2: ekstrakcja pierwszego kompletnego obiektu (naprawia "Extra data")
    fragment = _extract_first_json_object(clean)
    if fragment:
        try:
            result = json.loads(fragment)
            logger.warning(
                "[%s] JSON Extra data — użyto _extract_first_json_object", label
            )
            return result
        except json.JSONDecodeError:
            pass

    # Próba 3: naprawa fragmentu zaczynającego się od ','
    # AI zwraca środek tablicy/obiektu bez otwierającego nawiasu
    repaired = clean.strip()
    if repaired.startswith(","):
        inner = repaired.lstrip(",").strip()
        if inner.startswith("{"):
            # Fragment listy obiektów → owijamy w [...]
            wrapped = "[" + inner
            opens = wrapped.count("{") - wrapped.count("}")
            if opens > 0:
                wrapped += "}" * opens
            if not wrapped.endswith("]"):
                wrapped += "]"
            try:
                result = json.loads(wrapped)
                logger.warning("[%s] JSON naprawiony (fragment tablicy → list)", label)
                return result
            except json.JSONDecodeError:
                pass
        elif inner.startswith('"'):
            # Fragment klucz-wartość → owijamy w {...}
            wrapped = "{" + inner
            if not wrapped.endswith("}"):
                wrapped += "}"
            try:
                result = json.loads(wrapped)
                logger.warning("[%s] JSON naprawiony (fragment obiektu → dict)", label)
                return result
            except json.JSONDecodeError:
                pass
    else:
        # Brak przecinka — próba uzupełnienia brakującego zamknięcia
        if not repaired.endswith("}") and not repaired.endswith("]"):
            if "{" in repaired:
                repaired += "}"
            elif "[" in repaired:
                repaired += "]"
        try:
            result = json.loads(repaired)
            logger.warning("[%s] JSON naprawiony (ucięty)", label)
            return result
        except Exception:
            pass

    # Próba 4: raw_decode — wyciąga wszystkie kompletne obiekty {} z fragmentu
    # Pomija luźne stringi (tylko dict/list się kwalifikują)
    decoder = json.JSONDecoder()
    collected = []
    scan_text = clean.lstrip(", \t\n")
    pos = 0
    while pos < len(scan_text):
        if scan_text[pos] not in ("{", "["):
            pos += 1
            continue
        try:
            obj, end = decoder.raw_decode(scan_text, pos)
            if isinstance(obj, (dict, list)):
                collected.append(obj)
            pos = end
            while pos < len(scan_text) and scan_text[pos] in ", \t\n":
                pos += 1
        except json.JSONDecodeError:
            pos += 1
    if collected:
        if len(collected) == 1:
            result = collected[0]
        else:
            # Wiele obiektów — sprawdź czy to lista jednorodnych dictów (pytania/dni)
            # czy różne sekcje do scalenia (różne klucze)
            dict_items = [o for o in collected if isinstance(o, dict)]
            list_items_flat = []
            for o in collected:
                if isinstance(o, list):
                    list_items_flat.extend(o)

            # Heurystyka: jednorodne dict-y → lista elementów (nie scalaj kluczami)
            all_dicts = len(dict_items) == len(collected) and len(dict_items) > 1
            if all_dicts:
                first_keys = frozenset(dict_items[0].keys())
                are_homogeneous = all(
                    frozenset(d.keys()) == first_keys for d in dict_items[1:]
                )
            else:
                are_homogeneous = False

            if are_homogeneous:
                # Jednorodne obiekty → wrzuć jako listę pod key_guess
                key_guess = {
                    "ankieta": "pytania",
                    "horoskop": "dni",
                    "gra": "pytania",
                }.get(label.split("-")[0].split("_")[0], "items")
                result = {key_guess: dict_items}
                logger.warning(
                    "[%s] raw_decode: %d jednorod. obiektów → {%s: [...]}",
                    label,
                    len(dict_items),
                    key_guess,
                )
            elif dict_items:
                # Różne sekcje — scalaj (oryginalne zachowanie)
                merged = {}
                for obj in dict_items:
                    merged.update(obj)
                if list_items_flat:
                    key_guess = {
                        "ankieta": "pytania",
                        "horoskop": "dni",
                        "gra": "pytania",
                    }.get(label.split("-")[0].split("_")[0], "items")
                    if key_guess not in merged:
                        merged[key_guess] = list_items_flat
                result = merged
            elif list_items_flat:
                key_guess = {
                    "ankieta": "pytania",
                    "horoskop": "dni",
                    "gra": "pytania",
                }.get(label.split("-")[0].split("_")[0], "items")
                result = {key_guess: list_items_flat}
            else:
                result = collected
        logger.warning(
            "[%s] JSON naprawiony (raw_decode: %d obiektów)", label, len(collected)
        )
        return result

    logger.warning("[%s] JSON nienaprawialny (raw_len=%d)", label, len(raw))
    return None


def _load_prompt_json() -> dict:
    """
    Wczytuje prompt.json z katalogu prompts/.
    Fallback: minimalny słownik jeśli plik nie istnieje.
    """
    try:
        with open(PROMPT_JSON_PATH, encoding="utf-8") as f:
            data = json.load(f)
        logger.info("[zwykly] prompt.json wczytany OK")
        return data
    except FileNotFoundError:
        logger.error(
            "[zwykly] Brak prompt.json: %s — używam fallbacku", PROMPT_JSON_PATH
        )
    except json.JSONDecodeError as e:
        logger.error("[zwykly] Błąd JSON w prompt.json: %s", e)
    return _fallback_prompt_dict()


def _fallback_prompt_dict() -> dict:
    """Minimalny fallback gdyby prompt.json był niedostępny."""
    return {
        "system": "Odpowiadaj WYŁĄCZNIE w formacie JSON bez żadnego tekstu poza klamrami {}.",
        "output_schema": {
            "odpowiedz_tekstowa": "...",
            "kategoria_pdf": "Manifest Wolności",
            "emocja": "radosc|smutek|zlosc|lek|nuda|spokoj",
        },
        "instrukcje": {
            "sokrates": "Odpowiedz mądrze, max 4 zdania, podpisz: Sokrates.",
            "tyler": "Styl nihilistyczny Fight Club. Podpisz: Tyler Durden.",
            "zasady_nota": "Dostosuj zasady twórczo do spraw nadawcy.",
        },
        "zasady_tylera": [
            "Pierwsza zasada: Nie mówi się o tym.",
            "Druga zasada: Nie mówi się o tym.",
            "Trzecia zasada: Jeśli ktoś zawoła stop, walka się kończy.",
            "Czwarta zasada: Walczą tylko dwaj faceci.",
            "Piąta zasada: Jedna walka naraz.",
            "Szósta zasada: Żadnych koszul, żadnych butów.",
            "Siódma zasada: Walki trwają tak długo jak muszą.",
            "Ósma zasada: Jeśli to twoja pierwsza noc, musisz walczyć.",
        ],
        "manifesty": [
            {
                "temat": "KONSUMPCJONIZM",
                "tresc": "Rzeczy, które posiadasz, w końcu zaczynają posiadać ciebie.",
            },
            {"temat": "HISTORIA", "tresc": "Jesteśmy średnimi dziećmi historii."},
            {"temat": "SAMODOSKONALENIE", "tresc": "Samodoskonalenie to masturbacja."},
            {"temat": "TOŻSAMOŚĆ", "tresc": "Nie jesteś swoją pracą."},
            {
                "temat": "PROJEKT CHAOS",
                "tresc": "Pewnego dnia umrzesz. Jesteś trybem w maszynie.",
            },
        ],
        "formatowanie_adresata": "Użyj formy: Drogi [Imię]-[Przymiotnik]-[Przydomek].",
        "user_text_placeholder": "{{USER_TEXT}}",
    }


def _render_prompt(
    data: dict, body: str, previous_body: str = None, sender_name: str = ""
) -> str:
    """
    Buduje pełny string promptu z danych prompt.json.
    Obsługuje zarówno stary format (instrukcje/zasady_tylera/manifesty)
    jak i nowy (tyler_zasady_OBOWIAZKOWE / tyler_manifesty_OBOWIAZKOWE).
    Obsługuje previous_body — poprzednią wiadomość od nadawcy.
    Hard constraints umieszczone NA POCZĄTKU — żeby nie zostały ucięte przy długich emailach.
    sender_name — imię nadawcy przekazane z GAS/webhook (priorytet nad autodetekcją).
    """
    lines = []

    # ── Hard constraints PIERWSZE — krytyczne zakazy na samym początku ────────
    hard = data.get("hard_constraints", [])
    if hard:
        lines.append("### BEZWZGLĘDNE ZAKAZY I WYMOGI — NARUSZENIE = BŁĘDNA ODPOWIEDŹ:")
        for h in hard:
            lines.append(f"- {h}")
        lines.append("")

    # ── System ───────────────────────────────────────────────────────────────
    lines.append(data.get("system", ""))
    lines.append("")

    # ── Schemat wyjściowy ─────────────────────────────────────────────────────
    schema = data.get("output_schema", {})
    if schema:
        lines.append("### SCHEMAT JSON DO WYPEŁNIENIA:")
        lines.append(json.dumps(schema, ensure_ascii=False, indent=2))
        lines.append("")

    # ── Poprzednia wiadomość (jeśli dostępna) ─────────────────────────────────
    if previous_body and previous_body.strip():
        lines.append(
            "### POPRZEDNIA WIADOMOŚĆ OD TEJ OSOBY (Tyler i Sokrates MUSZĄ do niej nawiązać):"
        )
        lines.append(previous_body[:2000])
        lines.append("")
        # Instrukcja nawiązania z prompt.json
        poprzednia_instr = data.get("tyler_poprzednia_wiadomosc", "")
        if poprzednia_instr:
            lines.append("### INSTRUKCJA NAWIĄZANIA DO POPRZEDNIEJ WIADOMOŚCI:")
            lines.append(poprzednia_instr)
            lines.append("")

    # ── Tekst użytkownika ─────────────────────────────────────────────────────
    lines.append("### OBECNA WIADOMOŚĆ OD NADAWCY (na jej podstawie generuj WSZYSTKO):")
    lines.append(body)
    lines.append("")
    # ── Imię nadawcy — kluczowe! ──────────────────────────────────────────────
    detected_name = _detect_sender_name(body) or sender_name or ""
    if detected_name:
        lines.append("### KRYTYCZNE — IMIĘ NADAWCY TEGO EMAILA:")
        lines.append(f"Osoba która NAPISAŁA ten email ma na imię: {detected_name}")
        lines.append(
            f"Tyler i Sokrates MUSZĄ zwracać się wyłącznie do '{detected_name}' — "
            f"ZAKAZ zwracania się do innych osób wymienionych w treści emaila "
            f"(np. jeśli w emailu jest 'Drogi Pawle', to Paweł jest adresatem emaila nadawcy, "
            f"NIE nadawcą do nas)."
        )
        lines.append("")

    # ── Sokrates ──────────────────────────────────────────────────────────────
    sokrates = (
        data.get("sokrates_instrukcja")
        or data.get("instrukcje_person", {}).get("sokrates")
        or data.get("instrukcje", {}).get("sokrates")
    )
    if sokrates:
        lines.append("### SOKRATES — INSTRUKCJA:")
        lines.append(sokrates)
        lines.append("")

    # ── Tyler — odmowa rekrutacji ─────────────────────────────────────────────
    odmowa = data.get("tyler_odmowa_rekrutacji") or data.get(
        "instrukcje_person", {}
    ).get("tyler", {}).get("zasada_rekrutacji")
    if odmowa:
        lines.append("### TYLER — ODMOWA REKRUTACJI (OBOWIĄZKOWE):")
        lines.append(odmowa)
        lines.append("")

    # ── Tyler — zasady (nowy format) ──────────────────────────────────────────
    zasady_obj = data.get("tyler_zasady_OBOWIAZKOWE", {})
    if zasady_obj:
        lines.append("### TYLER — 8 PUNKTÓW/DOGMATÓW (OBOWIĄZKOWE, KONKRETNE):")
        lines.append(zasady_obj.get("opis", ""))
        lines.append(f"WYMÓG ZASADA 1=2: {zasady_obj.get('zasada_1_2_identyczne', '')}")
        lines.append(f"FORMAT: {zasady_obj.get('format', '')}")
        lines.append(f"PRZYKŁAD ZŁY:   {zasady_obj.get('przyklad_zly', '')}")
        lines.append(f"PRZYKŁAD DOBRY: {zasady_obj.get('przyklad_dobry', '')}")
        lines.append("")
    else:
        # stary format
        zasady = data.get("zasady_tylera", [])
        inst = data.get("instrukcje", {})
        nota = inst.get("zasady_nota", "")
        if zasady:
            lines.append("### ELEMENTY DLA TYLERA (Wpleć w wypowiedź):")
            if nota:
                lines.append(nota)
            for z in zasady:
                lines.append(f"- {z}")
            lines.append("")

    # ── Tyler — manifesty (nowy format) ───────────────────────────────────────
    manifesty_obj = data.get("tyler_manifesty_OBOWIAZKOWE", {})
    if manifesty_obj:
        lines.append("### TYLER — 8 MANIFESTÓW (OBOWIĄZKOWE, KONKRETNE):")
        lines.append(manifesty_obj.get("opis", ""))
        for t in manifesty_obj.get("tematy", []):
            lines.append(f"- {t}")
        lines.append("")
    else:
        # stary format
        manifesty = data.get("manifesty", [])
        if manifesty:
            lines.append("### MANIFESTY TYLERA (Dostosuj i wygłoś każdy):")
            for i, m in enumerate(manifesty, 1):
                lines.append(f"{i}. O {m.get('temat', '???')}: {m.get('tresc', '')}")
            lines.append("")

    # ── Formatowanie adresata ─────────────────────────────────────────────────
    fmt = data.get("formatowanie_adresata", "")
    if fmt:
        lines.append("### FORMATOWANIE ADRESATA (OBOWIĄZKOWE):")
        # fmt może być dict (nowy JSON) lub str (stary format)
        if isinstance(fmt, dict):
            for k, v in fmt.items():
                lines.append(f"{k}: {v}")
        else:
            lines.append(fmt)
        lines.append("")

    # ── Końcowe przypomnienie ─────────────────────────────────────────────────
    lines.append("### PRZYPOMNIENIE PRZED GENEROWANIEM:")
    lines.append(
        "Każde zdanie Tylera MUSI nawiązywać do konkretnych słów z wiadomości nadawcy."
    )
    lines.append("ZAKAZ ogólnych rad, coachingu, pozytywnego myślenia, pocieszania.")
    lines.append("ZASADA 1 I ZASADA 2 MUSZĄ BYĆ IDENTYCZNE SŁOWO W SŁOWO.")
    lines.append("ADRESAT: ZAKAZ 'Drogi/Droga' — tylko forma wołacza jak w instrukcji.")
    lines.append("Zwróć WYŁĄCZNIE poprawny JSON bez żadnego tekstu poza klamrami.")
    lines.append("")

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════════
# DeepSeek — główny model
# ═══════════════════════════════════════════════════════════════════════════════


def _call_ai_with_fallback(
    system: str, user: str, max_tokens: int = 6000
) -> tuple[str | None, str]:
    """
    DeepSeek jako główny model.
    Zwraca (tekst_odpowiedzi, nazwa_providera).
    """
    # Używa tylko DeepSeek
    result = call_deepseek(system, user, MODEL_TYLER, max_tokens=max_tokens)
    if result:
        return result, "deepseek"
    logger.error("[zwykly] DeepSeek zawiódł!")
    return None, "none"


# ═══════════════════════════════════════════════════════════════════════════════
# PARSOWANIE ODPOWIEDZI MODELU
# ═══════════════════════════════════════════════════════════════════════════════
def _clean_manifest_labels(text: str) -> str:
    """
    Usuwa etykiety manifestów które model wypisuje mimo zakazu.
    np. "KONSUMPCJONIZM: treść" → "treść"
    """
    if not text:
        return text
    labels = [
        "KONSUMPCJONIZM",
        "DNO",
        r"DNO \(Rock Bottom\)",
        r"BÓG/RELIGIA",
        "BÓG",
        "RELIGIA",
        "KLASA ROBOTNICZA",
        r"ŚMIERTELNOŚĆ",
        r"ODPUSZCZENIE \(Let Go\)",
        "ODPUSZCZENIE",
        "AUTENTYCZNOŚĆ",
        "ILUZJA BEZPIECZEŃSTWA",
        "HISTORIA",
        "SAMODOSKONALENIE",
        "TOŻSAMOŚĆ",
        "RYZYKO",
        "BUNT",
        "KONTROLA",
    ]
    pattern = r"^(?:" + "|".join(labels) + r")\s*:\s*"
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        cleaned.append(re.sub(pattern, "", line, flags=re.IGNORECASE))
    return "\n".join(cleaned)


def _parse_response(raw: str) -> tuple[str, str]:
    if not raw or not raw.strip():
        return "", ""

    data = _parse_json_safe(raw, "zwykly-parse")
    if data is None:
        logger.warning("[zwykly] JSON parse error, fallback to raw text")
        return raw.strip(), ""
    if not isinstance(data, dict):
        return raw.strip(), ""

    text = data.get("odpowiedz_tekstowa") or data.get("odpowiedz") or ""
    emotion = data.get("emocja") or data.get("emotion") or ""
    return str(text).strip(), str(emotion).strip()


# ═══════════════════════════════════════════════════════════════════════════════
# EMOTKA + PDF
# ═══════════════════════════════════════════════════════════════════════════════


def _get_emoticon_and_pdf(emotion_key: str) -> tuple:
    """Zwraca (png_b64, pdf_b64) dla danej emocji z fallbackiem na error."""
    png_b64 = read_file_base64(os.path.join(EMOTKI_DIR, f"{emotion_key}.png"))
    pdf_b64 = read_file_base64(os.path.join(PDF_DIR, f"{emotion_key}.pdf"))

    if not png_b64:
        logger.warning("[zwykly] Brak PNG dla %s, używam error.png", emotion_key)
        png_b64 = read_file_base64(os.path.join(EMOTKI_DIR, f"{FALLBACK_EMOT}.png"))
    if not pdf_b64:
        logger.warning("[zwykly] Brak PDF dla %s, używam error.pdf", emotion_key)
        pdf_b64 = read_file_base64(os.path.join(PDF_DIR, f"{FALLBACK_EMOT}.pdf"))

    return png_b64, pdf_b64


# ═══════════════════════════════════════════════════════════════════════════════
# ŁADOWANIE WYTYCZNYCH STYLU (zwykly_obrazek_tyler.js)
# ═══════════════════════════════════════════════════════════════════════════════


def _load_style_config() -> dict:
    """
    Wczytuje STYLE_CONFIG z pliku zwykly_panel_wytyczne.json.
    Czyta klucz STYLE_CONFIG bezpośrednio z JSON.
    """
    try:
        with open(STYLE_JS_PATH, encoding="utf-8") as f:
            data = json.load(f)
        config = data.get("STYLE_CONFIG", {})
        if not config:
            logger.warning("[zwykly-img] Brak bloku STYLE_CONFIG w %s", STYLE_JS_PATH)
            return {}
        logger.info("[zwykly-img] Wczytano STYLE_CONFIG OK")
        return config
    except FileNotFoundError:
        logger.warning("[zwykly-img] Brak pliku %s", STYLE_JS_PATH)
    except json.JSONDecodeError as e:
        logger.error("[zwykly-img] Błąd parsowania STYLE_CONFIG: %s", e)
    return {}


# ═══════════════════════════════════════════════════════════════════════════════
# KONFIGURACJA POSTACI, STYLÓW, AKCJI
# ═══════════════════════════════════════════════════════════════════════════════

FIGHT_CLUB_CHARACTERS = [
    # Tyler Durden - Esencja chaosu
    "Brad Pitt as Tyler Durden — raw, feral intensity. Post-fight appearance: blood-caked knuckles, a chipped front tooth, and a deep gash over a swollen eye. Wearing a scuffed, dirty red leather jacket over a bare, sweat-glistening chest marked with chemical soap-burn scars. His hair is a greasy, matted mess. He’s holding a smoldering cigarette, standing amidst the wreckage of a burned-out house. Hyper-realistic, 35mm film grain, 1990s grime.",
    # Narrator (Norton) - Symbol totalnego rozkładu psychicznego
    "Edward Norton as the Narrator — the look of total insomnia. Sunken, charcoal-rimmed eyes, pale sickly skin with visible veins. Wearing a sweat-stained, tattered white dress shirt with the sleeves ripped off, covered in dried blood and office coffee stains. A massive purple hematoma on his cheekbone and a split lip. He looks completely dissociated and broken, staring into the camera with a 'thousand-yard stare'. Dark, moody lighting.",
    # Marla Singer - Zniszczona panna młoda (zgodnie z prośbą)
    "Helena Bonham Carter as Marla Singer — wearing a shredded, soot-covered vintage bridesmaid/wedding dress from a thrift store. Her hair is an unwashed, bird's-nest tangle. Smudged, heavy black 'raccoon' eye makeup running down her face. She’s leaning against a peeling wallpaper wall in a derelict hallway, a pink feather boa hanging like a dead animal around her neck. Nihilistic smirk, blood on her teeth, holding a cigarette with trembling, ash-covered fingers.",
    # Angel Face (Jared Leto) - Zniszczone piękno
    "Jared Leto as Angel Face — once-ethereal, angelic features now pulverized into a pulp of gore. Both eyes swollen shut, nose shattered and crooked, blood dripping from a ruined mouth. His platinum blonde hair is soaked in crimson. A haunting contrast between his delicate bone structure and the absolute brutality of the beating. Extreme close-up, harsh fluorescent lighting.",
    # Bob - Rozpacz i fizyczna masa
    "Meat Loaf as Bob — a mountain of a man in a state of emotional collapse. Wearing a massive, sweat-drenched, grey XXXL sweatshirt. Tear-streaked face, puffy eyes, and the visible shape of gynecomastia. He looks like a tragic, broken giant. Surroundings: a dark, damp basement with cracked concrete and single bare lightbulb casting long, dramatic shadows.",
]

PANEL_STYLES = [
    "35mm film grain, high contrast, sickly green and amber tones, Fincher cinematography",
    "raw gritty street photography, harsh fluorescent light, 1990s documentary style",
    "extreme chiaroscuro, single bare bulb lighting, deep shadows, industrial decay",
    "handheld camera blur, motion, chaotic energy, smoke and sweat",
    "desaturated noir, cold blue shadows, cracked concrete textures",
    "overexposed bleach bypass, washed out whites, dark crushed blacks",
]

PANEL_ACTIONS = [
    # Kultowa scena z samochodem
    "releasing steering wheel of a speeding car, hands off, smiling maniacally while oncoming headlights reflect in glazed eyes, 35mm motion blur, chaos",
    # Marla w sukni ślubnej (zgodnie z Twoją prośbą)
    "standing in a scorched, derelict ballroom, arms spread wide in a ruined wedding dress, face turned toward black soot and smoke, liberated and destroyed ",
    # Scena z Raymondem K. Hesselem (pistolet do głowy/konfrontacja)
    "crouching over a terrified clerk pinned against a dumpster in a rain-slicked alley, forcing them to confront their meaningless life, steam rising from grates, rats scurrying in shadows",
    # Nihilizm konsumpcyjny
    "laughing maniacally with blood-caked teeth, standing amidst a bonfire of burning IKEA furniture and designer catalogs, high contrast",
    # Rock Bottom (Narrator)
    "sitting at the bottom of a dark, wet rocky pit, staring up at a tiny square of grey sky with hollow eyes, personifying 'hitting rock bottom' [cite: 22]",
    # Portret wściekłości
    "screaming directly into the lens with veins bulging on the neck, face inches from camera, splattered with sweat and grime, raw rage and contempt ",
    # Pisanie krwią/mydłem
    "writing a nihilistic manifesto on a cracked wall with bloody knuckles, chemical smoke and lye dust in the background, industrial setting",
    # Zniszczenie życia nadawcy (Jadzi)
    "standing over a pile of burning truskawka-themed objects and 12-zloty notes, pointing a judgmental finger at the camera, cold lighting",
    # Wyjście z wypadku
    "walking away from a twisted, flaming car wreck in slow motion, face smeared with blood, looking dead ahead without blinking, fire illuminating the night ",
    # Scena w kościele (z logu debug)
    "reading from a burning book in a dimly lit, empty church, surrounded by a congregation of rats, amidst 35mm film grain and heavy shadows ",
]


def _extract_nouns_from_body(body: str) -> list:
    """
    Wyciąga rzeczowniki/konkretne obiekty z treści emaila.
    Szuka słów pisanych z wielkiej litery (imiona, miejsca) oraz
    typowych rzeczowników codziennych.
    Zwraca listę max 6 słów.
    """
    # Słowa które zawsze wyrzucamy (stopwords)
    stopwords = {
        "się",
        "nie",
        "jak",
        "ale",
        "czy",
        "też",
        "już",
        "aby",
        "żeby",
        "tego",
        "tej",
        "ten",
        "tak",
        "jest",
        "był",
        "być",
        "mam",
        "mieć",
        "to",
        "i",
        "w",
        "z",
        "na",
        "do",
        "po",
        "za",
        "od",
        "przez",
        "że",
        "co",
        "gdy",
        "więc",
        "bo",
        "dla",
        "przy",
        "nad",
        "pod",
        "mój",
        "moja",
        "moje",
        "jego",
        "jej",
        "ich",
        "swój",
        "twój",
        "wszystko",
        "tylko",
        "jeszcze",
        "bardzo",
        "bardziej",
        "może",
        "chcę",
        "musi",
        "można",
        "który",
        "która",
        "które",
    }
    words = re.findall(r"[A-Za-zżźćńółęąśŻŹĆŃÓŁĘĄŚ]{4,}", body)
    seen = set()
    nouns = []
    for w in words:
        wl = w.lower()
        if wl not in stopwords and wl not in seen:
            seen.add(wl)
            nouns.append(w)
        if len(nouns) >= 6:
            break
    return nouns


def _extract_nouns_deepseek(body: str) -> dict:
    """
    Wysyła email do DeepSeek SEKWENCYJNIE (klucze po kolei).
    Zwraca dict {rzecz001: 'kopalnia', ...} lub {} przy błędzie.
    """
    NOUNS_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_znajdz_rzeczowniki.json")
    try:
        with open(NOUNS_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[rzeczowniki] Brak zwykly_znajdz_rzeczowniki.json: %s", e)
        return {}

    system_msg = cfg.get("system", "")
    user_prefix = cfg.get("user_prefix", "Wypisz WSZYSTKIE rzeczowniki z tekstu:\n")
    max_tokens = cfg.get("max_tokens", 3000)
    temperature = cfg.get("temperature", 0.1)
    user_msg = user_prefix + (body or "")

    raw = call_deepseek(system_msg, user_msg, MODEL_TYLER)

    if not raw:
        logger.error("[rzeczowniki] Brak odpowiedzi od AI")
        return {}

    # Parsuj JSON — obsługa ```json...``` wszędzie w tekście
    try:
        clean = _strip_json_markdown(raw)
        # Guard: jeśli model zwrócił tekst zamiast JSON
        if not clean.strip().startswith("{"):
            logger.warning(
                "[rzeczowniki] Odpowiedź nie jest JSON — fallback regex | raw: %.120s",
                raw,
            )
            return {}
        # Użyj raw_decode zamiast json.loads — obsługuje "Extra data"
        decoder = json.JSONDecoder()
        result = None
        try:
            result, _ = decoder.raw_decode(clean)
        except json.JSONDecodeError:
            for match in re.finditer(r"[{\[]", clean):
                start = match.start()
                try:
                    obj, end = decoder.raw_decode(clean[start:])
                    if obj is not None:
                        result = obj
                        break
                except json.JSONDecodeError:
                    continue
        if result is None:
            raise ValueError("Nie znaleziono JSON")
        if not isinstance(result, dict):
            raise ValueError(f"Oczekiwano dict, dostałem {type(result).__name__}")
        nouns_dict = {
            k: v
            for k, v in result.items()
            if re.match(r"^rzecz\d+$", k) and isinstance(v, str)
        }
        logger.info("[rzeczowniki] OK — %d rzeczowników", len(nouns_dict))
        return nouns_dict
    except Exception as e:
        logger.warning("[rzeczowniki] Błąd JSON: %s | raw: %.200s", e, raw)
        return {}


def _append_nouns_to_debug_txt(debug_txt_dict: dict, nouns_dict: dict) -> dict:
    """
    Dopisuje listę rzeczowników na końcu pliku _.txt (base64).
    Zwraca zaktualizowany dict debug_txt.
    """
    if not debug_txt_dict or not nouns_dict:
        return debug_txt_dict
    try:
        existing = base64.b64decode(debug_txt_dict["base64"]).decode("utf-8")
        lines = [
            "",
            "---------------------------------------------",
            "RZECZOWNIKI Z EMAILA (zwykly_znajdz_rzeczowniki.json)",
            "---------------------------------------------",
        ]
        for k in sorted(nouns_dict.keys()):
            lines.append(f"  {k} = {nouns_dict[k]}")
        lines.append("")
        appended = existing + "\n".join(lines)
        debug_txt_dict["base64"] = base64.b64encode(appended.encode("utf-8")).decode(
            "ascii"
        )
        logger.info("[rzeczowniki] Dopisano %d rzeczowników do _.txt", len(nouns_dict))
    except Exception as e:
        logger.warning("[rzeczowniki] Błąd dopisywania do _.txt: %s", e)
    return debug_txt_dict


def _detect_sender_name(body: str) -> str | None:
    """
    Próbuje wykryć imię nadawcy z treści emaila.
    Szuka podpisu na końcu lub zwrotu do siebie w pierwszej osobie.
    Zwraca imię lub None.
    """
    # Szukaj podpisu: linia z jednym słowem zaczynającym się wielką literą
    # na końcu wiadomości
    lines = [l.strip() for l in body.strip().splitlines() if l.strip()]
    for line in reversed(lines[-5:]):
        m = re.match(r"^([A-ZŁŻŹĆŃÓĘĄŚ][a-złżźćńóęąś]{2,12})$", line)
        if m:
            return m.group(1)

    # Szukaj "Pozdrawiam, Imię" lub "— Imię"
    m = re.search(
        r"(?:pozdrawiam|pozdrowienia|z poważaniem|regards)[,\s]+([A-ZŁŻŹĆŃÓĘĄŚ][a-złżźćńóęąś]{2,12})",
        body,
        re.IGNORECASE,
    )
    if m:
        return m.group(1)

    m = re.search(r"(?:^|\n)[—–-]\s*([A-ZŁŻŹĆŃÓĘĄŚ][a-złżźćńóęąś]{2,12})", body)
    if m:
        return m.group(1)

    return None


def _detect_gender(body: str, sender_name: str = "") -> str:
    """
    Wykrywa płeć nadawcy na podstawie treści emaila i sender_name.
    Kolejność:
      1. Regex na końcówkach czasowników/przymiotników w body
      2. DeepSeek — zapytanie o płeć na podstawie body + sender_name
      3. Fallback: 'nieznana'
    Zwraca 'kobieta', 'mezczyzna' lub 'nieznana'.
    """
    if not body and not sender_name:
        return "nieznana"

    text = (body or "").lower()

    # ── 1. Regex na końcówkach gramatycznych ─────────────────────────────────
    zenskie = [
        r"\bjeste[mś]\s+\w*a\b",
        r"\bby[łl]am\b",
        r"\bposz[łl]am\b",
        r"\bpracowa[łl]am\b",
        r"\bchcia[łl]am\b",
        r"\bpisa[łl]am\b",
        r"\bzrobi[łl]am\b",
        r"\bprzysz[łl]am\b",
        r"\bmia[łl]am\b",
        r"\bdosta[łl]am\b",
        r"\bwysz[łl]am\b",
        r"\bzmęczona\b",
        r"\bszczęśliwa\b",
        r"\bzdenerwowana\b",
        r"\bprzejęta\b",
        r"\bpoczułam\b",
        r"\bpani\b",
    ]
    meskie = [
        r"\bby[łl]em\b",
        r"\bposzed[łl]em\b",
        r"\bpracowa[łl]em\b",
        r"\bchcia[łl]em\b",
        r"\bpisa[łl]em\b",
        r"\bzrobi[łl]em\b",
        r"\bprzysz[łl]em\b",
        r"\bmia[łl]em\b",
        r"\bdosta[łl]em\b",
        r"\bwysz[łl]em\b",
        r"\bzmęczony\b",
        r"\bszczęśliwy\b",
        r"\bzdenerwowany\b",
        r"\bpoczułem\b",
        r"\bpan\b",
    ]

    score_k = sum(1 for p in zenskie if re.search(p, text))
    score_m = sum(1 for p in meskie if re.search(p, text))

    if score_k > score_m:
        return "kobieta"
    elif score_m > score_k:
        return "mezczyzna"

    # ── 2. Fallback: detekcja z końcówki imienia ──────────────────────────────
    if sender_name:
        imie = sender_name.split()[0].lower()
        # Typowe polskie imiona żeńskie kończą się na -a (Monika, Anna, Kasia...)
        # Wyjątki męskie na -a: Kuba, Barnaba, Kosma — mała lista
        meskie_na_a = {"kuba", "barnaba", "kosma", "bonawentura", "sasha", "misza"}
        if imie.endswith("a") and imie not in meskie_na_a:
            return "kobieta"
        # Imiona zakończone na spółgłoskę lub -o/-u → zazwyczaj męskie
        if imie and imie[-1] not in "aąę":
            return "mezczyzna"

    return "nieznana"


def _add_text_below_image(image_obj: dict, text: str, panel_index: int) -> dict:
    """
    Rozszerza obrazek o 18% na dole i dopisuje tekst Pillow.
    Zwraca nowy dict z zaktualizowanym base64/filename.
    """
    try:
        from PIL import Image, ImageDraw, ImageFont

        raw = base64.b64decode(image_obj["base64"])
        img = Image.open(io.BytesIO(raw)).convert("RGB")
        W, H = img.size

        # Pasek na dole — 18% wysokości, min 80px
        bar_h = max(80, int(H * 0.18))
        new_img = Image.new("RGB", (W, H + bar_h), (10, 10, 10))
        new_img.paste(img, (0, 0))

        draw = ImageDraw.Draw(new_img)

        PADDING = 24
        max_w = W - PADDING * 2

        def load_font(size):
            for font_path in [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
            ]:
                try:
                    return ImageFont.truetype(font_path, size)
                except Exception:
                    continue
            return ImageFont.load_default()

        def wrap_text(txt, fnt, max_px):
            words = txt.split()
            lines_out = []
            current = ""
            for word in words:
                test = (current + " " + word).strip()
                bbox = draw.textbbox((0, 0), test, font=fnt)
                if bbox[2] - bbox[0] <= max_px:
                    current = test
                else:
                    if current:
                        lines_out.append(current)
                    current = word
            if current:
                lines_out.append(current)
            return lines_out

        # Dobierz font_size tak żeby tekst zmieścił się w max 4 liniach w pasku
        font_size = max(10, bar_h // 4)
        for attempt in range(14):
            font = load_font(font_size)
            lines_out = wrap_text(text, font, max_w)
            line_h = font_size + 6
            total_h = len(lines_out) * line_h
            if total_h <= bar_h - 8 and len(lines_out) <= 4:
                break
            font_size = max(10, font_size - 2)

        lines_out = lines_out[:4]

        # Rysuj tekst — wyśrodkowany w pasku
        line_h = font_size + 6
        total_text_h = len(lines_out) * line_h
        y = H + (bar_h - total_text_h) // 2
        for line in lines_out:
            bbox = draw.textbbox((0, 0), line, font=font)
            tw = bbox[2] - bbox[0]
            x = (W - tw) // 2
            # cień
            draw.text((x + 1, y + 1), line, font=font, fill=(0, 0, 0))
            draw.text((x, y), line, font=font, fill=(220, 210, 180))
            y += line_h

        buf = io.BytesIO()
        new_img.save(buf, format="JPEG", quality=TYLER_JPG_QUALITY, optimize=True)
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tyler_{ts}_panel{panel_index}_txt.jpg"

        result = dict(image_obj)
        result["base64"] = b64
        result["filename"] = filename
        result["size_jpg"] = f"{len(buf.getvalue()) // 1024}KB"
        result["caption"] = text
        return result

    except Exception as e:
        logger.warning("[tyler-txt] Błąd dopisywania tekstu: %s", e)
        return image_obj


# ═══════════════════════════════════════════════════════════════════════════════
# GENEROWANIE PROMPTÓW DLA TRYPTYKU (DeepSeek)
# ═══════════════════════════════════════════════════════════════════════════════


def _extract_tyler_sentences(response_text: str) -> dict:
    """
    Wyciąga gotowe zdania z odpowiedzi Tylera do użycia w dymkach tryptyku.
    Priorytetyzuje najbardziej agresywne zdania — o Bogu, śmierci, dnie.
    Zwraca dict:
      panel1 — pierwsza zasada (identyczna 1=2)
      panel2 — manifest DNO/BÓG/ŚMIERTELNOŚĆ (priorytet nihilistyczny)
      panel3 — okrzyk końcowy lub ostatnie zdanie Tylera
    """
    if not response_text:
        return {
            "panel1": "Nie mówi się o tym.",
            "panel2": "Bóg cię nie lubi. Prawdopodobnie cię nienawidzi.",
            "panel3": "Puść kierownicę. Pozwól sobie na wypadek.",
        }

    # Wytnij sekcję Tylera
    tyler_section = response_text
    if "### TYLER DURDEN" in response_text:
        tyler_section = response_text.split("### TYLER DURDEN", 1)[1]

    lines = [l.strip() for l in tyler_section.splitlines() if l.strip()]

    # Panel 1 — pierwsza zasada w stylu Fight Club
    panel1 = None
    ordinal_re = re.compile(
        r"^(pierwsza|druga|trzecia|czwarta|pi[aą]ta|sz[oó]sta|si[oó]dma|[oó]sma)\s+zasada",
        re.IGNORECASE,
    )
    for line in lines:
        if ordinal_re.match(line):
            panel1 = line[:120]
            break
    if not panel1:
        for line in lines:
            if re.match(r"^[1-8][.)]", line):
                panel1 = re.sub(r"^[1-8][.)]\s*", "", line)[:120]
                break
    if not panel1:
        panel1 = "Pierwsza zasada: nie mówi się o tym."

    # Panel 2 — priorytet: DNO, BÓG, ŚMIERTELNOŚĆ, ODPUSZCZENIE (nihilistyczne)
    panel2 = None
    nihilist_priority = [
        "DNO",
        "BÓG",
        "ŚMIERTELNOŚĆ",
        "ODPUSZCZENIE",
        "AUTENTYCZNOŚĆ",
        "ILUZJA",
    ]
    for priority_word in nihilist_priority:
        for line in lines:
            if line.upper().startswith(priority_word):
                panel2 = line[:140]
                break
        if panel2:
            break
    # fallback: pierwsza linia z CAPS tematem manifestu
    if not panel2:
        for line in lines:
            if re.match(r"^[A-ZŻŹĆĄŚĘÓŁŃ]{4,}[\s:]", line):
                panel2 = line[:140]
                break
    if not panel2:
        for line in lines:
            if line.startswith("- ") and len(line) > 15:
                panel2 = line[2:][:140]
                break
    if not panel2:
        panel2 = "Bóg cię nie lubi. Jesteś niechcianym produktem historii."

    # Panel 3 — okrzyk końcowy lub ostatnie zdanie
    panel3 = None
    for line in lines:
        if "okrzyk" in line.lower():
            panel3 = re.sub(r"okrzyk[^:]*:\s*", "", line, flags=re.IGNORECASE).strip()[
                :120
            ]
            break
    if not panel3 and lines:
        for line in reversed(lines):
            if (
                line
                and not line.startswith("---")
                and not line.startswith("###")
                and len(line) > 15
            ):
                panel3 = line[:120]
                break
    if not panel3:
        panel3 = "Puść kierownicę. Pozwól sobie na wypadek."

    return {"panel1": panel1, "panel2": panel2, "panel3": panel3}


def _extract_tyler_rules(response_text: str) -> list:
    """
    Wyciąga 8 zasad Tylera z tekstu odpowiedzi AI.
    Zwraca listę stringów (max 8 zasad).
    Panel 1 = zasady 1+2 (identyczne), panele 2-7 = zasady 2-7 (indeks 1-6),
    panel 7 = zasada 8.
    Tak naprawdę zwracamy listę 7 zasad do 7 paneli:
      panel_rules[0] = zasada 1 (i 2 — są identyczne)
      panel_rules[1] = zasada 3
      ...
      panel_rules[6] = zasada 8
    """
    if not response_text:
        return []

    tyler_section = response_text
    if "### TYLER DURDEN" in response_text:
        tyler_section = response_text.split("### TYLER DURDEN", 1)[1]

    ordinal_map = {
        "pierwsza": 1,
        "druga": 2,
        "trzecia": 3,
        "czwarta": 4,
        "piąta": 5,
        "piata": 5,
        "szósta": 6,
        "szosta": 6,
        "siódma": 7,
        "siodma": 7,
        "ósma": 8,
        "osma": 8,
    }

    rules = {}
    lines = [l.strip() for l in tyler_section.splitlines() if l.strip()]
    for line in lines:
        m = re.match(
            r"^(pierwsza|druga|trzecia|czwarta|pi[aą]ta|sz[oó]sta|si[oó]dma|[oó]sma)\s+zasada",
            line,
            re.IGNORECASE,
        )
        if m:
            ordinal = m.group(1).lower().replace("ó", "o").replace("ą", "a")
            idx = ordinal_map.get(ordinal)
            if idx and idx not in rules:
                rules[idx] = line.strip()

    # Fallback: szukaj linii z numerem 1. 2. itd.
    if len(rules) < 4:
        for line in lines:
            m = re.match(r"^([1-8])[.)]\s*(.+)", line)
            if m:
                idx = int(m.group(1))
                if idx not in rules:
                    rules[idx] = line.strip()

    # Zbuduj listę 7 paneli: panel1=zasada1, panel2=zasada3, ..., panel7=zasada8
    # Mapowanie panel → numer zasady
    # Panel 1 = zasada 1 (i 2 — są identyczne)
    # Panel 2 = zasada 3
    # Panel 3 = zasada 4
    # Panel 4 = zasada 5
    # Panel 5 = zasada 6
    # Panel 6 = zasada 7
    # Panel 7 = zasada 8
    panel_to_rule = {1: 1, 2: 3, 3: 4, 4: 5, 5: 6, 6: 7, 7: 8}
    panel_rules = []
    for p in range(1, 8):
        rule_idx = panel_to_rule[p]
        rule_text = rules.get(rule_idx, rules.get(p, ""))
        panel_rules.append(rule_text)

    logger.info(
        "[tyler-rules] Wyciągnięto %d zasad z tekstu Tylera → %d paneli",
        len(rules),
        len([r for r in panel_rules if r]),
    )
    return panel_rules


PANEL_WYTYCZNE_JSON_PATH = os.path.join(PROMPTS_DIR, "zwykly_panel_wytyczne.json")


def _load_panel_wytyczne() -> dict:
    """
    Wczytuje wytyczne do generowania paneli z zwykly_panel_wytyczne.json.
    Wszystkie wytyczne stylistyczne, system prompt AI, szablony user promptu
    i logika odwrócenia są tam — Python nic nie hardkoduje.
    Fallback: minimalny dict jeśli plik niedostępny.
    """
    try:
        with open(PANEL_WYTYCZNE_JSON_PATH, encoding="utf-8") as f:
            data = json.load(f)
        logger.info("[panel-wytyczne] Wczytano zwykly_panel_wytyczne.json OK")
        return data
    except FileNotFoundError:
        logger.error(
            "[panel-wytyczne] Brak pliku %s — używam fallbacku",
            PANEL_WYTYCZNE_JSON_PATH,
        )
    except json.JSONDecodeError as e:
        logger.error("[panel-wytyczne] Błąd JSON w zwykly_panel_wytyczne.json: %s", e)
    return {
        "system_prompt_AI": (
            "You are a cinematic visual prompt engineer for FLUX image generation. "
            "Fight Club 1999 aesthetic. Given a Tyler Durden RULE, generate a scene showing "
            "the VIOLATION of that rule. Characters actively do what the rule forbids. "
            "Output: ONE paragraph, max 120 words, English only, just the FLUX prompt."
        ),
        "user_prompt_szablon": (
            "TYLER'S RULE (panel [PANEL_NR]/7):\n[ZASADA_TEKST]\n\n"
            "Objects from sender's email: [USER_OBJECTS]\n"
            "Generate a FLUX prompt showing the VIOLATION of this rule:"
        ),
        "style_variants": [
            "35mm film grain, high contrast, Fight Club 1999, David Fincher"
        ],
        "fallback_gdy_brak_zasady": (
            "Tyler Durden walking away from a burning wreck, 35mm film grain, "
            "Fight Club 1999 aesthetic, David Fincher, underexposed, gritty"
        ),
        "styl_globalny": {
            "zakazy_negatywne": "clean, polished, beautiful, anime, text, watermark"
        },
    }


# _generate_panel_prompt_from_rule usunięta — zastąpiona przez _generate_triptych_prompts_batch
# (1 call DeepSeek dla wszystkich 7 paneli naraz zamiast 7 osobnych calli)


def _detect_city(body: str) -> str:
    """
    Wykrywa miasto/miejscowość z treści emaila.
    Szuka znanych polskich miast oraz słów 'w [Miasto]', 'z [Miasto]'.
    """
    if not body:
        return ""
    known = [
        "Warszawa",
        "Kraków",
        "Wrocław",
        "Poznań",
        "Gdańsk",
        "Łódź",
        "Szczecin",
        "Bydgoszcz",
        "Lublin",
        "Katowice",
        "Białystok",
        "Gdynia",
        "Częstochowa",
        "Radom",
        "Sosnowiec",
        "Toruń",
        "Kielce",
        "Rzeszów",
        "Gliwice",
        "Zabrze",
        "Bogatynia",
        "Legnica",
        "Opole",
        "Zielona Góra",
        "Olsztyn",
        "Płock",
    ]
    for city in known:
        if city.lower() in body.lower():
            return city
    m = re.search(
        r"\b(?:w|z|do|ze|pod|nad|koło|przy)\s+([A-ZŁŻŹĆŃÓĘĄŚ][a-złżźćńóęąś]{3,})", body
    )
    if m:
        return m.group(1)
    return ""


def _detect_job(body: str) -> str:
    """
    Wykrywa zawód/profesję z treści emaila.
    Szuka typowych słów kluczowych.
    """
    if not body:
        return ""
    patterns = [
        r"\bpracuję\s+(?:jako|na\s+stanowisku)\s+([a-złżźćńóęąś\s]{3,60})",
        r"\bjeste[mś]\s+([a-złżźćńóęąś]{4,20}(?:em|iem|ą)?)\b",
        r"\bzawód[:\s]+([a-złżźćńóęąś\s]{3,25})",
        r"\binspektor\b",
        r"\binżynier\b",
        r"\bnauczyciel\b",
        r"\blekarz\b",
        r"\bkierowca\b",
        r"\bprogramista\b",
        r"\bksięgow\w+\b",
        r"\bsprzedaw\w+\b",
        r"\bpielęgniark\w+\b",
        r"\bstrażak\b",
        r"\bpolicjant\b",
        r"\bgórnik\b",
        r"\bdyrektor\b",
        r"\bprezes\b",
        r"\bmenedżer\b",
        r"\barchitekt\b",
    ]
    for p in patterns:
        m = re.search(p, body, re.IGNORECASE)
        if m:
            if m.lastindex:
                return m.group(1).strip()
            return m.group(0).strip()
    return ""


def _split_into_sentences(text: str) -> list:
    """
    Dzieli tekst na zdania. Pomija nagłówki (###), separatory (---),
    podpisy (— Sokrates) i linie krótsze niż 20 znaków.
    Zwraca listę zdań jako stringów.
    """
    if not text:
        return []
    sentences = []
    # Podziel po . ! ? ale nie po skrótach
    raw = re.split(r"(?<=[.!?])\s+", text)
    for s in raw:
        s = s.strip()
        if not s:
            continue
        if s.startswith("#") or s.startswith("—") or s.startswith("-"):
            continue
        if len(s) < 20:
            continue
        sentences.append(s)
    return sentences


def _build_session_vars(
    body: str,
    sender_email: str,
    sender_name: str,
    previous_body: str,
    res_text: str,
    emotion_key: str,
    provider: str,
    panel_assignments: list = None,
    nouns_dict: dict = None,
) -> dict:
    """
    Buduje słownik WSZYSTKICH zmiennych globalnych sesji.
    Klucze = nazwy zmiennych bez nawiasów kwadratowych.
    Wartości = stringi gotowe do podstawienia.

    Zmienne z GAS/webhook:
      SENDER, SENDER_NAME, BODY, PREVIOUS_BODY

    Wykryte z emaila:
      USER_PERSON, USER_OBJECTS, USER_GENDER, USER_CITY, USER_JOB, USER_EMOTION, USER_PROVIDER
      USER_OBJECTS pochodzi z nouns_dict (DeepSeek) jeśli dostępny, fallback na regex.

    Ze zdań Tylera:
      TEXT_1 .. TEXT_N

    Ze zdań Sokratesa:
      SOKRATES_1 .. SOKRATES_N
    """
    vars_dict = {}

    # ── Zmienne z webhook / GAS ───────────────────────────────────────────────
    vars_dict["SENDER"] = sender_email or ""
    vars_dict["SENDER_NAME"] = sender_name or ""
    vars_dict["BODY"] = body or ""
    # Jeśli PREVIOUS_BODY jest identyczny z BODY — to błąd webhooka, traktuj jako brak historii
    _prev = previous_body or ""
    if _prev.strip() and _prev.strip() == (body or "").strip():
        logger.warning(
            "[session_vars] PREVIOUS_BODY identyczny z BODY — traktuję jako brak historii"
        )
        _prev = ""
    vars_dict["PREVIOUS_BODY"] = _prev

    # ── USER_OBJECTS: DeepSeek nouns_dict (priorytet) → fallback regex ──────────
    if nouns_dict and isinstance(nouns_dict, dict):
        # nouns_dict = {rzecz001: 'kopalnia', rzecz002: 'pies', ...}
        # Bierzemy wartości w kolejności kluczy, max 15
        sorted_nouns = [v for k, v in sorted(nouns_dict.items()) if isinstance(v, str)]
        vars_dict["USER_OBJECTS"] = ", ".join(sorted_nouns[:15])
    else:
        nouns = _extract_nouns_from_body(body)
        vars_dict["USER_OBJECTS"] = ", ".join(nouns[:15]) if nouns else ""
    # KLUCZOWE: sender_name z webhooka (GAS) ma ABSOLUTNY PRIORYTET.
    # _detect_sender_name(body) wykrywa imię z TREŚCI emaila (np. "Mama / Anna") —
    # to jest imię osoby PODPISANEJ pod listem, nie nadawcy wiadomości do systemu.
    # Nadawcą do systemu jest zawsze SENDER_NAME z nagłówka From:.
    if sender_name and sender_name.strip():
        vars_dict["USER_PERSON"] = sender_name.strip()
    else:
        # Fallback na detekcję z body tylko gdy webhook nie przysłał sender_name
        vars_dict["USER_PERSON"] = _detect_sender_name(body) or ""
    # ── Zdrobnienie imienia — słownik (bez zewnętrznego AI) ─────────────────────────────
    _user_person = (
        vars_dict["USER_PERSON"].split()[0] if vars_dict["USER_PERSON"] else ""
    )
    _ZDROBNIENIA = {
        "monika": "Moniczka",
        "anna": "Ania",
        "katarzyna": "Kasia",
        "małgorzata": "Gosia",
        "agnieszka": "Aga",
        "barbara": "Basia",
        "krystyna": "Krysia",
        "magdalena": "Madzia",
        "joanna": "Asia",
        "aleksandra": "Ola",
        "maria": "Marysia",
        "teresa": "Tereska",
        "irena": "Irka",
        "elżbieta": "Ela",
        "halina": "Halinka",
        "zofia": "Zosia",
        "danuta": "Danka",
        "beata": "Beatka",
        "ewa": "Ewka",
        "weronika": "Wera",
        "patrycja": "Patka",
        "marta": "Martusia",
        "karolina": "Karolcia",
        "natalia": "Natka",
        "sylwia": "Sylwka",
        "dorota": "Dorotka",
        "iwona": "Iwonka",
        "renata": "Renata",
        "tomasz": "Tomek",
        "piotr": "Piotrek",
        "krzysztof": "Krzysiek",
        "andrzej": "Andrzej",
        "jan": "Janek",
        "stanisław": "Stasiek",
        "michał": "Michał",
        "adam": "Adasiek",
        "marek": "Marek",
        "robert": "Robert",
        "paweł": "Pawełek",
        "marcin": "Marcinek",
        "jacek": "Jacek",
        "rafał": "Rafałek",
        "grzegorz": "Grzesiek",
        "dariusz": "Darek",
        "łukasz": "Łukasz",
        "artur": "Artur",
        "kamil": "Kamil",
        "mateusz": "Mateusz",
        "bartłomiej": "Bartek",
        "bartosz": "Bartek",
        "maciej": "Maciej",
        "wojciech": "Wojtek",
        "sebastian": "Sebastian",
        "dawid": "Dawid",
        "filip": "Filip",
        "szymon": "Szymon",
        "dominik": "Dominik",
        "patryk": "Patryk",
        "jakub": "Kuba",
        "daniel": "Daniel",
    }
    if _user_person:
        vars_dict["USER_NAME_ZDROBNIENIE"] = _ZDROBNIENIA.get(
            _user_person.lower(), _user_person
        )
    else:
        vars_dict["USER_NAME_ZDROBNIENIE"] = ""
    vars_dict["USER_GENDER"] = _detect_gender(body, sender_name)
    vars_dict["USER_CITY"] = _detect_city(body)
    vars_dict["USER_JOB"] = _detect_job(body)
    vars_dict["USER_EMOTION"] = emotion_key or ""
    vars_dict["USER_PROVIDER"] = provider or ""

    # ── Zdania Tylera → TEXT_1 .. TEXT_N ─────────────────────────────────────
    # Konwersja res_text na string je\u015bli jest dict (safety check)
    if isinstance(res_text, dict):
        res_text = json.dumps(res_text, ensure_ascii=False)
        logger.warning(
            "[session_vars] res_text by\u0142 dict — konwertowano na JSON string"
        )
    elif not isinstance(res_text, str):
        res_text = str(res_text) if res_text else ""

    tyler_text = ""
    if res_text and isinstance(res_text, str) and "### TYLER DURDEN" in res_text:
        tyler_text = res_text.split("### TYLER DURDEN", 1)[1]
    elif res_text:
        tyler_text = res_text

    tyler_sentences = _split_into_sentences(tyler_text)
    for i, s in enumerate(tyler_sentences, 1):
        vars_dict[f"TEXT_{i}"] = s

    # ── Zdania Sokratesa → SOKRATES_1 .. SOKRATES_N ──────────────────────────
    sokrates_text = ""
    if res_text and isinstance(res_text, str) and "### SOKRATES" in res_text:
        part = res_text.split("### SOKRATES", 1)[1]
        if "---" in part:
            sokrates_text = part.split("---")[0]
        else:
            sokrates_text = part

    sokrates_sentences = _split_into_sentences(sokrates_text)
    for i, s in enumerate(sokrates_sentences, 1):
        vars_dict[f"SOKRATES_{i}"] = s

    return vars_dict


def _render_template(text: str, vars_dict: dict) -> tuple:
    """
    Podstawia wszystkie [ZMIENNA] w tekście na wartości ze słownika.
    Zwraca (tekst_po_podstawieniu, lista_użytych_zmiennych).
    Jeśli [TEXT_N] nie istnieje w słowniku — losuje z dostępnych TEXT_*.
    """
    if not text or not vars_dict:
        return text, []

    used = []

    # Znajdź wszystkie placeholdery w tekście
    placeholders = re.findall(r"\[([A-Z_0-9]+)\]", text)

    # Zbierz dostępne TEXT_* i SOKRATES_* do losowania fallback
    text_keys = sorted(
        [k for k in vars_dict if re.match(r"^TEXT_\d+$", k)],
        key=lambda x: int(x.split("_")[1]),
    )
    sokrates_keys = sorted(
        [k for k in vars_dict if re.match(r"^SOKRATES_\d+$", k)],
        key=lambda x: int(x.split("_")[1]),
    )

    result = text
    for ph in placeholders:
        if ph in vars_dict:
            result = result.replace(f"[{ph}]", vars_dict[ph], 1)
            used.append(ph)
        elif re.match(r"^TEXT_\d+$", ph) and text_keys:
            # fallback — losuj z dostępnych TEXT_*
            fallback_key = random.choice(text_keys)
            result = result.replace(f"[{ph}]", vars_dict[fallback_key], 1)
            used.append(f"{ph}→{fallback_key}(losowy)")
        elif re.match(r"^SOKRATES_\d+$", ph) and sokrates_keys:
            fallback_key = random.choice(sokrates_keys)
            result = result.replace(f"[{ph}]", vars_dict[fallback_key], 1)
            used.append(f"{ph}→{fallback_key}(losowy)")
        # jeśli zmienna nieznana — zostawiamy [ZMIENNA] bez zmian

    return result, used


# _generate_panel_prompt usunięta — zastąpiona przez _generate_triptych_prompts_batch


# Zarządzanie tokenami HF jest teraz w core/hf_token_manager


def _png_to_jpg(image_obj: dict, panel_index: int) -> dict:
    """
    Konwertuje PNG (base64) do JPG 95% jakości.
    Nazwa wynikowa: tyler_YYYYMMDD_HHMMSS_panel{N}.jpg
    Zwraca nowy dict z zaktualizowanymi polami base64 / content_type / filename.
    Przy błędzie zwraca oryginał (PNG) żeby nie tracić obrazka.
    """
    try:
        from PIL import Image

        raw_bytes = base64.b64decode(image_obj["base64"])
        img = Image.open(io.BytesIO(raw_bytes)).convert("RGB")

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=TYLER_JPG_QUALITY, optimize=True)
        jpg_b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tyler_{ts}_panel{panel_index}.jpg"

        size_png_kb = len(raw_bytes) // 1024
        size_jpg_kb = len(buf.getvalue()) // 1024

        logger.info(
            "[tyler-jpg] Panel %d: %dKB PNG → %dKB JPG (jakość=%d%%)",
            panel_index,
            size_png_kb,
            size_jpg_kb,
            TYLER_JPG_QUALITY,
        )

        result = {
            "base64": jpg_b64,
            "content_type": "image/jpeg",
            "filename": filename,
            "size_jpg": f"{size_jpg_kb}KB",
            "size_png_orig": f"{size_png_kb}KB",
        }
        # Zachowaj metadata z oryginału
        for key in ("seed", "token_name", "remaining_requests"):
            if key in image_obj:
                result[key] = image_obj[key]
        return result

    except ImportError:
        logger.error("[tyler-jpg] Pillow niedostępny — zwracam PNG")
        return image_obj
    except Exception as e:
        logger.warning("[tyler-jpg] Błąd konwersji: %s — zwracam PNG", e)
        return image_obj


def _load_substitute_image() -> dict | None:
    if not os.path.exists(SUBSTITUTE_IMAGE_PATH):
        logger.warning("[test-mode] Brak pliku zastępczego: %s", SUBSTITUTE_IMAGE_PATH)
        return None
    try:
        with open(SUBSTITUTE_IMAGE_PATH, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("ascii")
        return {
            "base64": b64,
            "content_type": "image/jpeg",
            "filename": "zastepczy.jpg",
        }
    except Exception as e:
        logger.warning("[test-mode] Błąd odczytu zastepczy.jpg: %s", e)
        return None


def _generate_flux_image(
    prompt: str, panel_index: int = 0, test_mode: bool = False
) -> dict | None:
    """
    Generuje jeden obrazek FLUX z losowym seed.
    Próbuje każdy token HF po kolei.
    Zwraca dict z base64 lub None.

    Parametr test_mode:
    - Jeśli test_mode=True (przychodzi z KEYWORDS_TEST via disable_flux),
      to zwracamy zastępczy obrazek zamiast generować FLUX.
    - To oszczędza tokeny HF_TOKEN.
    """
    # ── KEYWORDS_TEST (disable_flux) → test_mode ──────────────────────────────
    # Jeśli test_mode=True, wy generowanie FLUX i użyj zastępczego obrazka
    if test_mode:
        image = _load_substitute_image()
        if image:
            image = dict(image)
            image["filename"] = f"tyler_panel{panel_index}_zastepczy.jpg"
        return image
    tokens = get_active_tokens()
    if tokens and panel_index > 0:
        offset = (panel_index - 1) % len(tokens)
        tokens = tokens[offset:] + tokens[:offset]
    if not tokens:
        if hf_tokens.all_dead():
            logger.warning(
                "[flux-tyler] Wszystkie tokeny HF na czarnej liście (402/401/403) — "
                "używam zastepczy.jpg zamiast FLUX"
            )
        else:
            logger.error(
                "[flux-tyler] Brak tokenów HF w zmiennych środowiskowych — używam zastepczy.jpg"
            )
        # Fallback do zastępczego obrazka — tak jak test_mode
        substitute = _load_substitute_image()
        if substitute:
            substitute = dict(substitute)
            substitute["filename"] = f"tyler_panel{panel_index}_zastepczy.jpg"
            return substitute
        return None

    seed = random.randint(0, 2**32 - 1)
    payload = {
        "inputs": prompt,
        "parameters": {
            "num_inference_steps": HF_STEPS,
            "guidance_scale": HF_GUIDANCE,
            "seed": seed,
        },
    }

    logger.info(
        "[flux-tyler] Panel %d — %d tokenów dostępnych, seed=%d",
        panel_index,
        len(tokens),
        seed,
    )

    for name, token in tokens:
        headers = {"Authorization": f"Bearer {token}", "Accept": "image/png"}
        try:
            logger.info("[flux-tyler] Próbuję token: %s", name)
            resp = requests.post(
                HF_API_URL, headers=headers, json=payload, timeout=HF_TIMEOUT
            )

            remaining = resp.headers.get("X-Remaining-Requests")

            if resp.status_code == 200:
                logger.info(
                    "[flux-tyler] ✓ Token %s: sukces (PNG %d B, pozostało: %s)",
                    name,
                    len(resp.content),
                    remaining or "?",
                )
                return {
                    "base64": base64.b64encode(resp.content).decode("ascii"),
                    "content_type": "image/png",
                    "filename": f"tyler_panel{panel_index}_seed{seed}.png",
                    "seed": seed,
                    "token_name": name,
                    "remaining_requests": int(remaining) if remaining else None,
                }

            elif resp.status_code == 402:
                # Wyczerpane kredyty — dodaj do czarnej listy na całą sesję
                mark_dead(name)
                logger.warning(
                    "[flux-tyler] ✗ Token %s: wyczerpane kredyty (402) — "
                    "dodano do czarnej listy sesji",
                    name,
                )
            elif resp.status_code in (401, 403):
                # Nieważny token — też na czarną listę
                mark_dead(name)
                logger.warning(
                    "[flux-tyler] ✗ Token %s: nieważny (HTTP %d) — "
                    "dodano do czarnej listy sesji",
                    name,
                    resp.status_code,
                )
            elif resp.status_code in (503, 529):
                logger.warning(
                    "[flux-tyler] ⚠ Token %s: przeciążony (HTTP %d) — ponowna próba później",
                    name,
                    resp.status_code,
                )
            else:
                logger.warning(
                    "[flux-tyler] ✗ Token %s: HTTP %d: %s",
                    name,
                    resp.status_code,
                    resp.text[:100],
                )

        except requests.exceptions.Timeout:
            logger.warning("[flux-tyler] ⏱ Token %s: timeout (%ds)", name, HF_TIMEOUT)
        except requests.exceptions.ConnectionError as e:
            logger.warning(
                "[flux-tyler] 🔌 Token %s: connection error: %s", name, str(e)[:80]
            )
        except Exception as e:
            logger.warning("[flux-tyler] ❌ Token %s: wyjątek: %s", name, str(e)[:80])

    logger.error(
        "[flux-tyler] Wszystkie tokeny HF zawiodły dla panelu %d — używam zastepczy.jpg",
        panel_index,
    )
    substitute = _load_substitute_image()
    if substitute:
        substitute = dict(substitute)
        substitute["filename"] = f"tyler_panel{panel_index}_zastepczy.jpg"
        return substitute
    return None


# ═══════════════════════════════════════════════════════════════════════════════
# GENEROWANIE TRYPTYKU
# ═══════════════════════════════════════════════════════════════════════════════


def _generate_raw_email_image(body: str, test_mode: bool = False) -> dict | None:
    """
    Generuje obrazek FLUX bezpośrednio z treści emaila — BEZ udziału AI.
    Prompt = surowa treść emaila skrócona do 400 znaków.
    Obrazek jest konwertowany do JPG 95% i zmniejszony do 95% rozmiaru.
    """
    # Surowy prompt — tylko treść emaila, żadnego AI
    raw_prompt = body.strip()[:400]

    logger.info(
        "[raw-img] Generuję obrazek z surowej treści emaila (%.80s...)", raw_prompt
    )

    img = _generate_flux_image(raw_prompt, panel_index=97, test_mode=test_mode)
    if not img or not img.get("base64"):
        logger.warning("[raw-img] Brak obrazka z surowej treści")
        return None

    try:
        from PIL import Image as PILImage

        raw_bytes = base64.b64decode(img["base64"])
        pil = PILImage.open(io.BytesIO(raw_bytes)).convert("RGB")

        # Zmniejsz do 95% rozmiaru
        w, h = pil.size
        new_w = int(w * 0.95)
        new_h = int(h * 0.95)
        pil = pil.resize((new_w, new_h), PILImage.LANCZOS)

        buf = io.BytesIO()
        pil.save(buf, format="JPEG", quality=95, optimize=True)
        jpg_b64 = base64.b64encode(buf.getvalue()).decode("ascii")

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tyler_raw_email_{ts}.jpg"

        logger.info("[raw-img] OK: %s (%dKB)", filename, len(buf.getvalue()) // 1024)

        return {
            "base64": jpg_b64,
            "content_type": "image/jpeg",
            "filename": filename,
            "size_jpg": f"{len(buf.getvalue()) // 1024}KB",
        }

    except Exception as e:
        logger.warning("[raw-img] Błąd konwersji: %s", e)
        return img


def _generate_triptych_prompts_batch(
    panel_rules: list,
    session_vars: dict,
    style_config: dict,
) -> list:
    """
    Generuje prompty FLUX dla wszystkich 7 paneli w JEDNYM wywołaniu DeepSeek.
    Zamiast 7 osobnych calli → 1 call zwracający JSON z 7 promptami.
    Zwraca listę 7 stringów (promptów), fallback na puste stringi.
    """
    w = _load_panel_wytyczne()
    nouns_str = session_vars.get("USER_OBJECTS", "") or "debris, broken furniture, ash"
    panel_style = random.choice(
        w.get("style_variants", ["35mm film grain, Fight Club 1999"])
    )

    # Buduj listę zasad do promptu zbiorczego
    zasady_lines = []
    for i, rule in enumerate(panel_rules[:7], 1):
        zasady_lines.append(f"Panel {i}: {rule[:120] if rule else '(brak zasady)'}")
    zasady_str = "\n".join(zasady_lines)

    system_batch = (
        "You are a cinematic visual prompt engineer for FLUX image generation. "
        "Fight Club 1999 aesthetic, David Fincher style. "
        "Given 7 Tyler Durden RULES, generate 7 FLUX image prompts, one per rule. "
        "Each prompt: ONE paragraph, max 80 words, English only. "
        "Show the VIOLATION of each rule — characters actively doing what the rule forbids. "
        "Characters look damaged, unwashed, nihilistic. "
        f"Visual style: {panel_style}, 35mm film grain, gritty, underexposed. "
        'RESPOND ONLY with valid JSON: {"prompts": ["prompt1", "prompt2", ..., "prompt7"]} '
        "No other text, no markdown fences."
    )
    user_batch = (
        f"Objects from sender's email context: {nouns_str}\n\n"
        f"7 Tyler Durden Rules:\n{zasady_str}\n\n"
        "Generate exactly 7 FLUX prompts as JSON array under key 'prompts'."
    )

    raw, prov = _call_ai_with_fallback(system_batch, user_batch, max_tokens=2000)
    logger.info("[tryptyk-batch] Call %s → %d znaków odpowiedzi", prov, len(raw or ""))

    if not raw:
        return [""] * 7

    try:
        clean = _strip_json_markdown(raw)
        # Użyj raw_decode zamiast json.loads — obsługuje "Extra data"
        decoder = json.JSONDecoder()
        parsed = None
        try:
            parsed, _ = decoder.raw_decode(clean)
        except json.JSONDecodeError:
            for match in re.finditer(r"[{\[]", clean):
                start = match.start()
                try:
                    obj, end = decoder.raw_decode(clean[start:])
                    if obj is not None:
                        parsed = obj
                        break
                except json.JSONDecodeError:
                    continue
        if parsed is None:
            raise ValueError("Nie znaleziono JSON")
        # Obsłuż przypadek gdy model zwrócił tablicę bezpośrednio: ["p1","p2",...]
        # zamiast {"prompts": ["p1","p2",...]}
        if isinstance(parsed, list):
            prompts = parsed
            logger.info(
                "[tryptyk-batch] Model zwrócił tablicę bezpośrednio — akceptuję"
            )
        elif isinstance(parsed, dict):
            prompts = parsed.get("prompts", [])
        else:
            raise ValueError(f"Nieoczekiwany typ: {type(parsed).__name__}")
        if isinstance(prompts, list) and len(prompts) >= 1:
            # Uzupełnij do 7 jeśli model zwrócił mniej
            while len(prompts) < 7:
                prompts.append("")
            logger.info("[tryptyk-batch] OK — %d promptów", len(prompts))
            return [str(p)[:500] for p in prompts[:7]]
    except Exception as e:
        logger.warning("[tryptyk-batch] Błąd JSON: %s | raw: %.200s", e, raw)

    # Fallback: podziel raw po newlinach jeśli JSON nie wyszedł
    lines = [l.strip() for l in (raw or "").split("\n") if len(l.strip()) > 20]
    while len(lines) < 7:
        lines.append("")
    return lines[:7]


def _generate_triptych(
    response_text: str,
    prompt_data: dict,
    body: str,
    session_vars: dict = None,
    test_mode: bool = False,
) -> tuple:
    """
    Generuje 7 paneli — każdy odpowiada jednej zasadzie Tylera.
    OPTYMALIZACJA: 1 call DeepSeek dla wszystkich 7 promptów naraz (zamiast 7 calli).
    Obrazki FLUX generowane sekwencyjnie.
    Jeśli HF_TOKEN nie działa → panel pomijany, zwracamy ile wygenerowano.
    """
    if session_vars is None:
        session_vars = {}

    style_config = _load_style_config() or {}
    panel_rules = _extract_tyler_rules(response_text)

    # Fallback: brak zasad → 1 panel z wytycznych JSON
    if not any(panel_rules):
        logger.warning(
            "[zwykly-img] Brak zasad Tylera — fallback: 1 panel z wytycznych JSON"
        )
        w = _load_panel_wytyczne()
        fallback_prompt = w.get("fallback_gdy_brak_zasady", "").replace(
            "[USER_OBJECTS]", session_vars.get("USER_OBJECTS", "debris")
        )
        if not fallback_prompt:
            return [], [], []
        image = _generate_flux_image(
            fallback_prompt, panel_index=1, test_mode=test_mode
        )
        if not image:
            return [], [], []
        image = _png_to_jpg(image, panel_index=1)
        image = _add_text_below_image(image, "Tyler Durden", 1)
        return (
            [image],
            [fallback_prompt],
            [
                {
                    "panel": 1,
                    "caption": "fallback",
                    "used_vars": [],
                    "prompt_preview": fallback_prompt[:120],
                }
            ],
        )

    while len(panel_rules) < 7:
        panel_rules.append("")

    if test_mode:
        substitute = _load_substitute_image()
        if substitute:
            images = []
            panel_prompts = []
            panel_assignments = []
            for idx in range(1, 8):
                img = dict(substitute)
                img["filename"] = f"tyler_panel{idx}_zastepczy.jpg"
                # Dodaj napis z zasady na obrazek w test_mode
                rule_text = (
                    (panel_rules[idx - 1] or "")[:120]
                    if panel_rules[idx - 1]
                    else f"Zasada {idx}"
                )
                img = _add_text_below_image(img, rule_text, idx)
                images.append(img)
                panel_prompts.append("")
                panel_assignments.append(
                    {
                        "panel": idx,
                        "rule": (panel_rules[idx - 1] or "")[:100],
                        "caption": rule_text,
                        "used_vars": [],
                        "prompt_preview": "[test_mode substitute image z napisem zasady]",
                    }
                )
            logger.info(
                "[zwykly-img] test_mode — używam zastępczego obrazu dla 7 paneli (+napisy)"
            )
            return images, panel_prompts, panel_assignments
        logger.info("[zwykly-img] test_mode — brak zastepczy.jpg, pomijam FLUX")
        return [], [], []

    # ── 1 CALL: Generuj wszystkie 7 promptów naraz ───────────────────────────
    logger.info("[zwykly-img] Generuję 7 promptów FLUX w 1 callu DeepSeek")
    flux_prompts = _generate_triptych_prompts_batch(
        panel_rules, session_vars, style_config
    )

    # ── Generuj obrazki równolegle (bez dodatkowych calli DeepSeek) ──────────
    def _gen_panel(panel_idx):
        rule_text = panel_rules[panel_idx - 1]
        flux_prompt = flux_prompts[panel_idx - 1]

        if not flux_prompt and rule_text:
            # Prosty fallback bez AI — składamy ręcznie
            nouns_str = session_vars.get("USER_OBJECTS", "debris, ash")
            flux_prompt = (
                f"Fight Club 1999, Tyler Durden violating the rule: '{rule_text[:80]}', "
                f"surrounded by {nouns_str}, 35mm film grain, gritty, underexposed, Fincher"
            )

        caption = rule_text[:120] if rule_text else f"Zasada {panel_idx}"
        image = (
            _generate_flux_image(
                flux_prompt, panel_index=panel_idx, test_mode=test_mode
            )
            if flux_prompt
            else None
        )
        if image:
            image = _png_to_jpg(image, panel_index=panel_idx)
            image = _add_text_below_image(image, caption, panel_idx)
        return panel_idx, image, flux_prompt, [], caption

    results = {}
    logger.info("[zwykly-img] Generuję 7 paneli FLUX sekwencyjnie")

    for i in range(1, 8):
        try:
            idx, img, prompt, uvars, caption = _gen_panel(i)
            results[idx] = (img, prompt, uvars, caption)
            if img:
                logger.info("[zwykly-img] Panel %d/7 OK", idx)
            else:
                logger.warning(
                    "[zwykly-img] Panel %d/7 brak obrazka (HF limit lub brak zasady)",
                    idx,
                )
        except Exception as e:
            logger.error("[zwykly-img] Panel %d/7 błąd: %s", i, e)
            results[i] = (None, "", [], f"Zasada {i}")

    # Złóż w kolejności 1-7
    images, panel_prompts, panel_assignments = [], [], []
    for idx in range(1, 8):
        img, prompt, uvars, caption = results.get(idx, (None, "", [], f"Zasada {idx}"))
        if prompt:
            panel_prompts.append(prompt)
        panel_assignments.append(
            {
                "panel": idx,
                "rule": (panel_rules[idx - 1] or "")[:100],
                "caption": caption,
                "used_vars": uvars or [],
                "prompt_preview": (prompt or "")[:120],
            }
        )
        if img:
            images.append(img)

    logger.info("[zwykly-img] Wygenerowano %d/7 paneli", len(images))
    return images, panel_prompts, panel_assignments


# ═══════════════════════════════════════════════════════════════════════════════
# GŁÓWNA FUNKCJA
# ═══════════════════════════════════════════════════════════════════════════════


def _build_debug_txt(
    body: str,
    provider: str,
    emotion_key: str,
    res_raw: str,
    res_text: str,
    triptych_images: list,
    panel_prompts: list,
    system_msg: str = "",
    user_msg: str = "",
    session_vars: dict = None,
    panel_assignments: list = None,
) -> dict:
    """
    Buduje pełny log debug TXT do zapisu na Google Drive.
    Zawiera: statystyki długości, wszystkie prompty, odpowiedź AI,
    info o obrazkach, WSZYSTKIE zmienne sesji, przyporządkowania paneli,
    zestawienie końcowe co nadawca otrzyma.
    """
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    if session_vars is None:
        session_vars = {}
    if panel_assignments is None:
        panel_assignments = []

    body_len = len(body or "")
    user_msg_len = len(user_msg or "")
    res_raw_len = len(res_raw or "")
    res_text_len = len(res_text or "")
    system_msg_len = len(system_msg or "")

    # Zestawienie obrazków
    img_lines = []
    for i, img in enumerate(triptych_images or [], 1):
        fn = img.get("filename", "?")
        ct = img.get("content_type", "?")
        size = img.get("size_jpg", img.get("size_png_orig", "?"))
        img_lines.append(f"  Obrazek {i}: {fn} | format: {ct} | rozmiar: {size}")

    # Zestawienie "nadawca otrzyma"
    otrzyma = []
    if res_text:
        otrzyma.append("  v reply_html — odpowiedz Tylera i Sokratesa (HTML)")
    if triptych_images:
        otrzyma.append(
            f"  v triptych — {len(triptych_images)} obrazek(ow) JPG Fight Club"
        )
    otrzyma.append("  v emoticon — emotka PNG (FLUX)")
    otrzyma.append("  v cv_pdf — CV w stylu Tylera (PDF)")
    otrzyma.append("  v horoskop_pdf — Horoskop nihilistyczny (PDF)")
    otrzyma.append("  v karta_rpg_pdf — Karta postaci RPG (PDF)")
    otrzyma.append("  v raport_pdf — Raport psychiatryczny (PDF)")
    otrzyma.append("  v ankieta_pdf — Ankieta interaktywna AcroForm (PDF)")
    otrzyma.append("  v plakat_svg — Plakat Tyler Durden (SVG)")
    otrzyma.append("  v gra_html — Gra interaktywna (HTML)")
    otrzyma.append("  v wyjasnienie.txt — Wyjasnienie odpowiedzi (TXT)")
    otrzyma.append("  v _.txt — Ten log debugowania (TXT)")

    lines = [
        f"=== ZWYKLY DEBUG {ts} ===",
        f"provider:              {provider}",
        f"emocja:                {emotion_key}",
        f"panele wygenerowane:   {len(triptych_images)}",
        "",
        "---------------------------------------------",
        "STATYSTYKI DLUGOSCI",
        "---------------------------------------------",
        f"Email otrzymany:        {body_len} znakow",
        f"System prompt:          {system_msg_len} znakow",
        f"User prompt (do AI):    {user_msg_len} znakow (email + instrukcje)",
        f"Odpowiedz surowa (AI):  {res_raw_len} znakow",
        f"Odpowiedz tekstowa:     {res_text_len} znakow",
        "",
        "---------------------------------------------",
        "TRESC EMAILA (pelna)",
        "---------------------------------------------",
        (body or "(brak)"),
        "",
        "---------------------------------------------",
        "SYSTEM PROMPT (pelny)",
        "---------------------------------------------",
        (system_msg or "(brak)"),
        "",
        "---------------------------------------------",
        "USER PROMPT WYSLANY DO AI (pelny)",
        "---------------------------------------------",
        (user_msg or "(brak)"),
        "",
        "---------------------------------------------",
        "SUROWA ODPOWIEDZ AI (pelna)",
        "---------------------------------------------",
        (res_raw or "(brak)"),
        "",
        "---------------------------------------------",
        "ODPOWIEDZ TEKSTOWA (pelna)",
        "---------------------------------------------",
        (res_text or "(brak)"),
        "",
        "---------------------------------------------",
        "PROMPTY PANELI FLUX (pelne)",
        "---------------------------------------------",
    ]
    for i, p in enumerate(panel_prompts or [], 1):
        lines.append(f"Panel {i}:")
        lines.append(p)
        lines.append("")

    lines += [
        "---------------------------------------------",
        "OBRAZKI WYGENEROWANE",
        "---------------------------------------------",
    ]
    if img_lines:
        lines += img_lines
    else:
        lines.append("  (brak obrazkow)")

    # ── ZMIENNE GLOBALNE SESJI ────────────────────────────────────────────────
    lines += [
        "",
        "---------------------------------------------",
        "ZMIENNE GLOBALNE SESJI (dostepne jako [ZMIENNA] w plikach JSON)",
        "---------------------------------------------",
    ]

    # Grupuj: najpierw webhook/wykryte, potem TEXT_*, potem SOKRATES_*
    webhook_keys = ["SENDER", "SENDER_NAME", "BODY", "PREVIOUS_BODY"]
    detected_keys = [
        "USER_PERSON",
        "USER_NAME_ZDROBNIENIE",
        "USER_OBJECTS",
        "USER_GENDER",
        "USER_CITY",
        "USER_JOB",
        "USER_EMOTION",
        "USER_PROVIDER",
    ]
    text_keys = sorted(
        [k for k in session_vars if re.match(r"^TEXT_\d+$", k)],
        key=lambda x: int(x.split("_")[1]),
    )
    sokr_keys = sorted(
        [k for k in session_vars if re.match(r"^SOKRATES_\d+$", k)],
        key=lambda x: int(x.split("_")[1]),
    )

    lines.append("-- Z Google Apps Script / webhook:")
    for k in webhook_keys:
        v = session_vars.get(k, "")
        preview = str(v)[:120].replace("\n", " ")
        lines.append(f'  [{k}] = "{preview}"')

    lines.append("")
    lines.append("-- Wykryte z emaila:")
    for k in detected_keys:
        v = session_vars.get(k, "")
        lines.append(f'  [{k}] = "{v}"')

    lines.append("")
    lines.append(f"-- Zdania Tylera ({len(text_keys)} zdań):")
    for k in text_keys:
        v = session_vars.get(k, "")
        lines.append(f'  [{k}] = "{v}"')

    lines.append("")
    lines.append(f"-- Zdania Sokratesa ({len(sokr_keys)} zdań):")
    for k in sokr_keys:
        v = session_vars.get(k, "")
        lines.append(f'  [{k}] = "{v}"')

    # ── PRZYPORZĄDKOWANIE PANELI ──────────────────────────────────────────────
    lines += [
        "",
        "---------------------------------------------",
        "PRZYPORZĄDKOWANIE ZMIENNYCH DO PANELI",
        "---------------------------------------------",
    ]
    if panel_assignments:
        for pa in panel_assignments:
            used = ", ".join(pa.get("used_vars", [])) or "(brak podstawien)"
            lines.append(f"  Panel {pa['panel']}: uzyte zmienne: {used}")
            lines.append(f"    caption: \"{pa.get('caption', '')}\"")
            lines.append(f"    prompt (poczatek): \"{pa.get('prompt_preview', '')}\"")
            lines.append("")
    else:
        lines.append("  (brak danych o przyporządkowaniu)")

    lines += [
        "",
        "---------------------------------------------",
        "NADAWCA POWINIEN OTRZYMAC:",
        "---------------------------------------------",
    ]
    lines += otrzyma
    lines += [
        "",
        "=== KONIEC ===",
    ]

    content = "\n".join(lines)
    b64 = base64.b64encode(content.encode("utf-8")).decode("ascii")
    return {
        "base64": b64,
        "content_type": "text/plain",
        "filename": "_.txt",
        "filename_drive": f"zwykly_debug_{ts}.txt",
    }


def _generate_icon_flux(emotion_key: str, sender_name: str = "") -> str | None:
    """
    Zwraca emotkę PNG z katalogu EMOTKI_DIR — bez wywołania API/FLUX.
    HF tokeny są na czarnej liście — generowanie FLUX nie ma sensu.
    Jeśli plik istnieje → zwraca base64, jeśli nie → None.
    """
    emot_name = EMOCJA_MAP.get(emotion_key, FALLBACK_EMOT)
    path = os.path.join(EMOTKI_DIR, f"{emot_name}.png")
    if os.path.exists(path):
        try:
            with open(path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")
            logger.info("[icon] Emotka z pliku: %s", path)
            return b64
        except Exception as e:
            logger.warning("[icon] Błąd odczytu emotki %s: %s", path, e)
    # Spróbuj fallback na nazwę emotion_key bezpośrednio
    path2 = os.path.join(EMOTKI_DIR, f"{emotion_key}.png")
    if os.path.exists(path2):
        try:
            with open(path2, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("ascii")
            logger.info("[icon] Emotka z pliku (fallback): %s", path2)
            return b64
        except Exception as e:
            logger.warning("[icon] Błąd odczytu emotki fallback %s: %s", path2, e)
    logger.warning("[icon] Brak pliku emotki dla emocji: %s", emotion_key)
    return None


def _generate_cv_content(
    body: str, previous_body: str | None, sender_email: str, sender_name: str = ""
) -> dict | None:
    """
    Generuje treść CV w stylu Tylera przez DeepSeek AI.
    Zwraca dict z polami CV lub None przy błędzie.
    """
    try:
        with open(CV_CONTENT_JSON_PATH, encoding="utf-8") as f:
            cv_cfg = json.load(f)
    except Exception as e:
        logger.warning("[cv] Brak zwykly_cv_content.json: %s", e)
        cv_cfg = {}

    system_msg = cv_cfg.get(
        "system", "Generuj prześmiewcze CV w stylu Tylera Durdena. Zwróć TYLKO JSON."
    )
    schema = cv_cfg.get("output_schema", {})
    instrukcje = cv_cfg.get("instrukcje_dodatkowe", [])

    context_parts = [f"EMAIL:\n{body[:MAX_DLUGOSC_EMAIL]}"]
    if sender_name and sender_name.strip():
        context_parts.append(f"\nIMIĘ NADAWCY (SENDER_NAME): {sender_name}")
        context_parts.append(
            f"KRYTYCZNE: pole 'imie_nazwisko' w CV MUSI zaczynać się od '{sender_name}' "
            f"lub zawierać imię '{sender_name}' + wymyślone nazwisko nawiązujące do emaila. "
            f"ZAKAZ 'Anonim Bezdomny' lub jakiegokolwiek imienia bez związku z nadawcą."
        )
    if previous_body and previous_body.strip():
        context_parts.append(
            f"\nPOPRZEDNIA WIADOMOŚĆ:\n{previous_body[:MAX_DLUGOSC_EMAIL]}"
        )
    if sender_email:
        context_parts.append(f"\nEMAIL NADAWCY: {sender_email}")
    context_parts.append(
        f"\nSCHEMAT JSON DO WYPEŁNIENIA:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
    )
    if instrukcje:
        context_parts.append(
            f"\nINSTRUKCJE:\n" + "\n".join(f"- {i}" for i in instrukcje)
        )
    context_parts.append(
        "\nKRYTYCZNE: Zwróć TYLKO czysty JSON bez żadnego tekstu poza klamrami. "
        "WSZYSTKIE pola schematu MUSZĄ być wypełnione — doswiadczenie, wyksztalcenie, "
        "umiejetnosci, jezyki, zainteresowania, zyciorys, cytat_tylera. "
        "ZAKAZ zwracania pustych list [] lub pustych stringów dla tych pól.\n\n"
        "PRZYKŁAD PRAWIDŁOWEGO OUTPUT (każde pole MUSI być wypełnione):\n"
        "{\n"
        '  "imie_nazwisko": "Imię z konkretnym nazwiskiem",\n'
        '  "tytul_zawodowy": "Konkretny tytuł",\n'
        '  "doswiadczenie": [\n'
        '    {"firma": "Konkretna firma", "stanowisko": "Stanowisko", "okres": "Lata", "obowiazki": ["Obowiązek 1", "Obowiązek 2", "Obowiązek 3"]},\n'
        '    {"firma": "Druga firma", "stanowisko": "Stanowisko", "okres": "Lata", "obowiazki": ["Obowiązek 1", "Obowiązek 2"]}\n'
        "  ],\n"
        '  "umiejetnosci": ["Umiejętność 1", "Umiejętność 2", "Umiejętność 3", "Umiejętność 4", "Umiejętność 5"],\n'
        '  "jezyki": ["Język 1", "Język 2", "Język 3", "Język 4", "Język 5"],\n'
        '  "zainteresowania": ["Zainteresowanie 1", "Zainteresowanie 2"],\n'
        '  "zyciorys": "Kilka zdań",\n'
        '  "cytat_tylera": "Jedno zdanie podsumowania"\n'
        "}\n\n"
        "ZAKAZ ZWRACANIA CZEGO KOKOLWIEK POZA JSON. Każda lista MUSI mieć min 3 elementy. Każdy string MUSI mieć min 5 słów."
    )

    user_msg = "\n".join(context_parts)

    def _normalize_cv_data(cv_data):
        if isinstance(cv_data, list) and cv_data:
            first = cv_data[0]
            if isinstance(first, dict):
                logger.warning(
                    "[cv] _parse_cv: root list zamiast dict — używam pierwszego obiektu"
                )
                cv_data = first
        if not isinstance(cv_data, dict):
            return cv_data

        def _try_parse_json(value):
            if isinstance(value, str):
                try:
                    return json.loads(value)
                except Exception:
                    pass
            return value

        for key in [
            "doswiadczenie",
            "wyksztalcenie",
            "umiejetnosci",
            "jezyki",
            "zainteresowania",
        ]:
            if key in cv_data and isinstance(cv_data[key], str):
                parsed = _try_parse_json(cv_data[key])
                if isinstance(parsed, list):
                    logger.warning(
                        "[cv] _parse_cv: sparsowano stringowe %s jako listę",
                        key,
                    )
                    cv_data[key] = parsed

        return cv_data

    def _parse_cv(raw: str) -> dict | None:
        if not raw:
            return None
        try:
            clean = _strip_json_markdown(raw)
            decoder = json.JSONDecoder()
            cv_data = None
            try:
                cv_data, _ = decoder.raw_decode(clean)
            except json.JSONDecodeError:
                for match in re.finditer(r"[{\[]", clean):
                    start = match.start()
                    try:
                        obj, end = decoder.raw_decode(clean[start:])
                        if obj is not None:
                            cv_data = obj
                            break
                    except json.JSONDecodeError:
                        continue
            if cv_data is None:
                return None
            cv_data = _normalize_cv_data(cv_data)
            if cv_data is None or not isinstance(cv_data, dict):
                return None
            return cv_data
        except Exception:
            return None

    def _cv_is_complete(cv: dict) -> bool:
        """Sprawdza czy kluczowe tablicowe sekcje CV są wypełnione (z minimalną zawartością)."""
        if not cv:
            return False

        # Minimalne progi list
        list_minimums = {
            "doswiadczenie": 2,
            "umiejetnosci": 5,
            "jezyki": 5,
            "zainteresowania": 1,
        }
        for key, min_len in list_minimums.items():
            val = cv.get(key)
            if not val or not isinstance(val, list) or len(val) < min_len:
                logger.warning(
                    "[cv] _cv_is_complete: '%s' niekompletne (len=%d, wymagane>=%d)",
                    key,
                    len(val) if isinstance(val, list) else 0,
                    min_len,
                )
                return False

        # Sprawdzenie czy stringi nie są puste
        required_strings = ["imie_nazwisko", "zyciorys", "cytat_tylera"]
        for key in required_strings:
            val = cv.get(key, "").strip()
            if not val or len(val) < 5:  # Min 5 znaków
                logger.warning("[cv] _cv_is_complete: '%s' pusty lub za krótki", key)
                return False

        if cv.get("imie_nazwisko") in ("Anonim Bezdomny", ""):
            return False
        return True

    # Próba 1 — max_tokens=8000
    raw, _ = _call_ai_with_fallback(system_msg, user_msg, max_tokens=8000)
    if raw:
        cv_data = _parse_cv(raw)
        if cv_data and _cv_is_complete(cv_data):
            logger.info(
                "[cv] CV wygenerowane OK: %s", cv_data.get("imie_nazwisko", "?")
            )
            return cv_data
        elif cv_data:
            logger.warning(
                "[cv] CV niekompletne z próby 1 — będzie ponownie generowane w próbie 2. Pola: %s",
                {
                    k: len(v) if isinstance(v, list) else len(str(v))
                    for k in [
                        "doswiadczenie",
                        "umiejetnosci",
                        "jezyki",
                        "imie_nazwisko",
                    ]
                    if k in cv_data
                },
            )
            cv_data = None  # Nie zwracaj niekompletne — spróbuj jeszcze raz

    # Próba 2 — uproszczony prompt, więcej tokenów
    context_parts_retry = [
        f"EMAIL:\n{body[:MAX_DLUGOSC_EMAIL]}",
    ]
    if sender_name:
        context_parts_retry.append(f"\nIMIĘ NADAWCY: {sender_name}")
    context_parts_retry.append(
        f"\nWYGENERUJ CV W STYLU TYLERA DURDENA. Zwróć TYLKO czysty JSON z tymi dokładnie polami:\n"
        f"{json.dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"OBOWIĄZKOWE: doswiadczenie (lista 3 firm), umiejetnosci (lista 5), jezyki (lista min 5), "
        f"wyksztalcenie (lista 1), zainteresowania (lista 3), zyciorys (string), cytat_tylera (string). "
        f"ABSOLUTNY ZAKAZ pustych list. Zwróć TYLKO JSON, zaczynając od {{."
    )
    user_msg_retry = "\n".join(context_parts_retry)
    raw2, _ = _call_ai_with_fallback(system_msg, user_msg_retry, max_tokens=10000)
    if raw2:
        cv_data2 = _parse_cv(raw2)
        if cv_data2 and _cv_is_complete(cv_data2):
            logger.info(
                "[cv] CV (retry) wygenerowane OK: %s",
                cv_data2.get("imie_nazwisko", "?"),
            )
            return cv_data2
        elif cv_data2:
            logger.warning(
                "[cv] CV zwrócone z próby 2 ale niekompletne — będzie użyty fallback"
            )
            # Nie zwracaj niekompletne — przejdź do fallback

    # Obie próby zawiodły lub CV niekompletne — użyj fallback generator

    # Fallback: Jeśli obydwie próby zawiodły, wygeneruj minimalne CV
    logger.error("[cv] Fallback: Wygenerowanie minimalnego CV ze zmyślonymi danymi")
    return _generate_fallback_cv(body, sender_name, sender_email)


def _generate_fallback_cv(
    body: str, sender_name: str = "", sender_email: str = ""
) -> dict | None:
    """
    Fallback generator CV — tworzy minimalne ale kompletne CV
    na podstawie ekstrahowanych danych z emaila.
    """
    try:
        # Ekstrahuj podstawowe dane z emaila
        nouns = _extract_nouns_from_body(body)
        main_noun = nouns[0] if nouns else "Pracownik"

        # Wygeneruj imię i nazwisko
        if sender_name and sender_name.strip():
            imie_nazwisko = sender_name
        else:
            imie_nazwisko = f"{main_noun} {main_noun}ski"

        # Podstawowe dane
        job_title = "Specjalista ds. " + main_noun
        email_fallback = (
            sender_email or f"{imie_nazwisko.replace(' ', '.').lower()}@nieznany.pl"
        )
        city_detected = _detect_city(body) or "Polska"

        # Wygeneruj minimalne sekcje
        return {
            "imie_nazwisko": imie_nazwisko[:60],
            "tytul_zawodowy": job_title[:80],
            "email": email_fallback[:100],
            "telefon": f"666-{main_noun.upper()[:10]}-00{len(nouns)}",
            "miasto": city_detected[:50],
            "podsumowanie": (
                f"{imie_nazwisko} to specjalista od {main_noun.lower()}. "
                f"Wiele lat doświadczenia. Zawsze dostępny do pracy nad nowymi projektami."
            ),
            "doswiadczenie": [
                {
                    "firma": f"{main_noun} Solutions Sp. z o.o.",
                    "stanowisko": f"Senior {main_noun} Engineer",
                    "okres": "2020-2026",
                    "obowiazki": [
                        f"Wdrażanie rozwiązań związanych z {main_noun.lower()}",
                        f"Zarządzanie projektami z {main_noun.lower()}",
                        "Konsultacje i wsparcie techniczne dla klientów",
                    ],
                },
                {
                    "firma": f"{main_noun} & Co.",
                    "stanowisko": f"{main_noun} Consultant",
                    "okres": "2018-2020",
                    "obowiazki": [
                        f"Analiza i optymalizacja procesów {main_noun.lower()}",
                        "Szkolenia zespołu",
                        "Raportowanie wyników do zarządu",
                    ],
                },
            ],
            "wyksztalcenie": [
                {
                    "uczelnia": f"Akademia {main_noun} w Polsce",
                    "kierunek": f"Inżynieria {main_noun} i Systemów",
                    "rok": 2018,
                }
            ],
            "umiejetnosci": [
                f"{main_noun} - zaawansowany",
                "Zarządzanie projektami",
                "Komunikacja międzykulturowa",
                "Problem solving",
                "Analiza danych",
            ],
            "jezyki": [
                "Polski - ojczysty",
                "Angielski - biegły",
                f"Język {main_noun} - biegły",
                "Niemiecki - średnio zaawansowany",
                "Komunikacja niewerbal - zaawansowana",
            ],
            "zainteresowania": [
                f"Nowinki technologiczne w {main_noun}",
                "Rozwijanie umiejętności zawodowych",
                "Podróże biznesowe",
            ],
            "zyciorys": (
                f"{imie_nazwisko} to profesjonalista z bogatym doświadczeniem w dziedzinie {main_noun.lower()}. "
                f"Od wielu lat pracuje nad rozwojem innowacyjnych rozwiązań dla klientów z całej Polski. "
                f"Specjalizuje się w zarządzaniu złożonymi projektami i budowaniu efektywnych zespołów. "
                f"W wolnym czasie rozwija swoje umiejętności poprzez udział w konferencjach branżowych i szkoleniach."
            ),
            "cytat_tylera": (
                f"Jesteś egzempem zawodowca. Wiesz co robić, robisz to dobrze, ale nie wiesz dlaczego."
            ),
        }
    except Exception as e:
        logger.error("[cv-fallback] Błąd w fallback generatorze: %s", e)
        # Ostateczny fallback - puste ale kompletne CV
        return {
            "imie_nazwisko": sender_name or "Pracownik",
            "tytul_zawodowy": "Specjalista",
            "email": sender_email or "pracownik@nieznany.pl",
            "telefon": "666-000-000",
            "miasto": "Polska",
            "podsumowanie": "Profesjonalista o bogatym doświadczeniu.",
            "doswiadczenie": [
                {
                    "firma": "Firma A",
                    "stanowisko": "Stanowisko A",
                    "okres": "2020-2026",
                    "obowiazki": ["Obowiązek 1", "Obowiązek 2"],
                }
            ],
            "wyksztalcenie": [
                {"uczelnia": "Akademia", "kierunek": "Kierunek", "rok": 2018}
            ],
            "umiejetnosci": ["Umiejętność 1", "Umiejętność 2"],
            "jezyki": ["Polski", "Angielski"],
            "zainteresowania": ["Praca", "Rozwój"],
            "zyciorys": "Osoba o znacznym doświadczeniu zawodowym.",
            "cytat_tylera": "Pracujesz, ale nie wiesz po co.",
        }


def _generate_cv_photo(body: str, cv_data: dict, test_mode: bool = False) -> str | None:
    """
    Generuje zdjęcie profilowe do CV przez FLUX.
    Prompt budowany lokalnie (bez AI) — oszczędność 1 calla.
    Zwraca base64 PNG lub None.
    """
    try:
        with open(CV_PHOTO_FLUX_PATH, encoding="utf-8") as f:
            photo_cfg = json.load(f)
    except Exception as e:
        logger.warning("[cv-photo] Brak zwykly_cv_photo_flux.json: %s", e)
        photo_cfg = {}

    style_base = photo_cfg.get(
        "style_base", "professional CV headshot portrait, sharp focus"
    )

    imie = (
        cv_data.get("imie_nazwisko", "unknown person") if cv_data else "unknown person"
    )
    tytul = cv_data.get("tytul_zawodowy", "") if cv_data else ""
    plec = _detect_gender(body, imie)
    plec_en = {"kobieta": "woman", "mezczyzna": "man"}.get(plec, "person")

    # Prompt budowany lokalnie — bez AI
    photo_prompt = (
        f"Portrait of a {plec_en}, {tytul or 'office worker'}, "
        f"Fight Club 1999 aesthetic, exhausted, slightly damaged look, "
        f"professional headshot, dramatic lighting, film grain, {style_base}"
    )
    logger.info("[cv-photo] Prompt (lokalny): %.150s", photo_prompt)

    img = _generate_flux_image(photo_prompt, panel_index=98, test_mode=test_mode)
    if img and img.get("base64"):
        try:
            from PIL import Image as PILImage

            raw = base64.b64decode(img["base64"])
            pil = PILImage.open(io.BytesIO(raw)).convert("RGB")
            w, h = pil.size
            side = min(w, h)
            left = (w - side) // 2
            top = (h - side) // 2
            pil = pil.crop((left, top, left + side, top + side))
            pil = pil.resize((300, 300), PILImage.LANCZOS)
            buf = io.BytesIO()
            pil.save(buf, format="PNG")
            return base64.b64encode(buf.getvalue()).decode("ascii")
        except Exception as e:
            logger.warning("[cv-photo] Błąd resize: %s", e)
            return img["base64"]
    return None


def _build_cv_pdf(cv_data: dict, photo_b64: str | None) -> str | None:
    """
    Buduje PDF CV z reportlab z polskimi znakami (UTF-8).
    Zdjęcie w prawym górnym rogu.
    Zwraca base64 PDF lub None przy błędzie.
    """
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.utils import ImageReader
    except ImportError as e:
        logger.error("[cv-pdf] Brak reportlab: %s", e)
        return None

    FN, FB = _register_fonts()
    logger.info("[cv-pdf] Czcionki: FN=%s FB=%s", FN, FB)

    buf = io.BytesIO()
    W, H = A4
    c = rl_canvas.Canvas(buf, pagesize=A4)

    BLACK = (0.05, 0.05, 0.05)
    DARK = (0.15, 0.15, 0.15)
    GRAY = (0.45, 0.45, 0.45)
    LGRAY = (0.85, 0.85, 0.85)
    RED = (0.7, 0.1, 0.1)
    WHITE = (1.0, 1.0, 1.0)

    def set_color(rgb):
        c.setFillColorRGB(*rgb)

    def draw_text(txt, x, y, font=FN, size=10, color=BLACK, max_width=None):
        set_color(color)
        c.setFont(font, size)
        effective_width = max_width if max_width is not None else col_width
        if effective_width and effective_width < (right_margin - x):
            # Zawijanie — zawsze tnij do szerokości strony
            effective_width = min(effective_width, right_margin - x)
        words = str(txt).split()
        line = ""
        lines = []
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, font, size) <= effective_width:
                line = test
            else:
                if line:
                    lines.append(line)
                line = w
        if line:
            lines.append(line)
        for i, ln in enumerate(lines):
            c.drawString(x, y - i * (size + 2), ln)
        return len(lines) * (size + 2)

    c.setFillColorRGB(*BLACK)
    c.rect(0, H - 45 * mm, W, 45 * mm, fill=1, stroke=0)

    imie = cv_data.get("imie_nazwisko", "Anonim Bezdomny")
    set_color(WHITE)
    c.setFont(FB, 22)
    c.drawString(15 * mm, H - 20 * mm, imie)

    tytul = cv_data.get("tytul_zawodowy", "")
    set_color((0.8, 0.8, 0.8))
    c.setFont(FN, 11)
    c.drawString(15 * mm, H - 30 * mm, tytul)

    email_str = cv_data.get("email", "")
    tel_str = cv_data.get("telefon", "")
    miasto = cv_data.get("miasto", "")
    kontakt = " | ".join(filter(None, [email_str, tel_str, miasto]))
    set_color((0.65, 0.65, 0.65))
    c.setFont(FN, 9)
    c.drawString(15 * mm, H - 39 * mm, kontakt)

    if photo_b64:
        try:
            photo_bytes = base64.b64decode(photo_b64)
            photo_reader = ImageReader(io.BytesIO(photo_bytes))
            photo_size = 38 * mm
            c.drawImage(
                photo_reader,
                W - photo_size - 10 * mm,
                H - photo_size - 3.5 * mm,
                width=photo_size,
                height=photo_size,
                preserveAspectRatio=True,
                mask="auto",
            )
        except Exception as e:
            logger.warning("[cv-pdf] Błąd wklejania zdjęcia: %s", e)

    c.setStrokeColorRGB(*RED)
    c.setLineWidth(2)
    c.line(15 * mm, H - 48 * mm, W - 15 * mm, H - 48 * mm)

    y = H - 58 * mm
    left_margin = 15 * mm
    right_margin = W - 15 * mm
    # Jeśli jest zdjęcie, tekst nie może wchodzić pod zdjęcie w nagłówku
    # Zdjęcie zajmuje 38mm + 10mm margines = 48mm od prawej krawędzi
    photo_col_width = (
        (W - (38 * mm + 10 * mm + 15 * mm)) - left_margin
        if photo_b64
        else (right_margin - left_margin)
    )
    col_width = right_margin - left_margin  # pełna szerokość dla sekcji pod nagłówkiem

    def section_header(title, ypos):
        c.setFont(FB, 11)
        c.setFillColorRGB(*RED)
        c.drawString(left_margin, ypos, title.upper())
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(0.5)
        c.line(left_margin, ypos - 2, right_margin, ypos - 2)
        return ypos - 8 * mm

    def check_page_break(ypos, needed=20 * mm):
        if ypos < needed:
            c.showPage()
            return H - 20 * mm
        return ypos

    podsumowanie = cv_data.get("podsumowanie", "")
    if podsumowanie:
        y = section_header("Podsumowanie zawodowe", y)
        c.setFont(FN, 10)
        c.setFillColorRGB(*DARK)
        words = podsumowanie.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, FN, 10) <= col_width:
                line = test
            else:
                c.drawString(left_margin, y, line)
                y -= 5 * mm
                line = w
                y = check_page_break(y)
        if line:
            c.drawString(left_margin, y, line)
            y -= 5 * mm
        y -= 3 * mm

    doswiadczenie = cv_data.get("doswiadczenie", [])
    if doswiadczenie:
        y = check_page_break(y, 40 * mm)
        y = section_header("Doświadczenie zawodowe", y)
        for job in doswiadczenie:
            y = check_page_break(y, 30 * mm)
            firma = job.get("firma", "")
            stanowisko = job.get("stanowisko", "")
            okres = job.get("okres", "")
            obowiazki = job.get("obowiazki", [])

            c.setFont(FB, 10)
            c.setFillColorRGB(*BLACK)
            c.drawString(left_margin, y, firma)
            c.setFont(FN, 10)
            c.setFillColorRGB(*GRAY)
            c.drawRightString(right_margin, y, okres)
            y -= 5 * mm

            c.setFont(FN, 10)
            c.setFillColorRGB(*DARK)
            c.drawString(left_margin + 2 * mm, y, stanowisko)
            y -= 5 * mm

            c.setFont(FN, 9)
            c.setFillColorRGB(*DARK)
            for ob in obowiazki:
                y = check_page_break(y)
                words_ob = f"• {ob}".split()
                line_ob = ""
                for w in words_ob:
                    test = (line_ob + " " + w).strip()
                    if c.stringWidth(test, FN, 9) <= col_width - 4 * mm:
                        line_ob = test
                    else:
                        if line_ob:
                            c.drawString(left_margin + 4 * mm, y, line_ob)
                            y -= 4.5 * mm
                            y = check_page_break(y)
                        line_ob = w
                if line_ob:
                    c.drawString(left_margin + 4 * mm, y, line_ob)
                    y -= 4.5 * mm
            y -= 3 * mm

    wyksztalcenie = cv_data.get("wyksztalcenie", [])
    if wyksztalcenie:
        y = check_page_break(y, 25 * mm)
        y = section_header("Wykształcenie", y)
        for edu in wyksztalcenie:
            uczelnia = edu.get("uczelnia", "")
            kierunek = edu.get("kierunek", "")
            rok = edu.get("rok", "")
            c.setFont(FB, 10)
            c.setFillColorRGB(*BLACK)
            c.drawString(left_margin, y, uczelnia)
            c.setFont(FN, 9)
            c.setFillColorRGB(*GRAY)
            c.drawRightString(right_margin, y, str(rok))
            y -= 5 * mm
            c.setFont(FN, 10)
            c.setFillColorRGB(*DARK)
            c.drawString(left_margin + 2 * mm, y, kierunek)
            y -= 7 * mm

    umiejetnosci = cv_data.get("umiejetnosci", [])
    if umiejetnosci:
        y = check_page_break(y, 20 * mm)
        y = section_header("Umiejętności", y)
        half = len(umiejetnosci) // 2 + len(umiejetnosci) % 2
        col1 = umiejetnosci[:half]
        col2 = umiejetnosci[half:]
        col_w2 = col_width / 2
        y_start = y
        c.setFont(FN, 9)
        c.setFillColorRGB(*DARK)
        for i, um in enumerate(col1):
            c.drawString(left_margin, y_start - i * 5 * mm, f"• {um}")
        for i, um in enumerate(col2):
            c.drawString(left_margin + col_w2, y_start - i * 5 * mm, f"• {um}")
        y = y_start - max(len(col1), len(col2)) * 5 * mm - 3 * mm

    jezyki = cv_data.get("jezyki", [])
    if jezyki:
        y = check_page_break(y, 15 * mm)
        y = section_header("Języki", y)
        c.setFont(FN, 9)
        c.setFillColorRGB(*DARK)
        for j in jezyki:
            c.drawString(left_margin, y, f"• {j}")
            y -= 4.5 * mm
        y -= 3 * mm

    zainteresowania = cv_data.get("zainteresowania", [])
    if zainteresowania:
        y = check_page_break(y, 15 * mm)
        y = section_header("Zainteresowania", y)
        c.setFont(FN, 9)
        c.setFillColorRGB(*DARK)
        line_z = " | ".join(zainteresowania)
        words_z = line_z.split()
        cur_z = ""
        for w in words_z:
            test = (cur_z + " " + w).strip()
            if c.stringWidth(test, FN, 9) <= col_width:
                cur_z = test
            else:
                c.drawCentredString(W / 2, y, f'"{cur_z}"')
                y -= 4 * mm
                cur_z = w
        if cur_z:
            c.drawCentredString(W / 2, y, f'"{cur_z}"')

    # ── ŻYCIORYS ──────────────────────────────────────────────────────────────────
    zyciorys = cv_data.get("zyciorys", "")
    if zyciorys:
        y = check_page_break(y, 25 * mm)
        y = section_header("Życiorys", y)
        c.setFont(FN, 10)
        c.setFillColorRGB(*DARK)
        safe_w = col_width - 4 * mm  # margines bezpieczeństwa dla polskich znaków
        words = zyciorys.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, FN, 10) <= safe_w:
                line = test
            else:
                c.drawString(left_margin, y, line)
                y -= 5 * mm
                line = w
                y = check_page_break(y)
        if line:
            c.drawString(left_margin, y, line)
            y -= 8 * mm

    cytat = cv_data.get("cytat_tylera", "")
    if cytat:
        y = check_page_break(y, 20 * mm)
        c.setStrokeColorRGB(*LGRAY)
        c.setLineWidth(0.5)
        c.line(left_margin, y + 3 * mm, right_margin, y + 3 * mm)
        y -= 3 * mm
        c.setFont(FN, 8)
        c.setFillColorRGB(*RED)
        safe_w8 = col_width - 4 * mm
        words = cytat.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(test, FN, 8) <= safe_w8:
                line = test
            else:
                c.drawCentredString(W / 2, y, f"— {line}")
                y -= 4 * mm
                line = w
        if line:
            c.drawString(left_margin, y, f"— {line}")

    c.save()
    pdf_bytes = buf.getvalue()
    logger.info("[cv-pdf] PDF wygenerowany: %d B", len(pdf_bytes))
    ts_cv = datetime.now().strftime("%Y%m%d_%H%M%S")
    return {
        "base64": base64.b64encode(pdf_bytes).decode("ascii"),
        "content_type": "application/pdf",
        "filename": f"cv_tylera_{ts_cv}.pdf",
    }


def _build_explanation_txt(res_text: str, body: str) -> dict | None:
    """
    Generuje plik wyjaśnienie.txt — DeepSeek tłumaczy każde zdanie
    Tylera i Sokratesa prostym językiem po polsku.
    Zwraca dict {base64, content_type, filename} lub None przy błędzie.
    """
    if not res_text or not res_text.strip():
        return None

    system_msg = (
        "Jesteś pomocnym asystentem który wyjaśnia odpowiedzi Tylera Durdena i Sokratesa. "
        "Otrzymasz odpowiedź napisaną do nadawcy emaila. "
        "Twoje zadanie: wyjaśnij PO POLSKU każde zdanie lub akapit z tej odpowiedzi — "
        "co autor miał na myśli, dlaczego tak napisał, do czego nawiązuje. "
        "Pisz prosto i zrozumiale, jakbyś tłumaczył przyjacielowi. "
        "Zachowaj kolejność — najpierw wyjaśnij Sokratesa, potem Tylera. "
        "Dla każdego zdania/akapitu napisz: ZDANIE: [cytat] → WYJAŚNIENIE: [co to znaczy]. "
        "Nie używaj markdownu. Tylko czysty tekst."
    )

    user_msg = (
        f"Email który otrzymał program (kontekst):\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"Odpowiedź do wyjaśnienia:\n{res_text}"
    )

    raw, provider = _call_ai_with_fallback(system_msg, user_msg, max_tokens=3000)

    if not raw or not raw.strip():
        logger.warning("[zwykly] Brak wyjaśnienia od AI")
        return None

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"wyjasnienie_{ts}.txt"

    # Nagłówek pliku
    header = (
        f"=== WYJAŚNIENIE ODPOWIEDZI TYLERA I SOKRATESA ===\n"
        f"Data: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        f"Provider: {provider}\n"
        f"{'=' * 50}\n\n"
    )

    content = header + raw.strip()
    b64 = base64.b64encode(content.encode("utf-8")).decode("ascii")

    logger.info("[zwykly] Wyjaśnienie wygenerowane: %d znaków", len(content))

    return {
        "base64": b64,
        "content_type": "text/plain",
        "filename": filename,
    }


# ═══════════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════════


def _build_ankieta(res_text: str, body: str) -> tuple[dict | None, dict | None]:
    """
    Generuje ankietę wiedzy o odpowiedzi Tylera.
    Zwraca (html_dict, pdf_dict) lub (None, None) przy błędzie.
    """
    try:
        with open(ANKIETA_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[ankieta] Brak JSON: %s", e)
        return None, None

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    user_msg = (
        f"Odpowiedź Tylera do nadawcy:\n{res_text}\n\n"
        f"Email nadawcy (kontekst):\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"SCHEMAT JSON — użyj DOKŁADNIE tych kluczy:\n{__import__('json').dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"Zwróć TYLKO czysty JSON. Klucz listy pytań MUSI być 'pytania'."
    )

    raw = call_deepseek(_js(system_msg), _ju(user_msg), MODEL_TYLER, max_tokens=4500)

    if not raw:
        logger.warning("[ankieta] Brak danych od AI")
        return None, None

    logger.info("[ankieta] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        data = _parse_json_safe(raw, "ankieta")
        if data is None:
            raise ValueError("JSON nienaprawialny")
        # AI zwróciło tablicę pytań bez wrappera — owijamy
        if isinstance(data, list):
            if len(data) > 0 and isinstance(data[0], dict):
                logger.warning(
                    "[ankieta] AI zwróciło listę — owijam w {pytania: [...]}"
                )
                data = {"pytania": data}
            else:
                raise ValueError(
                    f"Oczekiwano dict, dostałem list (pusta lub bez dictów)"
                )
        if not isinstance(data, dict):
            raise ValueError(f"Oczekiwano dict, dostałem {type(data).__name__}")
        if not data.get("pytania"):
            # Szukaj listy pytań pod zagnieżdżonymi kluczami
            for v in data.values():
                if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                    data["pytania"] = v
                    logger.info("[ankieta] wyciągnięto pytania z zagnieżdżonej listy")
                    break
        if not data.get("pytania"):
            logger.warning("[ankieta] JSON OK ale brak pytań — raw: %.200s", raw)
            return None, None
    except Exception as e:
        logger.warning("[ankieta] Błąd JSON: %s | raw: %.200s", e, raw)
        return None, None

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    tytul = data.get("tytul", "Test Tylera Durdena")
    pytania = data.get("pytania", [])

    # ── Buduj HTML ────────────────────────────────────────────────────────────
    html = f"""<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="UTF-8">
<title>{tytul}</title>
<style>
  body {{ font-family: 'Courier New', monospace; background: #0a0a0a; color: #e0d0b0; margin: 0; padding: 20px; }}
  h1 {{ color: #8b0000; text-align: center; font-size: 1.8em; border-bottom: 2px solid #8b0000; padding-bottom: 10px; }}
  .intro {{ color: #888; font-style: italic; text-align: center; margin: 20px 0; }}
  .pytanie {{ background: #111; border-left: 4px solid #8b0000; margin: 20px 0; padding: 15px; border-radius: 0 4px 4px 0; }}
  .pytanie h3 {{ color: #c8b89a; margin: 0 0 8px 0; font-size: 0.9em; }}
  .cytat {{ color: #666; font-style: italic; font-size: 0.85em; margin-bottom: 10px; }}
  .opcje label {{ display: block; margin: 8px 0; cursor: pointer; }}
  .opcje input {{ margin-right: 8px; accent-color: #8b0000; }}
  .wyjasnienie {{ display: none; background: #1a0a0a; border: 1px solid #8b0000; padding: 10px; margin-top: 10px; font-size: 0.85em; color: #c8b89a; }}
  button {{ background: #8b0000; color: white; border: none; padding: 12px 30px; font-size: 1em; cursor: pointer; margin: 20px auto; display: block; font-family: 'Courier New', monospace; }}
  button:hover {{ background: #a00000; }}
  #wynik {{ text-align: center; font-size: 1.2em; color: #8b0000; margin: 20px; display: none; }}
  .nr {{ color: #8b0000; font-weight: bold; }}
</style>
</head>
<body>
<h1>{tytul}</h1>
<p class="intro">{data.get("wprowadzenie", "")}</p>
<form id="quiz">
"""
    for p in pytania:
        nr = p.get("nr", "?")
        cytat = p.get("cytat_tylera", "")
        pytanie = p.get("pytanie", "")
        odp = p.get("odpowiedzi", {})
        if isinstance(odp, list):
            # model zwrócił listę [{"klucz":"a","tresc":"..."}] zamiast {"a":"..."}
            odp = {
                str(item.get("klucz", item.get("key", chr(97 + i)))): str(
                    item.get("tresc", item.get("text", ""))
                )
                for i, item in enumerate(odp)
            }
        elif not isinstance(odp, dict):
            odp = {}
        wyjasnienie = p.get("wyjasnienie", "")
        html += f"""
<div class="pytanie">
  <h3><span class="nr">Pytanie {nr}:</span> {pytanie}</h3>
  <div class="cytat">"{cytat}"</div>
  <div class="opcje">
    <label><input type="radio" name="q{nr}" value="a"> a) {odp.get("a", "")}</label>
    <label><input type="radio" name="q{nr}" value="b"> b) {odp.get("b", "")}</label>
    <label><input type="radio" name="q{nr}" value="c"> c) {odp.get("c", "")}</label>
  </div>
  <div class="wyjasnienie" id="w{nr}">{wyjasnienie}</div>
</div>"""

    zakonczenie = data.get("zakonczenie", "— Tyler Durden")
    html += f"""
</form>
<button onclick="sprawdz()">Sprawdź wynik</button>
<div id="wynik"></div>
<p style="text-align:center;color:#666;font-style:italic;margin-top:40px">{zakonczenie}</p>
<script>
function sprawdz() {{
  var poprawne = 0;
  var total = {len(pytania)};
  for (var i = 1; i <= total; i++) {{
    var sel = document.querySelector('input[name="q'+i+'"]:checked');
    var wyn = document.getElementById('w'+i);
    if (sel) {{
      if (sel.value === 'b') {{ poprawne++; wyn.style.background='#0a1a0a'; }}
      else {{ wyn.style.background='#1a0a0a'; }}
      wyn.style.display = 'block';
    }}
  }}
  var wynikDiv = document.getElementById('wynik');
  wynikDiv.style.display = 'block';
  wynikDiv.innerHTML = 'Wynik: ' + poprawne + '/' + total + ' — ' + 
    (poprawne < 4 ? 'Nie rozumiesz nic. Typowe.' : 
     poprawne < 7 ? 'Trochę rozumiesz. To niepokojące.' : 
     'Rozumiesz Tylera. Powinieneś się tym martwić.');
}}
</script>
</body>
</html>"""

    html_dict = _to_zip(html.encode("utf-8"), f"ankieta_{ts}.html", f"ankieta_{ts}.zip")

    # ── Buduj PDF AcroForm (interaktywny z checkboxami) ──────────────────────
    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm

        FN, FB = _register_fonts()

        buf = io.BytesIO()
        W, H = A4
        c = rl_canvas.Canvas(buf, pagesize=A4)
        c.setTitle(tytul)
        lm, rm = 15 * mm, W - 15 * mm
        cw = rm - lm
        RED = (0.6, 0.1, 0.1)
        DARK = (0.1, 0.1, 0.1)
        GRAY = (0.4, 0.4, 0.4)

        # Obramowanie karty
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(3)
        c.rect(8 * mm, 8 * mm, W - 16 * mm, H - 16 * mm, fill=0, stroke=1)
        c.setLineWidth(1)
        c.rect(10 * mm, 10 * mm, W - 20 * mm, H - 20 * mm, fill=0, stroke=1)

        # Nagłówek ankiety
        c.setFillColorRGB(*DARK)
        c.rect(10 * mm, H - 40 * mm, W - 20 * mm, 30 * mm, fill=1, stroke=0)
        c.setFont(FB, 8)
        c.setFillColorRGB(0.6, 0.6, 0.6)
        c.drawCentredString(W / 2, H - 18 * mm, "TEST WIEDZY — TYLER DURDEN")

        c.setFont(FB, 16)
        c.setFillColorRGB(1, 1, 1)
        tytul_pdf = tytul[:60] if tytul else "Test Tylera Durdena"
        c.drawCentredString(W / 2, H - 28 * mm, tytul_pdf)
        c.setFont(FN, 9)
        c.setFillColorRGB(0.7, 0.5, 0.5)
        c.drawCentredString(
            W / 2, H - 35 * mm, f"{len(pytania)} pytań | odpowiedź poprawna: b"
        )

        y = H - 50 * mm

        # Wprowadzenie
        wprowadzenie = data.get("wprowadzenie", "")
        if wprowadzenie:
            c.setFont(FN, 9)
            c.setFillColorRGB(*GRAY)
            words = wprowadzenie.split()
            line_txt = ""
            for w in words:
                test = (line_txt + " " + w).strip()
                if c.stringWidth(test, FN, 9) <= cw:
                    line_txt = test
                else:
                    c.drawString(lm, y, line_txt)
                    y -= 5 * mm
                    line_txt = w
            if line_txt:
                c.drawString(lm, y, line_txt)
            y -= 8 * mm

        # Pytania
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(0.5)
        for p in pytania:
            if y < 30 * mm:
                c.showPage()
                y = H - 20 * mm
            nr = p.get("nr", "?")
            pytanie_txt = p.get("pytanie", "")
            odp = p.get("odpowiedzi", {})
            if isinstance(odp, list):
                odp = {
                    str(it.get("klucz", chr(97 + i))): str(it.get("tresc", ""))
                    for i, it in enumerate(odp)
                }
            wyjasnienie = p.get("wyjasnienie", "")

            # Nagłówek pytania
            c.setFont(FB, 9)
            c.setFillColorRGB(*RED)
            c.drawString(lm, y, f"Pytanie {nr}:")
            y -= 4 * mm

            # Treść pytania — zawijanie
            c.setFont(FN, 8)
            c.setFillColorRGB(*DARK)
            words = pytanie_txt.split()
            line_txt = ""
            for w in words:
                test = (line_txt + " " + w).strip()
                if c.stringWidth(test, FN, 8) <= cw:
                    line_txt = test
                else:
                    c.drawString(lm + 3 * mm, y, line_txt)
                    y -= 4 * mm
                    line_txt = w
            if line_txt:
                c.drawString(lm + 3 * mm, y, line_txt)
            y -= 5 * mm

            # Odpowiedzi a/b/c
            for klucz in ("a", "b", "c"):
                tresc = str(odp.get(klucz, ""))
                if not tresc:
                    continue
                marker = "◆" if klucz == "b" else "○"
                c.setFont(FN, 8)
                c.setFillColorRGB(*DARK)
                line_txt = f"{marker} {klucz}) {tresc}"
                words = line_txt.split()
                txt_line = ""
                for w in words:
                    test = (txt_line + " " + w).strip()
                    if c.stringWidth(test, FN, 8) <= cw - 5 * mm:
                        txt_line = test
                    else:
                        c.drawString(lm + 5 * mm, y, txt_line)
                        y -= 4 * mm
                        txt_line = w
                if txt_line:
                    c.drawString(lm + 5 * mm, y, txt_line)
                y -= 4 * mm

            # Wyjaśnienie (mniejsze, szare)
            if wyjasnienie:
                c.setFont(FN, 7)
                c.setFillColorRGB(*GRAY)
                words = wyjasnienie.split()
                line_txt = "→ "
                for w in words:
                    test = (line_txt + " " + w).strip()
                    if c.stringWidth(test, FN, 7) <= cw - 3 * mm:
                        line_txt = test
                    else:
                        c.drawString(lm + 3 * mm, y, line_txt)
                        y -= 3.5 * mm
                        line_txt = w
                if line_txt:
                    c.drawString(lm + 3 * mm, y, line_txt)
                y -= 4 * mm

            c.line(lm, y, rm, y)
            y -= 5 * mm

        # Zakończenie
        zakonczenie = data.get("zakonczenie", "— Tyler Durden")
        if y < 20 * mm:
            c.showPage()
            y = H - 20 * mm
        c.setFont(FN, 9)
        c.setFillColorRGB(*RED)
        c.drawCentredString(W / 2, y, f'"{zakonczenie}"')

        c.save()
        pdf_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        pdf_dict = {
            "base64": pdf_b64,
            "content_type": "application/pdf",
            "filename": f"ankieta_{ts}.pdf",
        }
        logger.info("[ankieta] PDF OK: %d pytań", len(pytania))
        return html_dict, pdf_dict

    except Exception as e:
        logger.error("[ankieta] Błąd PDF: %s", e)
        return html_dict, None


# ═══════════════════════════════════════════════════════════════════════════════
# HOROSKOP PDF — styl gazety lat 60
# ═══════════════════════════════════════════════════════════════════════════════


def _build_horoskop(body: str, res_text: str) -> dict | None:
    """Generuje horoskop nihilistyczny na 7 dni w stylu gazety lat 60."""
    try:
        with open(HOROSKOP_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[horoskop] Brak JSON: %s", e)
        return None

    # Oblicz daty
    today = datetime.now()
    daty = [
        (
            today.replace(day=today.day) + __import__("datetime").timedelta(days=i)
        ).strftime("%d.%m.%Y")
        for i in range(7)
    ]

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    daty_str = "\n".join(f"Dzień {i + 1} ({d})" for i, d in enumerate(daty))
    user_msg = (
        f"Email nadawcy:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"Odpowiedź Tylera (kontekst):\n{res_text[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"WAŻNE: W polu 'data' każdego dnia użyj DOKŁADNIE tych dat:\n{daty_str}\n\n"
        f"SCHEMAT JSON — użyj DOKŁADNIE tych kluczy:\n{json.dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"Zwróć TYLKO czysty JSON. Klucz listy dni MUSI być 'dni'."
    )

    raw = call_deepseek(_js(system_msg), _ju(user_msg), MODEL_TYLER, max_tokens=4000)
    if not raw:
        return None

    logger.info("[horoskop] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        data = _parse_json_safe(raw, "horoskop")
        if data is None:
            raise ValueError("[horoskop] JSON nienaprawialny")
        # AI zwróciło tablicę dni bez wrappera — owijamy
        if isinstance(data, list):
            if len(data) > 0 and isinstance(data[0], dict):
                logger.warning("[horoskop] AI zwróciło listę — owijam w {dni: [...]}")
                data = {"dni": data}
            else:
                raise ValueError(
                    f"[horoskop] Oczekiwano dict, dostałem list (pusta lub bez dictów)"
                )
        if not isinstance(data, dict):
            raise ValueError(
                f"[horoskop] Oczekiwano dict, dostałem {type(data).__name__}"
            )
        KEY_MAP_HOROSKOP = {
            "horoskop": "dni",
            "days": "dni",
            "forecast": "dni",
            "prognozy": "dni",
            "przepowiednie": "dni",
            "lista": "dni",
        }
        for wrong, right in KEY_MAP_HOROSKOP.items():
            if wrong in data and right not in data:
                data[right] = data.pop(wrong)
                logger.info("[horoskop] znormalizowano '%s' → '%s'", wrong, right)
        if not data.get("dni"):
            for v in data.values():
                if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                    data["dni"] = v
                    logger.info("[horoskop] wyciągnięto dni z zagnieżdżonej listy")
                    break
        if not data.get("dni"):
            logger.warning("[horoskop] JSON OK ale brak dni — raw: %.200s", raw)
            return None
    except Exception as e:
        logger.warning("[horoskop] Błąd JSON: %s | raw: %.200s", e, raw)
        return None

    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        FN, FB = _register_fonts()

        buf = io.BytesIO()
        W, H = A4
        c = rl_canvas.Canvas(buf, pagesize=A4)
        lm, rm = 12 * mm, W - 12 * mm
        cw = rm - lm
        RED = (0.6, 0.1, 0.1)
        DARK = (0.1, 0.1, 0.1)
        GRAY = (0.4, 0.4, 0.4)

        # Obramowanie karty
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(3)
        c.rect(8 * mm, 8 * mm, W - 16 * mm, H - 16 * mm, fill=0, stroke=1)
        c.setLineWidth(1)
        c.rect(10 * mm, 10 * mm, W - 20 * mm, H - 20 * mm, fill=0, stroke=1)

        # Nagłówek horoskopu
        c.setFillColorRGB(*DARK)
        c.rect(10 * mm, H - 40 * mm, W - 20 * mm, 30 * mm, fill=1, stroke=0)
        c.setFont(FB, 8)
        c.setFillColorRGB(0.6, 0.6, 0.6)
        c.drawCentredString(
            W / 2, H - 18 * mm, "GAZETA NIHILISTYCZNA — HOROSKOP TYLERA"
        )

        znak = str(data.get("znak_zodiaku", "Nieznany Znak"))[:50]
        c.setFont(FB, 16)
        c.setFillColorRGB(1, 1, 1)
        c.drawCentredString(W / 2, H - 28 * mm, znak)
        motto = str(data.get("motto", ""))[:80]
        c.setFont(FN, 9)
        c.setFillColorRGB(0.7, 0.5, 0.5)
        c.drawCentredString(W / 2, H - 35 * mm, motto)

        y = H - 50 * mm

        def wrap_text(txt, font, size, max_w, x, y_pos, color=DARK, indent=0):
            c.setFont(font, size)
            c.setFillColorRGB(*color)
            words = str(txt).split()
            line_txt = ""
            for w in words:
                test = (line_txt + " " + w).strip()
                if c.stringWidth(test, font, size) <= max_w:
                    line_txt = test
                else:
                    c.drawString(x + indent, y_pos, line_txt)
                    y_pos -= size + 2
                    line_txt = w
            if line_txt:
                c.drawString(x + indent, y_pos, line_txt)
                y_pos -= size + 2
            return y_pos

        # Dni
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(0.5)
        for dzien in data.get("dni", []):
            if y < 35 * mm:
                c.showPage()
                y = H - 20 * mm

            nr = dzien.get("dzien", "?")
            data_str = dzien.get("data", "")
            naglowek = str(dzien.get("naglowek", ""))
            tresc = str(dzien.get("tresc", ""))
            rada = str(dzien.get("rada_tylera", ""))

            # Dzień nagłówek
            c.setFont(FB, 10)
            c.setFillColorRGB(*RED)
            c.drawString(lm, y, f"DZIEŃ {nr}  {data_str}")
            y -= 4 * mm
            c.line(lm, y, rm, y)
            y -= 4 * mm

            # Nagłówek sensacyjny
            c.setFont(FB, 9)
            c.setFillColorRGB(*DARK)
            y = wrap_text(naglowek, FB, 9, cw, lm, y, color=DARK)
            y -= 2 * mm

            # Treść
            y = wrap_text(tresc, FN, 8, cw, lm, y, color=GRAY, indent=2 * mm)
            y -= 2 * mm

            # Rada Tylera
            if rada:
                c.setFont(FN, 8)
                c.setFillColorRGB(*RED)
                y = wrap_text(
                    f"→ {rada}", FN, 8, cw - 3 * mm, lm + 3 * mm, y, color=RED
                )

            y -= 6 * mm

        # Przepowiednia ogólna
        przepowiednia = str(data.get("przepowiednia_ogolna", ""))
        if przepowiednia:
            if y < 25 * mm:
                c.showPage()
                y = H - 20 * mm
            c.setStrokeColorRGB(*RED)
            c.line(lm, y, rm, y)
            y -= 5 * mm
            c.setFont(FB, 9)
            c.setFillColorRGB(*RED)
            c.drawString(lm, y, "PRZEPOWIEDNIA OGÓLNA:")
            y -= 5 * mm
            y = wrap_text(przepowiednia, FN, 9, cw, lm, y, color=GRAY)

        c.save()
        pdf_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        logger.info("[horoskop] PDF OK: %d dni", len(data.get("dni", [])))
        return {
            "base64": pdf_b64,
            "content_type": "application/pdf",
            "filename": f"horoskop_{ts}.pdf",
        }

    except Exception as e:
        logger.error("[horoskop] Błąd PDF: %s", e)
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# KARTA RPG PDF
# ═══════════════════════════════════════════════════════════════════════════════


def _build_karta_rpg(body: str, res_text: str) -> dict | None:
    """Generuje kartę postaci RPG."""
    try:
        with open(KARTA_RPG_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[karta-rpg] Brak JSON: %s", e)
        return None

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    user_msg = (
        f"Email:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"Odpowiedź Tylera:\n{res_text[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"SCHEMAT JSON — użyj DOKŁADNIE tych polskich kluczy:\n{__import__('json').dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"Zwróć TYLKO czysty JSON. ZAKAZ angielskich kluczy (name/stats/age) — używaj nazwa_postaci/statystyki."
    )

    raw = call_deepseek(_js(system_msg), _ju(user_msg), MODEL_TYLER, max_tokens=3500)
    if not raw:
        logger.warning("[karta-rpg] Brak odpowiedzi od AI")
        return None

    logger.info("[karta-rpg] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        data = _parse_json_safe(raw, "karta-rpg")
        if data is None:
            raise ValueError("[karta-rpg] JSON nienaprawialny")
        if isinstance(data, list) and len(data) > 0:
            logger.warning("[karta-rpg] Model zwrócił listę — biorę pierwszy element")
            data = data[0]
        if not isinstance(data, dict):
            raise ValueError(
                f"[karta-rpg] Oczekiwano dict, dostałem {type(data).__name__}"
            )
        KEY_MAP_RPG = {
            "name": "nazwa_postaci",
            "character_name": "nazwa_postaci",
            "imie": "nazwa_postaci",
            "class": "klasa_postaci",
            "klasa": "klasa_postaci",
            "character_class": "klasa_postaci",
            "level": "poziom",
            "stats": "statystyki",
            "statistics": "statystyki",
            "attributes": "statystyki",
            "skills": "umiejetnosci_specjalne",
            "abilities": "umiejetnosci_specjalne",
            "equipment": "ekwipunek",
            "items": "ekwipunek",
            "weakness": "slabosci",
            "weaknesses": "slabosci",
            "quest": "quest_glowny",
            "main_quest": "quest_glowny",
            "quote": "cytat_postaci",
            "character_quote": "cytat_postaci",
        }
        for wrong, right in KEY_MAP_RPG.items():
            if wrong in data and right not in data:
                data[right] = data.pop(wrong)
                logger.info("[karta-rpg] znormalizowano '%s' → '%s'", wrong, right)
        if not data.get("nazwa_postaci") and not data.get("statystyki"):
            # Poluzowany warunek — wystarczy że jest jakikolwiek klucz z wartością
            has_any_content = any(
                v for v in data.values() if v not in (None, "", [], {})
            )
            if not has_any_content:
                logger.warning("[karta-rpg] JSON pusty — raw: %.200s", raw)
                return None
            logger.warning(
                "[karta-rpg] Brak nazwa_postaci/statystyki ale są inne dane — kontynuuję"
            )
    except Exception as e:
        logger.warning("[karta-rpg] Błąd JSON: %s | raw: %.200s", e, raw)
        return None

    try:
        from reportlab.pdfgen import canvas as rl_canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        FN, FB = _register_fonts()

        buf = io.BytesIO()
        W, H = A4
        c = rl_canvas.Canvas(buf, pagesize=A4)
        lm, rm = 15 * mm, W - 15 * mm
        cw = rm - lm
        RED = (0.6, 0.1, 0.1)
        DARK = (0.1, 0.1, 0.1)
        GRAY = (0.4, 0.4, 0.4)

        # Obramowanie karty
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(3)
        c.rect(8 * mm, 8 * mm, W - 16 * mm, H - 16 * mm, fill=0, stroke=1)
        c.setLineWidth(1)
        c.rect(10 * mm, 10 * mm, W - 20 * mm, H - 20 * mm, fill=0, stroke=1)

        # Nagłówek
        c.setFillColorRGB(*DARK)
        c.rect(10 * mm, H - 40 * mm, W - 20 * mm, 30 * mm, fill=1, stroke=0)
        c.setFont(FB, 8)
        c.setFillColorRGB(0.6, 0.6, 0.6)
        c.drawCentredString(W / 2, H - 18 * mm, "KARTA POSTACI — PROJEKT TYLER DURDEN")
        c.setFont(FB, 18)
        c.setFillColorRGB(1, 1, 1)
        c.drawCentredString(
            W / 2, H - 28 * mm, data.get("nazwa_postaci", "ANONIM")[:30]
        )
        c.setFont(FN, 10)
        c.setFillColorRGB(0.7, 0.5, 0.5)
        c.drawCentredString(W / 2, H - 35 * mm, data.get("klasa_postaci", "")[:50])

        y = H - 50 * mm

        # Poziom
        poziom = data.get("poziom", "?")
        c.setFont(FB, 11)
        c.setFillColorRGB(*RED)
        c.drawString(lm, y, f"POZIOM: {poziom}")
        c.setStrokeColorRGB(*RED)
        c.setLineWidth(0.5)
        c.line(lm, y - 2, rm, y - 2)
        y -= 8 * mm

        # Statystyki — 2 kolumny
        stats = data.get("statystyki", {})
        stat_list = list(stats.items())
        half = len(stat_list) // 2 + len(stat_list) % 2
        col1 = stat_list[:half]
        col2 = stat_list[half:]
        col_w = cw / 2 - 5 * mm

        c.setFont(FB, 9)
        c.setFillColorRGB(*RED)
        c.drawString(lm, y, "STATYSTYKI")
        y -= 5 * mm

        # Krok między statystykami — 18pt = etykieta(7) + wartość(7) + odstęp(4)
        STAT_STEP = 18
        col_half = cw / 2 - 3 * mm

        def draw_stat_col(items, x_base):
            sy = y_stat
            for sk, sv in items:
                label = sk.replace("_", " ").upper()
                c.setFont(FB, 7)
                c.setFillColorRGB(*DARK)
                c.drawString(x_base, sy, label + ":")
                sy -= 8
                # Zawijaj wartość jeśli długa
                val_str = str(sv)
                c.setFont(FN, 7)
                c.setFillColorRGB(*GRAY)
                words = val_str.split()
                line = ""
                for w in words:
                    test = (line + " " + w).strip()
                    if c.stringWidth(test, FN, 7) <= col_half:
                        line = test
                    else:
                        c.drawString(x_base + 2 * mm, sy, line)
                        sy -= 8
                        line = w
                if line:
                    c.drawString(x_base + 2 * mm, sy, line)
                sy -= STAT_STEP - 8
            return sy

        y_stat = y
        sy1 = draw_stat_col(col1, lm)
        sy2 = draw_stat_col(col2, lm + cw / 2)
        y = min(sy1, sy2) - 8 * mm

        # Umiejętności
        c.setFont(FB, 9)
        c.setFillColorRGB(*RED)
        c.drawString(lm, y, "UMIEJĘTNOŚCI SPECJALNE")
        c.line(lm, y - 2, rm, y - 2)
        y -= 6 * mm
        for um in data.get("umiejetnosci_specjalne", []):
            c.setFont(FN, 8)
            c.setFillColorRGB(*DARK)
            c.drawString(lm + 3 * mm, y, f"◆ {um}")
            y -= 5 * mm

        y -= 3 * mm

        # Ekwipunek
        c.setFont(FB, 9)
        c.setFillColorRGB(*RED)
        c.drawString(lm, y, "EKWIPUNEK")
        c.line(lm, y - 2, rm, y - 2)
        y -= 6 * mm
        for item in data.get("ekwipunek", []):
            c.setFont(FN, 8)
            c.setFillColorRGB(*DARK)
            c.drawString(lm + 3 * mm, y, f"⚔ {item}")
            y -= 5 * mm

        y -= 3 * mm

        # Quest + cytat
        c.setFont(FB, 9)
        c.setFillColorRGB(*RED)
        c.drawString(lm, y, "QUEST GŁÓWNY:")
        c.setFont(FN, 8)
        c.setFillColorRGB(*DARK)
        c.drawString(lm + 30 * mm, y, data.get("quest_glowny", ""))
        y -= 8 * mm

        # Cytat na dole
        c.setStrokeColorRGB(0.7, 0.7, 0.7)
        c.line(lm, y, rm, y)
        y -= 5 * mm
        c.setFont(FN, 8)
        c.setFillColorRGB(*RED)
        cytat = data.get("cytat_postaci", "")
        words = cytat.split()
        line = ""
        for w in words:
            test = (line + " " + w).strip()
            if c.stringWidth(f'"{test}"', FN, 8) <= cw:
                line = test
            else:
                c.drawCentredString(W / 2, y, f'"{line}"')
                y -= 4 * mm
                line = w
        if line:
            c.drawCentredString(W / 2, y, f'"{line}"')

        c.save()
        pdf_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        logger.info("[karta-rpg] OK")
        return {
            "base64": pdf_b64,
            "content_type": "application/pdf",
            "filename": f"karta_rpg_{ts}.pdf",
        }

    except Exception as e:
        logger.error("[karta-rpg] Błąd PDF: %s", e)
        return None


def _generate_psychiatric_photo(
    body: str, nouns_dict: dict, sender_name: str = "", test_mode: bool = False
) -> str | None:
    """
    Generuje zdjęcie pacjenta psychiatrycznego w kaftanie bezpieczeństwa przez FLUX.
    Używa promptu z zwykly_psychiatryczny_obrazek.json.
    Podmienia {{OBJECTS}} na rzeczowniki z emaila.
    Zwraca base64 JPG lub None.

    Parametr test_mode:
    - Jeśli test_mode=True (disable_flux=True z KEYWORDS_TEST),
      zwracamy zastępczy obrazek zamiast odpytywać tokeny HF.
    """
    # ── KEYWORDS_TEST (disable_flux) → test_mode ─────────────────────────────
    if test_mode:
        logger.info("[psych-photo] test_mode=True — pomijam FLUX, używam zastepczy.jpg")
        sub = _load_substitute_image()
        if sub:
            return sub.get("base64")
        return None

    try:
        with open(PSYCHIATRYCZNY_OBRAZEK_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[psych-photo] Brak zwykly_psychiatryczny_obrazek.json: %s", e)
        return None

    prompt_template = cfg.get("prompt_template", "")
    fallback_objects = cfg.get(
        "fallback_objects", "everyday objects, papers, worn shoes"
    )
    hf_params = cfg.get("hf_parameters", {})

    # ── Buduj listę obiektów z rzeczowników ──────────────────────────────────
    if nouns_dict:
        objects_list = list(nouns_dict.values())[:8]  # max 8 rzeczowników
        objects_str = ", ".join(objects_list)
    else:
        # fallback: wyciągnij z body regexem
        nouns_fallback = _extract_nouns_from_body(body)
        objects_str = (
            ", ".join(nouns_fallback[:6]) if nouns_fallback else fallback_objects
        )

    # ── Płeć — do opisu pacjenta ─────────────────────────────────────────────
    gender = _detect_gender(body, sender_name)
    gender_desc = {
        "kobieta": "female patient, woman",
        "mezczyzna": "male patient, man",
    }.get(gender, "patient")

    # ── Podmień {{OBJECTS}} w szablonie promptu ───────────────────────────────
    prompt = prompt_template.replace("{{OBJECTS}}", objects_str)
    # Podmień opcjonalne zmienne jeśli są w szablonie
    prompt = prompt.replace("{{GENDER}}", gender_desc)
    prompt = prompt.replace("{{NAME}}", sender_name or "unknown")

    logger.info("[psych-photo] Prompt (pierwsze 200 znaków): %.200s", prompt)
    logger.info("[psych-photo] Obiekty: %s | Płeć: %s", objects_str, gender)

    # ── Wywołaj FLUX z parametrami z JSON ────────────────────────────────────
    tokens = get_active_tokens()
    if not tokens:
        logger.error("[psych-photo] Brak tokenów HF")
        return None

    seed = random.randint(0, 2**32 - 1)
    payload = {
        "inputs": prompt,
        "parameters": {
            "num_inference_steps": hf_params.get("num_inference_steps", 4),
            "guidance_scale": hf_params.get("guidance_scale", 3.0),
            "width": hf_params.get("width", 768),
            "height": hf_params.get("height", 1024),
            "seed": seed,
        },
    }

    raw_img = None
    for name, token in tokens:
        headers = {"Authorization": f"Bearer {token}", "Accept": "image/png"}
        try:
            resp = requests.post(
                HF_API_URL, headers=headers, json=payload, timeout=HF_TIMEOUT
            )
            if resp.status_code == 200:
                raw_img = resp.content
                logger.info("[psych-photo] FLUX OK token=%s (%d B)", name, len(raw_img))
                break
            elif resp.status_code == 402:
                mark_dead(name)
                logger.warning(
                    "[psych-photo] 402 token=%s — wyczerpane kredyty, dodano do czarnej listy",
                    name,
                )
            elif resp.status_code in (401, 403):
                mark_dead(name)
                logger.warning(
                    "[psych-photo] HTTP %d token=%s — nieważny, dodano do czarnej listy",
                    resp.status_code,
                    name,
                )
            elif resp.status_code == 429:
                logger.warning("[psych-photo] 429 token=%s → następny", name)
            else:
                logger.warning("[psych-photo] HTTP %d token=%s", resp.status_code, name)
        except Exception as e:
            logger.warning("[psych-photo] Wyjątek token=%s: %s", name, e)

    if not raw_img:
        logger.error("[psych-photo] Wszystkie tokeny HF zawiodły")
        return None

    # ── Konwertuj PNG → JPG, zachowaj proporcje polaroid ─────────────────────
    try:
        from PIL import Image as PILImage

        pil = PILImage.open(io.BytesIO(raw_img)).convert("RGB")
        buf = io.BytesIO()
        pil.save(buf, format="JPEG", quality=92, optimize=True)
        b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        logger.info(
            "[psych-photo] Konwersja JPG OK (%dKB)", len(buf.getvalue()) // 1024
        )
        return b64
    except Exception as e:
        logger.warning("[psych-photo] Błąd konwersji: %s — zwracam PNG b64", e)
        return base64.b64encode(raw_img).decode("ascii")


# ═══════════════════════════════════════════════════════════════════════════════
# RAPORT PSYCHIATRYCZNY DOCX (zastępuje PDF)
# ═══════════════════════════════════════════════════════════════════════════════


def _build_raport_psychiatryczny(
    body: str,
    previous_body: str | None,
    res_text: str,
    nouns_dict: dict = None,
    sender_name: str = "",
    test_mode: bool = False,
) -> dict | None:
    """
    Generuje raport psychiatryczny jako DOCX (python-docx).
    Na końcu dokumentu wkleja zdjęcie FLUX pacjenta w kaftanie bezpieczeństwa.
    Zwraca dict {base64, content_type, filename} lub None.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error("[raport] Brak python-docx: %s", e)
        return None

    try:
        with open(RAPORT_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[raport] Brak JSON: %s", e)
        return None

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    context = f"EMAIL PACJENTA:\n{body[:MAX_DLUGOSC_EMAIL]}"
    if previous_body:
        context += f"\n\nPOPRZEDNI EMAIL (historia choroby):\n{previous_body[:MAX_DLUGOSC_EMAIL]}"
    context += f"\n\nODPOWIEDŹ TYLERA (materiał diagnostyczny):\n{res_text[:MAX_DLUGOSC_EMAIL]}"
    context += f"\n\nSCHEMAT JSON — użyj DOKŁADNIE tych kluczy:\n{json.dumps(schema, ensure_ascii=False, indent=2)}"
    context += "\n\nKLUCZ dane_pacjenta (dict) i diagnoza_wstepna MUSZĄ istnieć. Zwróć TYLKO czysty JSON."

    # DeepSeek dla raportu
    raw = call_deepseek(_js(system_msg), _ju(context), MODEL_TYLER)

    if not raw:
        logger.warning("[raport] Brak odpowiedzi od AI")
        return None

    logger.info("[raport] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        data = _parse_json_safe(raw, "raport")
        if data is None:
            raise ValueError("[raport] JSON nienaprawialny")
        if not isinstance(data, dict):
            raise ValueError(
                f"[raport] Oczekiwano dict, dostałem {type(data).__name__}"
            )
        KEY_MAP_RAPORT = {
            "pacjent": "dane_pacjenta",
            "patient": "dane_pacjenta",
            "patient_data": "dane_pacjenta",
            "dane": "dane_pacjenta",
            "diagnoza": "diagnoza_wstepna",
            "diagnosis": "diagnoza_wstepna",
            "primary_diagnosis": "diagnoza_wstepna",
            "rozpoznanie": "diagnoza_wstepna",
            "historia_choroby": "wywiad",
            "history": "wywiad",
            "powod": "powod_przyjecia",
            "reason": "powod_przyjecia",
            "symptoms": "objawy",
            "symptomy": "objawy",
            "recommendations": "zalecenia",
            "treatment": "zalecenia",
            "prognosis": "rokowanie",
            "notatka": "notatka_oddzialu",
            "note": "notatka_oddzialu",
        }
        for wrong, right in KEY_MAP_RAPORT.items():
            if wrong in data and right not in data:
                data[right] = data.pop(wrong)
                logger.info("[raport] znormalizowano '%s' → '%s'", wrong, right)
        if isinstance(data.get("dane_pacjenta"), str):
            data["dane_pacjenta"] = {"imie_nazwisko": data["dane_pacjenta"]}
        if not data.get("dane_pacjenta"):
            flat_keys = ["imie_nazwisko", "wiek", "zawod", "adres", "stan_cywilny"]
            found_flat = {k: data.pop(k) for k in flat_keys if k in data}
            if found_flat:
                data["dane_pacjenta"] = found_flat
        if not data.get("diagnoza_wstepna") and not data.get("dane_pacjenta"):
            logger.warning("[raport] JSON pusty — raw: %.200s", raw)
            return None
    except Exception as e:
        logger.warning("[raport] Błąd JSON: %s | raw: %.200s", e, raw)
        return None

    # ── Buduj DOCX ────────────────────────────────────────────────────────────
    try:
        doc = Document()

        # Marginesy
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        szpital_cfg = cfg.get("szpital", {})

        # ── Nagłówek szpitala ─────────────────────────────────────────────────
        h = doc.add_heading(
            szpital_cfg.get("nazwa", "Szpital Psychiatryczny im. Tylera Durdena"),
            level=1,
        )
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

        sub = doc.add_paragraph(szpital_cfg.get("adres", ""))
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.size = Pt(9)
        sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if szpital_cfg.get("oddzial"):
            od = doc.add_paragraph(szpital_cfg["oddzial"])
            od.alignment = WD_ALIGN_PARAGRAPH.CENTER
            od.runs[0].font.size = Pt(9)
            od.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # odstęp

        # ── Tytuł dokumentu ───────────────────────────────────────────────────
        tyt = doc.add_heading("HISTORIA CHOROBY — KARTA PRZYJĘCIA", level=2)
        tyt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in tyt.runs:
            run.font.size = Pt(12)

        nr = data.get("numer_historii_choroby", "NY-2026-00000")
        data_przyjecia = data.get("data_przyjecia", datetime.now().strftime("%d.%m.%Y"))
        nr_par = doc.add_paragraph(f"Nr: {nr}  |  Data: {data_przyjecia}")
        nr_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr_par.runs[0].font.size = Pt(9)
        nr_par.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        def sekcja(tytul_sek):
            p = doc.add_heading(tytul_sek.upper(), level=3)
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
            return p

        def pole(label, wartosc):
            if not wartosc:
                return
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_val = p.add_run(str(wartosc))
            run_val.font.size = Pt(10)

        def tekst_blok(zawartosc):
            if not zawartosc:
                return
            p = doc.add_paragraph(str(zawartosc))
            p.runs[0].font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)

        def lista_punktow(items):
            for item in items or []:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(str(item))
                run.font.size = Pt(10)

        # ── Dane pacjenta ─────────────────────────────────────────────────────
        sekcja("Dane Pacjenta")
        dp = data.get("dane_pacjenta", {})
        pole("Imię i nazwisko", dp.get("imie_nazwisko", ""))
        pole("Wiek", dp.get("wiek", ""))
        pole("Adres", dp.get("adres", ""))
        pole("Zawód", dp.get("zawod", ""))
        pole("Stan cywilny", dp.get("stan_cywilny", ""))
        doc.add_paragraph()

        # ── Powód przyjęcia ───────────────────────────────────────────────────
        sekcja("Powód Przyjęcia")
        tekst_blok(data.get("powod_przyjecia", ""))
        doc.add_paragraph()

        # ── Wywiad ────────────────────────────────────────────────────────────
        sekcja("Wywiad z Pacjentem")
        tekst_blok(data.get("wywiad", ""))
        doc.add_paragraph()

        # ── Objawy ────────────────────────────────────────────────────────────
        sekcja("Objawy")
        lista_punktow(data.get("objawy", []))
        doc.add_paragraph()

        # ── Diagnoza ──────────────────────────────────────────────────────────
        sekcja("Diagnoza")
        p_diag = doc.add_paragraph()
        run_diag = p_diag.add_run(data.get("diagnoza_wstepna", ""))
        run_diag.bold = True
        run_diag.font.size = Pt(11)
        run_diag.font.color.rgb = RGBColor(0x99, 0x1A, 0x1A)
        if data.get("diagnoza_dodatkowa"):
            p_dd = doc.add_paragraph()
            run_dd = p_dd.add_run(f"Diagnoza dodatkowa: {data['diagnoza_dodatkowa']}")
            run_dd.font.size = Pt(10)
            run_dd.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)
        doc.add_paragraph()

        # ── Zalecenia Terapeutyczne ─────────────────────────────────────────────────
        sekcja("Zalecenia Terapeutyczne")
        lista_punktow(data.get("zalecenia", []))
        doc.add_paragraph()

        # ── Rokowanie ─────────────────────────────────────────────────────────
        sekcja("Rokowanie")
        p_rok = doc.add_paragraph()
        run_rok = p_rok.add_run(data.get("rokowanie", ""))
        run_rok.font.size = Pt(10)
        run_rok.font.color.rgb = RGBColor(0x99, 0x1A, 0x1A)
        doc.add_paragraph()

        # ── Podpis ────────────────────────────────────────────────────────────
        p_podpis = doc.add_paragraph()
        run_podpis = p_podpis.add_run(data.get("podpis_lekarza", "Dr. T. Durden, MD"))
        run_podpis.font.size = Pt(10)
        run_podpis.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        if data.get("notatka_oddzialu"):
            p_not = doc.add_paragraph()
            run_not = p_not.add_run(f"Notatka pielęgniarki: {data['notatka_oddzialu']}")
            run_not.font.size = Pt(9)
            run_not.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # ── Zdjęcie psychiatryczne na końcu ───────────────────────────────────
        doc.add_paragraph()
        doc.add_page_break()

        photo_title = doc.add_heading("DOKUMENTACJA FOTOGRAFICZNA PACJENTA", level=2)
        photo_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in photo_title.runs:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)

        photo_sub = doc.add_paragraph("Zdjęcie wykonane przy przyjęciu — Oddział B")
        photo_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        photo_sub.runs[0].font.size = Pt(9)
        photo_sub.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()

        # Generuj zdjęcie przez FLUX
        photo_b64 = _generate_psychiatric_photo(
            body=body,
            nouns_dict=nouns_dict or {},
            sender_name=sender_name,
            test_mode=test_mode,
        )

        if photo_b64:
            try:
                photo_bytes = base64.b64decode(photo_b64)
                photo_stream = io.BytesIO(photo_bytes)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(photo_stream, width=Cm(12))
                logger.info("[raport] Zdjęcie wklejone do DOCX OK")
            except Exception as e:
                logger.warning("[raport] Błąd wklejania zdjęcia do DOCX: %s", e)
                p_no_img = doc.add_paragraph("[Zdjęcie niedostępne]")
                p_no_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            logger.warning("[raport] Brak zdjęcia psychiatrycznego — pomijam")
            p_no_img = doc.add_paragraph("[Zdjęcie niewygenertowane — błąd FLUX]")
            p_no_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ── Zapisz DOCX do BytesIO ────────────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        docx_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        logger.info("[raport] DOCX OK")
        return {
            "base64": docx_b64,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "filename": f"raport_psychiatryczny_{ts}.docx",
        }

    except Exception as e:
        logger.error("[raport] Błąd DOCX: %s", e)
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# PLAKAT SVG
# ═══════════════════════════════════════════════════════════════════════════════


def _build_plakat_svg(res_text: str, body: str) -> dict | None:
    """Generuje plakat motywacyjny SVG."""
    try:
        with open(PLAKAT_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[plakat] Brak JSON: %s", e)
        return None

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    user_msg = (
        f"Odpowiedź Tylera:\n{res_text[:MAX_DLUGOSC_EMAIL]}\n\nEmail:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"SCHEMAT JSON — użyj DOKŁADNIE tych kluczy na GÓRNYM POZIOMIE:\n{__import__('json').dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"Zwróć TYLKO czysty JSON. KLUCZ glowne_zdanie MUSI być na górnym poziomie — nie zagnieżdżaj w 'plakat'."
    )

    raw = call_deepseek(_js(system_msg), _ju(user_msg), MODEL_TYLER)
    if not raw:
        logger.warning("[plakat] Brak odpowiedzi od AI")
        return None

    logger.info("[plakat] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        data = _parse_json_safe(raw, "plakat")
        if data is None:
            raise ValueError("[plakat] JSON nienaprawialny")
        if not isinstance(data, dict):
            raise ValueError(
                f"[plakat] Oczekiwano dict, dostałem {type(data).__name__}"
            )
        if not data.get("glowne_zdanie") and isinstance(data.get("plakat"), dict):
            data.update(data.pop("plakat"))
            logger.info("[plakat] wyciągnięto dane z zagnieżdżonego 'plakat'")
        KEY_MAP_PLAKAT = {
            "zdanie": "glowne_zdanie",
            "main_sentence": "glowne_zdanie",
            "sentence": "glowne_zdanie",
            "tekst": "glowne_zdanie",
            "text": "glowne_zdanie",
            "cytat": "glowne_zdanie",
            "quote": "glowne_zdanie",
            "tresc": "glowne_zdanie",
            "subtitle": "podtytul",
            "podtytul": "podtytul",
            "background": "tlo_opis",
            "tlo": "tlo_opis",
            "color": "kolor_dominujacy",
            "kolor": "kolor_dominujacy",
            "keyword": "slowo_klucz",
            "slowo": "slowo_klucz",
        }
        for wrong, right in KEY_MAP_PLAKAT.items():
            if wrong in data and right not in data:
                data[right] = data.pop(wrong)
                logger.info("[plakat] znormalizowano '%s' → '%s'", wrong, right)
        if not data.get("glowne_zdanie"):
            logger.warning("[plakat] JSON bez glowne_zdanie — raw: %.200s", raw)
            return None
    except Exception as e:
        logger.warning("[plakat] Błąd JSON: %s | raw: %.200s", e, raw)
        return None

    glowne = data.get("glowne_zdanie", "Nie jesteś wyjątkowy.")
    podtytul = data.get("podtytul", "")
    autor = data.get("autor", "— Tyler Durden")
    kolor_tlo = data.get("kolor_dominujacy", "#0a0a0a")
    kolor_tekst = data.get("kolor_tekstu", "#ffffff")
    slowo = data.get("slowo_klucz", "PUSTKA").upper()

    def wrap_words(text, max_chars):
        """Zawija tekst na linie max_chars znaków."""
        words = text.split()
        lines_out, cur = [], ""
        for w in words:
            test = (cur + " " + w).strip()
            if len(test) <= max_chars:
                cur = test
            else:
                if cur:
                    lines_out.append(cur)
                cur = w
        if cur:
            lines_out.append(cur)
        return lines_out

    # Główne zdanie — max 26 znaków na linię, font 48
    lines = wrap_words(glowne, 26)
    line_height = 62
    text_start_y = 320 - (len(lines) * line_height) // 2
    text_lines_svg = ""
    for i, line in enumerate(lines):
        text_lines_svg += (
            f'<text x="420" y="{text_start_y + i * line_height}" '
            f'font-family="Georgia, serif" font-size="48" font-weight="bold" '
            f'fill="{kolor_tekst}" text-anchor="middle" dominant-baseline="middle" '
            f'letter-spacing="1">{line}</text>\n'
        )

    # Podtytuł — zawijany, max 40 znaków na linię, font 22
    sub_lines = wrap_words(podtytul, 40) if podtytul else []
    sub_start_y = text_start_y + len(lines) * line_height + 50
    sub_svg = ""
    for i, sl in enumerate(sub_lines):
        sub_svg += (
            f'<text x="420" y="{sub_start_y + i * 30}" '
            f'font-family="Georgia, serif" font-size="22" '
            f'fill="{kolor_tekst}" text-anchor="middle" opacity="0.75">{sl}</text>\n'
        )

    # Dodatkowy tekst — fragment odpowiedzi Tylera jako cytat pod spodem
    # Bierzemy pierwsze zdanie które nie jest nagłówkiem (max 120 znaków)
    extra_quote = ""
    for sentence in res_text.replace("\n", " ").split("."):
        s = sentence.strip()
        if len(s) > 30 and not s.startswith("#") and not s.startswith("—"):
            extra_quote = s[:120]
            break
    extra_lines = wrap_words(extra_quote, 45) if extra_quote else []
    extra_start_y = sub_start_y + max(len(sub_lines), 1) * 30 + 60
    extra_svg = ""
    for i, el in enumerate(extra_lines):
        extra_svg += (
            f'<text x="420" y="{extra_start_y + i * 26}" '
            f'font-family="Georgia, serif" font-size="18" font-style="italic" '
            f'fill="{kolor_tekst}" text-anchor="middle" opacity="0.55">{el}</text>\n'
        )

    autor_y = extra_start_y + max(len(extra_lines), 1) * 26 + 50
    linia_y = autor_y - 20
    svg = f"""<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="840" height="1188" viewBox="0 0 840 1188">
  <!-- Tło -->
  <rect width="840" height="1188" fill="{kolor_tlo}"/>

  <!-- Słowo klucz — watermark -->
  <text x="420" y="594" font-family="Arial Black, sans-serif" font-size="200" font-weight="bold"
        fill="{kolor_tekst}" text-anchor="middle" dominant-baseline="middle"
        opacity="0.04" transform="rotate(-30, 420, 594)">{slowo}</text>

  <!-- Linia górna dekoracyjna -->
  <rect x="60" y="80" width="720" height="3" fill="#8b0000"/>
  <rect x="60" y="87" width="720" height="1" fill="#8b0000" opacity="0.5"/>

  <!-- Główny tekst -->
  {text_lines_svg}

  <!-- Podtytuł (zawijany) -->
  {sub_svg}

  <!-- Cytat dodatkowy z odpowiedzi Tylera -->
  {extra_svg}

  <!-- Linia przed autorem -->
  <rect x="160" y="{linia_y}" width="520" height="1" fill="{kolor_tekst}" opacity="0.3"/>

  <!-- Autor -->
  <text x="420" y="{autor_y}"
        font-family="Georgia, serif" font-size="20" font-style="italic"
        fill="#8b0000" text-anchor="middle">{autor}</text>

  <!-- Linia dolna dekoracyjna -->
  <rect x="60" y="1100" width="720" height="1" fill="#8b0000" opacity="0.5"/>
  <rect x="60" y="1104" width="720" height="3" fill="#8b0000"/>

  <!-- Małe logo projektu -->
  <text x="420" y="1150" font-family="Arial, sans-serif" font-size="11"
        fill="{kolor_tekst}" text-anchor="middle" opacity="0.3" letter-spacing="3">PROJEKT TYLER DURDEN</text>
</svg>"""

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    logger.info("[plakat] OK")
    return _to_zip(svg.encode("utf-8"), f"plakat_{ts}.svg", f"plakat_{ts}.zip")


# ═══════════════════════════════════════════════════════════════════════════════
# DIAGRAM PRZEPŁYWU SVG
# ═══════════════════════════════════════════════════════════════════════════════


def _build_flow_diagram_svg(exec_logger) -> dict | None:
    """Generuje diagram przepływu pokazujący INPUT → API CALLS → SECTIONS."""
    try:
        # Pobierz dane z exec_logger metadata (ExecutionLogger, nie logging.Logger)
        metadata = getattr(exec_logger, "metadata", {}) or {}
        api_calls = metadata.get("api_calls", [])
        sections_completed = metadata.get("sections_completed", [])
        in_history = metadata.get("in_history", "nieznany")
        in_requiem = metadata.get("in_requiem", "nieznany")

        # Przygotuj dane do wizualizacji
        deepseek_count = sum(
            1 for call in api_calls if call.get("provider") == "deepseek"
        )
        total_tokens = sum(call.get("tokens", 0) for call in api_calls)

        sections_list = ", ".join(sections_completed) if sections_completed else "brak"

        svg = f"""<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">
  <!-- Tło -->
  <rect width="800" height="600" fill="#1a1a1a"/>

  <!-- Tytuł -->
  <text x="400" y="40" font-family="Arial, sans-serif" font-size="24" font-weight="bold"
        fill="#ffffff" text-anchor="middle">DIAGRAM PRZEPŁYWU</text>

  <!-- INPUT -->
  <rect x="50" y="100" width="120" height="60" fill="#4CAF50" rx="10"/>
  <text x="110" y="135" font-family="Arial, sans-serif" font-size="14" font-weight="bold"
        fill="#ffffff" text-anchor="middle">INPUT</text>

  <!-- Strzałka 1 -->
  <polygon points="180,130 200,125 200,135" fill="#ffffff"/>
  <line x1="170" y1="130" x2="200" y2="130" stroke="#ffffff" stroke-width="2"/>

  <!-- API CALLS -->
  <rect x="220" y="80" width="140" height="100" fill="#2196F3" rx="10"/>
  <text x="290" y="110" font-family="Arial, sans-serif" font-size="14" font-weight="bold"
        fill="#ffffff" text-anchor="middle">API CALLS</text>
  <text x="290" y="135" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff" text-anchor="middle">DeepSeek: {deepseek_count}</text>
  <text x="290" y="155" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff" text-anchor="middle">Tokens: {total_tokens}</text>

  <!-- Strzałka 2 -->
  <polygon points="370,130 390,125 390,135" fill="#ffffff"/>
  <line x1="360" y1="130" x2="390" y2="130" stroke="#ffffff" stroke-width="2"/>

  <!-- SECTIONS -->
  <rect x="410" y="100" width="140" height="60" fill="#FF9800" rx="10"/>
  <text x="480" y="125" font-family="Arial, sans-serif" font-size="14" font-weight="bold"
        fill="#ffffff" text-anchor="middle">SECTIONS</text>
  <text x="480" y="140" font-family="Arial, sans-serif" font-size="10"
        fill="#ffffff" text-anchor="middle">{sections_list[:20]}</text>

  <!-- Status użytkownika -->
  <rect x="580" y="80" width="160" height="100" fill="#9C27B0" rx="10"/>
  <text x="660" y="105" font-family="Arial, sans-serif" font-size="14" font-weight="bold"
        fill="#ffffff" text-anchor="middle">STATUS</text>
  <text x="660" y="125" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff" text-anchor="middle">Historia: {in_history}</text>
  <text x="660" y="140" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff" text-anchor="middle">Requiem: {in_requiem}</text>

  <!-- Szczegóły API calls -->
  <text x="50" y="220" font-family="Arial, sans-serif" font-size="16" font-weight="bold"
        fill="#ffffff">SZCZEGÓŁY API CALLS:</text>

  <text x="50" y="250" font-family="Arial, sans-serif" font-size="12"
        fill="#cccccc">• Łącznie wywołań: {len(api_calls)}</text>
  <text x="50" y="270" font-family="Arial, sans-serif" font-size="12"
        fill="#cccccc">• Sekcje wykonane: {len(sections_completed)}</text>
  <text x="50" y="290" font-family="Arial, sans-serif" font-size="12"
        fill="#cccccc">• Czas przetwarzania: ~{len(api_calls) * 2}s</text>

  <!-- Legenda -->
  <text x="50" y="340" font-family="Arial, sans-serif" font-size="14" font-weight="bold"
        fill="#ffffff">LEGENDA:</text>

  <rect x="50" y="360" width="15" height="15" fill="#4CAF50"/>
  <text x="75" y="372" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff">Wejście email</text>

  <rect x="50" y="385" width="15" height="15" fill="#2196F3"/>
  <text x="75" y="397" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff">Wywołania AI</text>

  <rect x="50" y="410" width="15" height="15" fill="#FF9800"/>
  <text x="75" y="422" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff">Generowane sekcje</text>

  <rect x="50" y="435" width="15" height="15" fill="#9C27B0"/>
  <text x="75" y="447" font-family="Arial, sans-serif" font-size="12"
        fill="#ffffff">Status użytkownika</text>

  <!-- Stopka -->
  <text x="400" y="580" font-family="Arial, sans-serif" font-size="10"
        fill="#666666" text-anchor="middle">Wygenerowano: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</text>
</svg>"""

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        logger.info("[flow_diagram] OK")
        return _to_zip(
            svg.encode("utf-8"), f"flow_diagram_{ts}.svg", f"flow_diagram_{ts}.zip"
        )

    except Exception as e:
        logger.warning("[flow_diagram] Błąd: %s", e)
        return None


# ═══════════════════════════════════════════════════════════════════════════════
# GRA HTML
# ═══════════════════════════════════════════════════════════════════════════════


def _build_gra_html(body: str, res_text: str) -> dict | None:
    """Generuje grę interaktywną HTML z wyborami Tylera."""
    try:
        with open(GRA_JSON_PATH, encoding="utf-8") as f:
            cfg = json.load(f)
    except Exception as e:
        logger.warning("[gra] Brak JSON: %s", e)
        return None

    system_msg = cfg.get("system", "")
    schema = cfg.get("output_schema", {})
    user_msg = (
        f"Email:\n{body[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"Odpowiedź Tylera:\n{res_text[:MAX_DLUGOSC_EMAIL]}\n\n"
        f"SCHEMAT JSON — użyj DOKŁADNIE tych kluczy:\n{__import__('json').dumps(schema, ensure_ascii=False, indent=2)}\n\n"
        f"Zwróć TYLKO czysty JSON. Klucz listy pytań MUSI być 'pytania'."
    )

    # max_tokens=4000 — zwiększone, 10 pytań × ~200 tokenów = min 3500 potrzebnych
    raw = call_deepseek(_js(system_msg), _ju(user_msg), MODEL_TYLER, max_tokens=4000)
    if not raw:
        logger.warning("[gra] Brak odpowiedzi od AI")
        return None

    logger.info("[gra] raw AI (pierwsze 300 znaków): %.300s", raw)

    try:
        clean = _strip_json_markdown(raw)
        # Naprawa: jeśli zaczyna się od przecinka, dodaj { na początku
        if clean.startswith(","):
            clean = "{" + clean
        # Naprawa: jeśli brakuje końcowego }, dodaj go
        if clean.count("{") > clean.count("}") and not clean.endswith("}"):
            clean += "}"
        if clean.count("[") > clean.count("]") and not clean.endswith("]"):
            clean += "]"

        # Użyj raw_decode zamiast json.loads — obsługuje "Extra data"
        # (gdy AI zwróci JSON + dodatkowy tekst poza klamrami)
        decoder = json.JSONDecoder()
        data = None
        try:
            data, _ = decoder.raw_decode(clean)
        except json.JSONDecodeError:
            # Fallback: szukaj największego JSON w tekście
            for match in re.finditer(r"[\[{]", clean):
                start = match.start()
                try:
                    obj, end = decoder.raw_decode(clean[start:])
                    if obj is not None:
                        data = obj
                        break
                except json.JSONDecodeError:
                    continue
        if data is None:
            raise ValueError(f"[gra] Nie znaleziono JSON w odpowiedzi")
        # AI zwróciło listę pytań bez wrappera — owijamy
        if isinstance(data, list):
            if len(data) > 0 and isinstance(data[0], dict):
                logger.warning("[gra] AI zwróciło listę — owijam w {pytania: [...]}")
                data = {"pytania": data}
            else:
                raise ValueError(f"[gra] Oczekiwano dict, dostałem list")
        if not isinstance(data, dict):
            raise ValueError(f"[gra] Oczekiwano dict, dostałem {type(data).__name__}")
        if not data.get("pytania"):
            # Szukaj pytań pod alternatywnymi kluczami
            KEY_MAP_GRA = {
                "questions": "pytania",
                "tasks": "pytania",
                "scenarios": "pytania",
                "choices": "pytania",
                "quests": "pytania",
            }
            for wrong, right in KEY_MAP_GRA.items():
                if wrong in data:
                    data[right] = data.pop(wrong)
                    logger.info("[gra] znormalizowano '%s' → '%s'", wrong, right)
                    break
        if not data.get("pytania"):
            # Szukaj zagnieżdżonej listy dictów
            for v in data.values():
                if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                    data["pytania"] = v
                    logger.info("[gra] wyciągnięto pytania z zagnieżdżonej listy")
                    break
        if not data.get("pytania"):
            logger.warning("[gra] JSON OK ale brak pytań — raw: %.200s", raw)
            return None
    except Exception as e:
        logger.warning("[gra] Błąd JSON: %s | raw: %.200s", e, raw)
        return None

    tytul = data.get("tytul_gry", "Gra Tylera Durdena")
    wstep = data.get("wstep", "")
    pytania = data.get("pytania", [])
    wyniki = data.get("wyniki", {})
    zakonczenie = data.get("zakonczenie", "— Tyler Durden")

    # Buduj komentarze JS
    komentarze_b = {}
    komentarze_inne = {}
    for p in pytania:
        nr = p.get("nr", 0)
        komentarze_b[nr] = p.get("komentarz_po_wyborze_b", "Dobrze.")
        komentarze_inne[nr] = p.get("komentarz_po_wyborze_innym", "Typowe.")

    kb_js = json.dumps(komentarze_b)
    ki_js = json.dumps(komentarze_inne)
    w_js = json.dumps(wyniki)

    pytania_html = ""
    for p in pytania:
        nr = p.get("nr", "?")
        sytuacja = p.get("sytuacja", "")
        pytanie_txt = p.get("pytanie", "")
        odp = p.get("odpowiedzi", {})
        if isinstance(odp, list):
            odp = {
                str(item.get("klucz", item.get("key", chr(97 + i)))): str(
                    item.get("tresc", item.get("text", ""))
                )
                for i, item in enumerate(odp)
            }
        elif not isinstance(odp, dict):
            odp = {}
        pytania_html += f"""
<div class="pytanie" id="p{nr}" style="display:none">
  <div class="nr">Pytanie {nr} / {len(pytania)}</div>
  <div class="sytuacja">{sytuacja}</div>
  <div class="pytanie-txt">{pytanie_txt}</div>
  <div class="opcje">
    <button class="opcja" onclick="odpowiedz({nr},'a')">a) {odp.get('a', '')}</button>
    <button class="opcja" onclick="odpowiedz({nr},'b')">b) {odp.get('b', '')}</button>
    <button class="opcja" onclick="odpowiedz({nr},'c')">c) {odp.get('c', '')}</button>
  </div>
  <div class="komentarz" id="k{nr}"></div>
</div>"""

    html = f"""<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="UTF-8">
<title>{tytul}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Courier New', monospace; background: #050505; color: #d0c0a0; min-height: 100vh; display: flex; flex-direction: column; align-items: center; padding: 20px; }}
  h1 {{ color: #8b0000; text-align: center; font-size: 1.6em; margin: 30px 0 10px; border-bottom: 2px solid #8b0000; padding-bottom: 10px; width: 100%; max-width: 700px; }}
  .wstep {{ color: #666; font-style: italic; text-align: center; margin: 15px 0 30px; max-width: 600px; }}
  .pytanie {{ background: #0f0f0f; border: 1px solid #2a1a1a; border-left: 4px solid #8b0000; padding: 25px; max-width: 700px; width: 100%; border-radius: 0 4px 4px 0; }}
  .nr {{ color: #8b0000; font-size: 0.8em; margin-bottom: 10px; letter-spacing: 2px; }}
  .sytuacja {{ color: #888; font-style: italic; margin-bottom: 12px; font-size: 0.9em; line-height: 1.5; }}
  .pytanie-txt {{ font-size: 1.1em; color: #c8b89a; margin-bottom: 20px; font-weight: bold; }}
  .opcje {{ display: flex; flex-direction: column; gap: 10px; }}
  .opcja {{ background: #1a1a1a; border: 1px solid #333; color: #c8b89a; padding: 12px 18px; text-align: left; cursor: pointer; font-family: 'Courier New', monospace; font-size: 0.9em; transition: all 0.2s; border-radius: 2px; }}
  .opcja:hover {{ background: #2a1a1a; border-color: #8b0000; }}
  .opcja:disabled {{ opacity: 0.5; cursor: not-allowed; }}
  .opcja.tyler {{ background: #1a0a0a; border-color: #8b0000; color: #ff6666; }}
  .komentarz {{ margin-top: 15px; padding: 10px; background: #0a0a0a; border-left: 2px solid #8b0000; color: #8b0000; font-style: italic; display: none; font-size: 0.85em; }}
  #wynik {{ display: none; background: #0f0f0f; border: 2px solid #8b0000; padding: 30px; max-width: 700px; width: 100%; text-align: center; margin-top: 20px; }}
  #wynik h2 {{ color: #8b0000; margin-bottom: 15px; }}
  #wynik .punkty {{ font-size: 2em; color: #c8b89a; margin: 10px 0; }}
  #wynik .komentarz-wynik {{ color: #888; font-style: italic; }}
  #start {{ background: #8b0000; color: white; border: none; padding: 15px 40px; font-size: 1.1em; cursor: pointer; font-family: 'Courier New', monospace; margin: 20px 0; letter-spacing: 1px; }}
  #start:hover {{ background: #a00000; }}
  .pasek {{ background: #1a1a1a; height: 4px; max-width: 700px; width: 100%; margin: 10px 0; }}
  .pasek-fill {{ background: #8b0000; height: 100%; transition: width 0.3s; width: 0%; }}
  footer {{ color: #333; font-size: 0.75em; margin-top: 40px; text-align: center; }}
</style>
</head>
<body>
<h1>{tytul}</h1>
<p class="wstep">{wstep}</p>
<div class="pasek"><div class="pasek-fill" id="pasek"></div></div>
<button id="start" onclick="startGra()">ROZPOCZNIJ GRĘ</button>
{pytania_html}
<div id="wynik">
  <h2>KONIEC GRY</h2>
  <div class="punkty" id="punkty-wynik"></div>
  <div class="komentarz-wynik" id="komentarz-wynik"></div>
</div>
<footer>{zakonczenie}</footer>
<script>
var bieżace = 0;
var punkty = 0;
var total = {len(pytania)};
var kb = {kb_js};
var ki = {ki_js};
var wyniki = {w_js};

function startGra() {{
  document.getElementById('start').style.display = 'none';
  pokazPytanie(1);
}}

function pokazPytanie(nr) {{
  bieżace = nr;
  var el = document.getElementById('p' + nr);
  if (el) el.style.display = 'block';
  document.getElementById('pasek').style.width = ((nr-1)/total*100) + '%';
  window.scrollTo(0, document.body.scrollHeight);
}}

function odpowiedz(nr, wybor) {{
  var btns = document.querySelectorAll('#p' + nr + ' .opcja');
  btns.forEach(function(b) {{ b.disabled = true; }});
  
  if (wybor === 'b') {{
    punkty++;
    btns[1].classList.add('tyler');
    var k = document.getElementById('k' + nr);
    k.innerHTML = '— ' + (kb[nr] || 'Dobrze.');
    k.style.display = 'block';
  }} else {{
    var k = document.getElementById('k' + nr);
    k.innerHTML = '— ' + (ki[nr] || 'Typowe.');
    k.style.display = 'block';
    k.style.borderLeftColor = '#444';
    k.style.color = '#555';
  }}

  setTimeout(function() {{
    if (nr < total) {{
      pokazPytanie(nr + 1);
    }} else {{
      pokazWynik();
    }}
  }}, 1800);
}}

function pokazWynik() {{
  document.getElementById('pasek').style.width = '100%';
  var wynikDiv = document.getElementById('wynik');
  wynikDiv.style.display = 'block';
  document.getElementById('punkty-wynik').innerHTML = punkty + ' / ' + total + ' punktów Tylera';
  var komentarz = '';
  if (punkty <= 3) komentarz = wyniki['0_3'] || 'Rozczarowujące.';
  else if (punkty <= 6) komentarz = wyniki['4_6'] || 'Trochę lepiej.';
  else if (punkty <= 9) komentarz = wyniki['7_9'] || 'Prawie.';
  else komentarz = wyniki['10'] || 'Jesteś gotowy.';
  document.getElementById('komentarz-wynik').innerHTML = komentarz;
  window.scrollTo(0, document.body.scrollHeight);
}}
</script>
</body>
</html>"""

    html_b64 = base64.b64encode(html.encode("utf-8")).decode("ascii")
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    logger.info("[gra] OK: %d pytań", len(pytania))
    return _to_zip(html.encode("utf-8"), f"gra_{ts}.html", f"gra_{ts}.zip")


# ═══════════════════════════════════════════════════════════════════════════════
# GŁÓWNA FUNKCJA RESPONDERA
# ═══════════════════════════════════════════════════════════════════════════════


def build_zwykly_section(
    body: str,
    previous_body: str = "",
    sender_email: str = "",
    sender_name: str = "",
    test_mode: bool = False,
    attachments: list = None,
) -> dict:
    """
    Główna funkcja respondera ZWYKLY.
    Zwraca dict z reply_html + załącznikami oraz meta pola do historii/Drive.
    """
    body = body or ""
    previous_body = previous_body or ""
    sender_email = sender_email or ""
    sender_name = sender_name or ""
    prompt_data = _load_prompt_json()
    user_msg = _render_prompt(prompt_data, body, previous_body, sender_name)
    system_msg = prompt_data.get("system", "")

    raw, provider = _call_ai_with_fallback(
        _js(system_msg), _ju(user_msg), max_tokens=6500
    )
    if not raw:
        logger.warning("[zwykly] Brak odpowiedzi AI")
        return {
            "reply_html": "<p>Przepraszam, nie udało się wygenerować odpowiedzi.</p>",
            "triptych": [],
            "images": [],
            "docs": [],
        }

    res_text = ""
    emotion_key = ""
    pdf_category = ""
    try:
        data = _parse_json_safe(raw, "zwykly-build")
        if isinstance(data, dict):
            res_text = str(
                data.get("odpowiedz_tekstowa")
                or data.get("odpowiedz")
                or data.get("tekst")
                or raw
            )
            emotion_key = str(data.get("emocja") or "")
            pdf_category = str(data.get("kategoria_pdf") or "")
        else:
            logger.warning(
                "[zwykly] Oczekiwano dict z AI, otrzymano %s — próba regex fallback",
                type(data).__name__ if data is not None else "None",
            )
            # Fallback: wyciągnij odpowiedz_tekstowa bezpośrednio z ucietego JSON-a
            m = re.search(
                r'"odpowiedz_tekstowa"\s*:\s*"((?:[^"\]|\.)*)',
                raw,
                re.DOTALL,
            )
            if m:
                extracted = m.group(1)
                # Odkoduj escape sequences (\n, \t itp.)
                try:
                    extracted = extracted.encode("utf-8").decode("unicode_escape")
                except Exception:
                    pass
                res_text = extracted
                logger.warning(
                    "[zwykly] regex fallback OK — wyciągnięto %d znaków", len(res_text)
                )
                # Spróbuj też wyciągnąć emocję i kategorię
                m_em = re.search(r'"emocja"\s*:\s*"([^"]*)"', raw)
                if m_em:
                    emotion_key = m_em.group(1)
                m_cat = re.search(r'"kategoria_pdf"\s*:\s*"([^"]*)"', raw)
                if m_cat:
                    pdf_category = m_cat.group(1)
            else:
                res_text = raw
    except Exception as e:
        logger.warning("[zwykly] Błąd parsowania JSON: %s | raw: %.200s", e, raw)
        res_text = raw

    main_section_html = _wrap_section_html(
        build_html_reply(res_text),
        title="Tyler Durden + Sokrates",
    )

    nouns = _extract_nouns_from_body(body)
    nouns_dict = {
        f"rzecz{str(i + 1).zfill(3)}": noun for i, noun in enumerate(nouns[:15])
    }

    session_vars = _build_session_vars(
        body,
        sender_email,
        sender_name,
        previous_body,
        res_text,
        emotion_key,
        provider,
        panel_assignments=[],
        nouns_dict=nouns_dict,
    )

    triptych_images, triptych_prompts, panel_assignments = _generate_triptych(
        res_text,
        prompt_data,
        body,
        session_vars=session_vars,
        test_mode=test_mode,
    )

    emoticon_b64 = _generate_icon_flux(emotion_key, sender_name)
    emoticon = None
    if emoticon_b64:
        emoticon = {
            "base64": emoticon_b64,
            "content_type": "image/png",
            "filename": f"emocja_{emotion_key or 'default'}.png",
        }

    cv_pdf = None
    cv_data = _generate_cv_content(body, previous_body, sender_email, sender_name)
    if cv_data:
        cv_photo = _generate_cv_photo(body, cv_data, test_mode=test_mode)
        cv_pdf = _build_cv_pdf(cv_data, cv_photo)

    ankieta_html, ankieta_pdf = _build_ankieta(res_text, body)
    horoskop_pdf = _build_horoskop(body, res_text)
    karta_rpg_pdf = _build_karta_rpg(body, res_text)

    raport_pdf = None
    psych_photo_1 = None
    psych_photo_2 = None
    try:
        raport_result = build_raport(
            body,
            previous_body,
            res_text,
            nouns_dict,
            sender_name=sender_name,
            gender=_detect_gender(body, sender_name),
            test_mode=test_mode,
        )
        if isinstance(raport_result, dict):
            raport_pdf = raport_result.get("raport_pdf")
            psych_photo_1 = raport_result.get("psych_photo_1")
            psych_photo_2 = raport_result.get("psych_photo_2")
        else:
            logger.warning(
                "[zwykly] build_raport zwrócił %s zamiast dict",
                type(raport_result).__name__,
            )
    except Exception as e:
        logger.warning("[zwykly] Błąd raportu psychiatrycznego: %s", e)

    plakat_svg = _build_plakat_svg(res_text, body)
    gra_html = _build_gra_html(body, res_text)
    explanation_txt = _build_explanation_txt(res_text, body)
    debug_txt = _build_debug_txt(
        body,
        provider,
        emotion_key,
        raw,
        res_text,
        triptych_images or [],
        triptych_prompts or [],
        system_msg,
        user_msg,
        session_vars,
        panel_assignments or [],
    )

    docs: list[dict] = []
    images: list[dict] = []
    docx_list: list[dict] = []
    emocje_section_html = ""
    dociekliwy_section_html = ""
    scrabble_section_html = ""

    try:
        from responders.emocje import build_emocje_section

        emocje_output = build_emocje_section(
            body=body,
            sender_name=sender_name,
            sender_email=sender_email,
            attachments=attachments,
            test_mode=test_mode,
        )
        if isinstance(emocje_output, dict):
            emocje_section_html = _wrap_section_html(
                emocje_output.get("reply_html", ""),
                title="Emocje",
            )
            _collect_section_attachments(emocje_output, docs, docx_list, images)
    except Exception as e:
        logger.warning("[zwykly] Błąd emocje: %s", e)

    try:
        from responders.dociekliwy import build_dociekliwy_section

        dociekliwy_output = build_dociekliwy_section(
            body=body,
            attachments=attachments,
            sender_email=sender_email,
            sender_name=sender_name,
            test_mode=test_mode,
        )
        if isinstance(dociekliwy_output, dict):
            dociekliwy_section_html = _wrap_section_html(
                dociekliwy_output.get("reply_html", ""),
                title="Dociekliwy",
            )
            _collect_section_attachments(dociekliwy_output, docs, docx_list, images)
    except Exception as e:
        logger.warning("[zwykly] Błąd dociekliwy: %s", e)

    try:
        from responders.scrabble import build_scrabble_section

        scrabble_output = build_scrabble_section(body)
        if isinstance(scrabble_output, dict):
            scrabble_section_html = _wrap_section_html(
                scrabble_output.get("reply_html", ""),
                title="Scrabble",
            )
            _collect_section_attachments(scrabble_output, docs, docx_list, images)
    except Exception as e:
        logger.warning("[zwykly] Błąd scrabble: %s", e)

    reply_html = _render_body_sections(
        main_section_html,
        emocje_section_html,
        dociekliwy_section_html,
        scrabble_section_html,
    )

    result = {
        "reply_html": reply_html,
        "triptych": triptych_images or [],
        "triptych_for_drive": triptych_images or [],
    }
    if emoticon:
        result["emoticon"] = emoticon
    if cv_pdf:
        result["cv_pdf"] = cv_pdf
    if ankieta_html:
        result["ankieta_html"] = ankieta_html
    if ankieta_pdf:
        result["ankieta_pdf"] = ankieta_pdf
    if horoskop_pdf:
        result["horoskop_pdf"] = horoskop_pdf
    if karta_rpg_pdf:
        result["karta_rpg_pdf"] = karta_rpg_pdf
    if raport_pdf:
        result["raport_pdf"] = raport_pdf
    if psych_photo_1:
        result["psych_photo_1"] = psych_photo_1
    if psych_photo_2:
        result["psych_photo_2"] = psych_photo_2
    if plakat_svg:
        result["plakat_svg"] = plakat_svg
    if gra_html:
        result["gra_html"] = gra_html
    if explanation_txt:
        result["explanation_txt"] = explanation_txt
    if debug_txt:
        result["debug_txt"] = debug_txt
    if images:
        result["images"] = images
    else:
        result["images"] = []
    if docs:
        result["docs"] = docs
    else:
        result["docs"] = []
    if docx_list:
        result["docx_list"] = docx_list
    return result
