"""
responders/analiza.py
Responder KEYWORDS3 — analiza powtórzeń tekstu.

Obsługuje:
- Wiele załączników DOCX / ODT / PDF
- Treść maila (body) — podświetlona w HTML odpowiedzi
- Każdy DOCX/ODT → przetworzony i odesłany jako załącznik
- PDF → tekst → DOCX z podświetleniami → załącznik
- Stopka pogrubiona na końcu treści HTML
"""
import os
import io
import re
import base64
from collections import defaultdict
from flask import current_app

try:
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import RGBColor
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

# ── Paleta kolorów (identyczna jak w 1.py) ────────────────────────────────────
_COLOR_DATA = [
    (WD_COLOR_INDEX.YELLOW,       (255, 255,   0), False),
    (WD_COLOR_INDEX.BRIGHT_GREEN, (  0, 255,   0), False),
    (WD_COLOR_INDEX.TURQUOISE,    (  0, 255, 255), False),
    (WD_COLOR_INDEX.PINK,         (255,   0, 255), False),
    (WD_COLOR_INDEX.BLUE,         (  0,   0, 255), True),
    (WD_COLOR_INDEX.RED,          (255,   0,   0), True),
    (WD_COLOR_INDEX.DARK_YELLOW,  (128, 128,   0), True),
    (WD_COLOR_INDEX.TEAL,         (  0, 128, 128), True),
    (WD_COLOR_INDEX.VIOLET,       (128,   0, 128), True),
    (WD_COLOR_INDEX.GREEN,        (  0, 128,   0), True),
    (WD_COLOR_INDEX.DARK_BLUE,    (  0,   0, 128), True),
    (WD_COLOR_INDEX.DARK_RED,     (128,   0,   0), True),
    (WD_COLOR_INDEX.GRAY_50,      (128, 128, 128), True),
    (WD_COLOR_INDEX.GRAY_25,      (192, 192, 192), False),
]
_PALETTE = (_COLOR_DATA * 11)[:150]

# Mapowanie koloru Word na CSS dla HTML
_COLOR_CSS = {
    WD_COLOR_INDEX.YELLOW:       ("#ffff00", "#000"),
    WD_COLOR_INDEX.BRIGHT_GREEN: ("#00ff00", "#000"),
    WD_COLOR_INDEX.TURQUOISE:    ("#00ffff", "#000"),
    WD_COLOR_INDEX.PINK:         ("#ff00ff", "#000"),
    WD_COLOR_INDEX.BLUE:         ("#0000ff", "#fff"),
    WD_COLOR_INDEX.RED:          ("#ff0000", "#fff"),
    WD_COLOR_INDEX.DARK_YELLOW:  ("#808000", "#fff"),
    WD_COLOR_INDEX.TEAL:         ("#008080", "#fff"),
    WD_COLOR_INDEX.VIOLET:       ("#800080", "#fff"),
    WD_COLOR_INDEX.GREEN:        ("#008000", "#fff"),
    WD_COLOR_INDEX.DARK_BLUE:    ("#000080", "#fff"),
    WD_COLOR_INDEX.DARK_RED:     ("#800000", "#fff"),
    WD_COLOR_INDEX.GRAY_50:      ("#808080", "#fff"),
    WD_COLOR_INDEX.GRAY_25:      ("#c0c0c0", "#000"),
}

X_VAL = 2200  # domyślny zasięg powtórzeń w znakach


# ── Stemizacja ────────────────────────────────────────────────────────────────
def _get_smart_root(word: str) -> str:
    w = word.lower()
    if len(w) < 3:
        return w
    if "śmie" in w:
        if "śmier" in w: return "śmierć"
        if "śmiet" in w: return "śmietana"
        return "śmiech/uśmiech"
    suffixes = [
        'ego', 'emu', 'ach', 'ami', 'ych', 'ich', 'owi', 'om',
        'em', 'am', 'ia', 'ie', 'y', 'a', 'u', 'i',
        'ł', 'ła', 'li', 'ły', 'cie', 'sz'
    ]
    stem = w
    for s in sorted(suffixes, key=len, reverse=True):
        if stem.endswith(s) and len(stem) - len(s) >= 3:
            stem = stem[:-len(s)]
            break
    return stem


# ── Mapa podświetleń ──────────────────────────────────────────────────────────
def _build_highlight_map(full_text: str, x_val: int = X_VAL) -> dict:
    """Zwraca {pozycja: kolor_info} dla powtarzających się słów."""
    matches = list(re.finditer(r'\b\w+\b', full_text))
    groups  = defaultdict(list)
    for m in matches:
        w = m.group().lower()
        if len(w) > 2:
            groups[_get_smart_root(w)].append(m)

    highlight_map = {}
    color_idx = 0

    for root, m_list in groups.items():
        if len(m_list) < 2:
            continue
        group_has_rep = False
        for i, m1 in enumerate(m_list):
            for j, m2 in enumerate(m_list):
                if i != j and abs(m1.start() - m2.start()) <= x_val:
                    group_has_rep = True
                    if m2.start() not in highlight_map:
                        highlight_map[m2.start()] = _PALETTE[color_idx % len(_PALETTE)]
        if group_has_rep:
            if m_list[0].start() not in highlight_map:
                highlight_map[m_list[0].start()] = _PALETTE[color_idx % len(_PALETTE)]
            color_idx += 1

    return highlight_map


# ── Nakładanie podświetleń na DOCX ────────────────────────────────────────────
def _apply_highlights_to_doc(doc, highlight_map: dict):
    full_text = "".join([p.text + " \n " for p in doc.paragraphs])
    for para in doc.paragraphs:
        txt = para.text
        if not txt:
            continue
        try:
            para_start = full_text.index(txt)
        except ValueError:
            continue
        tokens = re.split(r'(\b\w+\b)', txt)
        para.clear()
        offset = 0
        for t in tokens:
            run = para.add_run(t)
            pos = para_start + offset
            if pos in highlight_map:
                color_info = highlight_map[pos]
                run.font.highlight_color = color_info[0]
                if color_info[2]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
            offset += len(t)
    return doc


# ── Podświetlenie treści maila jako HTML ──────────────────────────────────────
def _highlight_text_to_html(text: str) -> str:
    """
    Analizuje tekst i zwraca HTML z podświetlonymi powtórzeniami
    (kolory CSS odpowiadające kolorom Word).
    """
    highlight_map = _build_highlight_map(text, X_VAL)
    if not highlight_map:
        # Brak powtórzeń — zwróć czysty tekst
        return "<p>" + text.replace("\n", "<br>") + "</p>"

    tokens = re.split(r'(\b\w+\b)', text)
    html   = "<p>"
    offset = 0

    for t in tokens:
        pos = offset
        if pos in highlight_map:
            color_info = highlight_map[pos]
            bg, fg     = _COLOR_CSS.get(color_info[0], ("#ffff00", "#000"))
            safe_t     = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html += f'<span style="background-color:{bg};color:{fg};">{safe_t}</span>'
        else:
            safe_t = t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html  += safe_t.replace("\n", "<br>")
        offset += len(t)

    html += "</p>"
    return html


# ── Pomocnicze: tekst → DOCX ──────────────────────────────────────────────────
def _text_to_doc(text: str):
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    return doc


def _doc_to_base64(doc) -> str:
    buf = io.BytesIO()
    doc.save(buf)
    return base64.b64encode(buf.getvalue()).decode("ascii")


# ── Odczyt tekstu z PDF ───────────────────────────────────────────────────────
def _pdf_bytes_to_text(pdf_bytes: bytes) -> str:
    """Wyciąga tekst z PDF. Próbuje pdfplumber, fallback pypdf."""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        if text.strip():
            return text
    except Exception as e:
        current_app.logger.warning("pdfplumber failed: %s", e)

    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(pdf_bytes))
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text
    except Exception as e:
        current_app.logger.warning("pypdf failed: %s", e)

    return ""


# ── Przetwórz jeden załącznik ─────────────────────────────────────────────────
def _process_attachment(att_b64: str, att_name: str) -> dict | None:
    """
    Przetwarza jeden załącznik (DOCX/ODT/PDF).
    Zwraca dict z base64 przetworzonego DOCX lub None przy błędzie.
    """
    if not _HAS_DOCX:
        return None

    name_lower = (att_name or "").lower()
    try:
        raw_bytes = base64.b64decode(att_b64)
    except Exception as e:
        current_app.logger.warning("Błąd dekodowania załącznika %s: %s", att_name, e)
        return None

    try:
        # PDF → tekst → DOCX
        if name_lower.endswith(".pdf"):
            text = _pdf_bytes_to_text(raw_bytes)
            if not text.strip():
                current_app.logger.warning("PDF %s — brak tekstu", att_name)
                return None
            doc = _text_to_doc(text)
        else:
            # DOCX / ODT
            doc = Document(io.BytesIO(raw_bytes))

        full_text     = "".join([p.text + " \n " for p in doc.paragraphs])
        highlight_map = _build_highlight_map(full_text, X_VAL)
        doc           = _apply_highlights_to_doc(doc, highlight_map)

        # Nazwa wynikowego pliku
        base_name = os.path.splitext(att_name)[0] if att_name else "analiza"
        out_name  = f"{base_name}_analiza.docx"

        return {
            "base64":   _doc_to_base64(doc),
            "filename": out_name,
        }
    except Exception as e:
        current_app.logger.exception("Błąd przetwarzania załącznika %s: %s", att_name, e)
        return None


# ── Główna funkcja responderu ─────────────────────────────────────────────────
def build_analiza_section(body: str,
                           attachments: list = None) -> dict:
    """
    Buduje sekcję 'analiza':
    - Podświetla treść maila (body) w HTML
    - Przetwarza wszystkie załączniki (DOCX/ODT/PDF)
    - Zwraca HTML z podświetloną treścią + listę DOCX załączników

    attachments — lista dict: [{"base64": ..., "name": ...}, ...]
    """
    if not _HAS_DOCX:
        current_app.logger.error("python-docx nie jest zainstalowane!")
        return {
            "reply_html": "<p>Błąd serwera: brak biblioteki python-docx.</p>",
            "docx_list":  [],
        }

    # ── 1. Podświetl treść maila w HTML ──────────────────────────────────────
    body_html    = _highlight_text_to_html(body) if body and body.strip() else ""

    # ── 2. Przetwórz załączniki ───────────────────────────────────────────────
    docx_list    = []
    sources      = []

    if attachments:
        for att in attachments:
            att_b64  = att.get("base64")
            att_name = att.get("name", "dokument")
            if not att_b64:
                continue
            result = _process_attachment(att_b64, att_name)
            if result:
                docx_list.append(result)
                sources.append(att_name)
                current_app.logger.info("Przetworzono załącznik: %s", att_name)

    # Jeśli brak załączników — utwórz DOCX z treści maila
    if not docx_list and body and body.strip():
        try:
            doc           = _text_to_doc(body)
            full_text     = "".join([p.text + " \n " for p in doc.paragraphs])
            highlight_map = _build_highlight_map(full_text, X_VAL)
            doc           = _apply_highlights_to_doc(doc, highlight_map)
            docx_list.append({
                "base64":   _doc_to_base64(doc),
                "filename": "analiza_tresci_maila.docx",
            })
            sources.append("treść maila")
        except Exception as e:
            current_app.logger.exception("Błąd tworzenia DOCX z treści maila: %s", e)

    # ── 3. Zbuduj HTML odpowiedzi ─────────────────────────────────────────────
    source_str = ", ".join(sources) if sources else "treść maila"
    stopka     = (
        f"<p><strong>Analiza powtórzeń wykonana na podstawie: {source_str}. "
        f"Wynik w załączniku DOCX.</strong></p>"
    )

    reply_html = ""
    if body_html:
        reply_html += body_html
    reply_html += stopka

    current_app.logger.info(
        "Analiza: źródła=%s | załączniki DOCX=%d",
        source_str, len(docx_list)
    )

    return {
        "reply_html": reply_html,
        "docx_list":  docx_list,
    }
